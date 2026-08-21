#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Программа для автоматического заполнения карт Excel (заключений эксперта)
на основе данных из постановлений (Word-документы .docx или .doc).
"""

import os
import re
import shutil
import logging
import sys
import win32com.client
import win32print
import win32api
import time
import openpyxl
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional, Tuple, List
from docx import Document
from pymorphy2 import MorphAnalyzer
# from docx.oxml import OxmlElement, qn


# ========== НАСТРОЙКИ ==========
# Путь для сохранения создаваемых карт для Тобольска
OUTPUT_CARDS_DIR_DEFAULT = r"C:\Users\user53\3_Тобольск"

# Путь к папке с исходными данными для режима "Ростов"
ROSTOV_DIR_DEFAULT = r"C:\Users\user53\4_Ростов"

# Путь к папке с исходными данными для режима "БВП"
BVP_DIR_DEFAULT = r"C:\Users\user53\6_БВП"

# Путь к папке с исходными данными для режима "ФЗО"
FZO_DIR_DEFAULT = r"C:\Users\user53\7_ФЗО"

# Путь для сохранения сгенерированных постановлений при массовой печати
MASSPRINT_SAVEPATH_DEFAULT = r"U:\Михайлова\СВО\Постановления"

# Путь к папке с шаблонами Word для постановлений
TEMPLATES_FOLDER_DEFAULT = r"U:\ШАБЛОНЫ\Заключения\СВО\Образцы"

# Путь к папке с постановлениями Тобольска
POSTANOVLENIYA_DIR_DEFAULT = r"U:\КАРТЫ - ПОСТАНОВЛЕНИЯ\ПОСТАНОВЛЕНИЯ\Тобольск"

# Путь к папке с постановлениями ЭКЦ
EKC_POSTANOVLENIYA_DIR_DEFAULT = r"U:\КАРТЫ - ПОСТАНОВЛЕНИЯ\ПОСТАНОВЛЕНИЯ\ЭКЦ\по делу Щегловой"

# Путь для сохранения создаваемых Карт для ЭКЦ
EKC_OUTPUT_DIR_DEFAULT = r"U:\КАРТЫ - ПОСТАНОВЛЕНИЯ\ПОСТАНОВЛЕНИЯ\ЭКЦ\по делу Щегловой\!ПРОВЕРЕНО"

# Имя файла лога
LOG_FILENAME = "card_filler.log"


# ========== НАСТРОЙКА ЛОГИРОВАНИЯ ==========
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOG_FILENAME, encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Инициализация морфологического анализатора
morph = MorphAnalyzer()


# ========== КОНФИГУРАЦИЯ СЛЕДОВАТЕЛЕЙ ==========
# Добавляйте новых следователей в этот список.
# Каждый элемент – словарь с ключами:
#   "fio_patterns" – список строк (вариантов написания ФИО),
#   "naz", "naz3", "naz4" – значения, которые нужно записать в карту.
INVESTIGATORS_CONFIG = [
    {
        "fio_patterns": ["Чекан П.Д.", "Чекан Павел Дмитриевич"],
        "naz": "следователя-криминалиста отдела криминалистики СУ СК РФ по Тюменской области подполковника юстиции Чекана П.Д.",
        "naz3": "Следователю-криминалисту отдела криминалистики СУ СК РФ по Тюменской области подполковнику юстиции",
        "naz4": "Чекану П.Д."
    },
    {
        "fio_patterns": ["А.А. Баннов", "Баннов А.А."],
        "naz": "начальника отделения ОУР УМВД России по г. Тюмени, старшего лейтенанта полиции А.А. Баннова",
        "naz3": "Начальнику отделения ОУР УМВД России по г. Тюмени, старшему лейтенанту полиции",
        "naz4": "Баннову А.А."
    },
    #{
    #    "fio_patterns": ["Петров П.П.", "Петров Петр Петрович"],
    #    "naz": "Следователя-криминалиста отдела криминалистики СУ СК России по Тюменской области Петрова П.П.",
    #    "naz3": "Следователю-криминалисту отдела криминалистики СУ СК России по Тюменской области",
    #    "naz4": "Петрову П.П."
    #},
]


# ========== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ==========
# Маппинг степени родства -> значение для ячейки РОД (родительный падеж)
ROD_MAPPING = {
    'отец': 'отца',
    'мать': 'матери',
    'брат': 'брата',
    'сестра': 'сестры',
    'сын': 'сына',
    'дочь': 'дочери',
    'дядя': 'дяди',
    'тётя': 'тёти',
    'бабушка': 'бабушки',
    'дедушка': 'дедушки',
    'неполнородная сестра': 'неполнородной сестры',
    'неполнородный брат': 'неполнородного брата',
    'не указано': 'родство не указано'
}


# ========== ИЗМЕНЕНИЯ В ФУНКЦИИ select_work_type() ==========
def select_work_type() -> str:
    """
    Предлагает пользователю выбрать тип работы.
    Возвращает строку-идентификатор выбранного варианта.
    """
    print("\nВыберите тип работы:")
    print("\n1. Заведующая: Заполнение Карт по постановлениям из ТОБОЛЬСКА по новорожденному\n(поместить в любую папку с образцом нужной Карты)")
    print("\n2. Регистратура: Заполнение Карт с КОСТЯМИ для РАСПРЕДЕЛЕНИЯ между экспертами\n(поместить в корневую папку '1_Кости' с папками экспертов)")
    print("\n3. Регистратура: Заполнение Карт с ОБРАЗЦАМИ в случае ПРЯМОЙ идентификации\n(поместить в папку с Картами, предназначенными для преобразования)")
    print("\n4. Регистратура: Массовое заполнение Карт с ОБРАЗЦАМИ (преобразование ФИО из родительного падежа в именительный)\n(поместить в папку с Картами образцов)")
    print("\n5. Регистратура: Массовая печать постановлений\n(поместить в папку с распечатываемыми Картами)")
    print("\n6. Эксперты: Заполнение Карт с КОСТЯМИ при ЗАВЕРШЕНИИ работы с заключением (дата окончания и объект)\n(поместить в папку с Картами костей)")
    print("\n7. Заведующая: Заполнение Карт по СПИСКУ ОБРАЗЦОВ ИЗ РОСТОВА\n(поместить в любую папку с образцом нужной Карты, обрабатываемый Excel - в папку 4_Ростов)")
    print("\n8. Заведующая: Заполнение Карт по постановлениям из ЭКЦ\n(поместить в любую папку с образцом нужной Карты)")
    print("\n9. Заведующая: Заполнение Карт по СПИСКУ ОБРАЗЦОВ БВП\n(поместить в любую папку с образцом нужной Карты, обрабатываемый Excel - в папку 6_БВП)")
    print("\n10. Заведующая: Заполнение Карт по СПИСКУ ОБРАЗЦОВ ФЗО\n(поместить в любую папку с образцом нужной Карты, обрабатываемый Excel - в папку 7_ФЗО)")

    while True:
        choice = input("\nВведите номер: ").strip()
        if choice == "1":
            return "tobolsk_newborn"
        elif choice == "2":
            return "distribution_cards"
        elif choice == "3":
            return "registry_cards"
        elif choice == "4":
            return "mass_fill_samples"
        elif choice == "5":
            return "mass_print_postanovleniy"
        elif choice == "6":
            return "bone_cards"
        elif choice == "7":
            return "rostov_cards"
        elif choice == "8":
            return "ekc_postanovleniya"
        elif choice == "9":
            return "bvp_cards"
        elif choice == "10":
            return "fzo_cards"
        else:
            print("Неверный выбор, попробуйте снова.")


def normalize_text(text: Optional[str]) -> str:
    """Удаляет лишние пробелы и приводит к строке"""
    if text is None:
        return ""
    return str(text).strip()


def format_two_digits(value) -> str:
    """
    Преобразует число (или строку с числом) в строку с ведущим нулём,
    если число состоит из одной цифры. Иначе возвращает строковое представление.
    Пример: 8 -> '08', 12 -> '12', '8' -> '08', '08' -> '08'.
    Если value не является числом (и не может быть преобразовано), возвращается исходная строка.
    """
    if value is None or value == "":
        return ""
    try:
        # Пробуем преобразовать в целое
        num = int(str(value).strip())
        # Форматируем с ведущим нулём, если число от 0 до 9
        return f"{num:02d}"
    except (ValueError, TypeError):
        # Не число – возвращаем как есть
        logger.debug(f"format_two_digits: не удалось преобразовать '{value}' в число, возвращено как есть")
        return str(value)


def format_date_ddmmyyyy(date_str: Optional[str]) -> str:
    """
    Принимает строку даты в формате DD.MM.YYYY (день и месяц могут быть однозначными).
    Возвращает строку с день и месяц, дополненными ведущим нулём до двух цифр.
    Если строка не соответствует формату (не три части, разделённые точками),
    возвращает исходную строку с предупреждением в лог.
    """
    if not date_str or not isinstance(date_str, str):
        return date_str if date_str is not None else ""
    parts = date_str.strip().split('.')
    if len(parts) != 3:
        logger.warning(f"format_date_ddmmyyyy: строка '{date_str}' не соответствует формату ДД.ММ.ГГГГ, возвращена без изменений")
        return date_str
    day, month, year = parts
    formatted_day = format_two_digits(day)
    formatted_month = format_two_digits(month)
    result = f"{formatted_day}.{formatted_month}.{year}"
    if result != date_str:
        logger.debug(f"Дата отформатирована: '{date_str}' -> '{result}'")
    return result


def add_dot_to_initials(text: str) -> str:
    """Если строка состоит из одной буквы (не содержит точки), добавляет точку."""
    if not text:
        return text
    text = text.strip()
    if len(text) == 1 and text.isalpha():
        return text + "."
    return text


def parse_date_flexible(date_value) -> Optional[datetime]:
    """
    Пытается преобразовать входное значение (строка или datetime) в объект datetime.
    Поддерживает различные форматы, включая наличие времени.
    Возвращает datetime или None, если не удалось.
    """
    if date_value is None:
        return None
    
    # Если уже datetime, возвращаем как есть
    if isinstance(date_value, datetime):
        return date_value
    
    # Преобразуем в строку
    date_str = str(date_value).strip()
    if not date_str:
        return None
    
    # Список возможных форматов (без времени и с временем)
    formats = [
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M:%S.%f",
        "%d.%m.%Y %H:%M:%S",
        "%d.%m.%Y %H:%M",
        "%d.%m.%Y",
        "%d.%m.%y",
        "%Y-%m-%d",
        "%d/%m/%Y",
        "%d/%m/%y",
        "%m/%d/%Y",
        "%m/%d/%y",
        "%Y/%m/%d",
    ]
    
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    
    # Если стандартные форматы не помогли, пробуем извлечь компоненты с помощью регулярного выражения
    patterns = [
        r'(\d{1,2})[./-](\d{1,2})[./-](\d{2,4})',   # DD.MM.YYYY или MM.DD.YYYY
        r'(\d{4})[./-](\d{1,2})[./-](\d{1,2})',     # YYYY-MM-DD
    ]
    
    for pattern in patterns:
        match = re.search(pattern, date_str)
        if match:
            parts = match.groups()
            # Если первый компонент длиной 4, то это год
            if len(parts[0]) == 4:
                year, month, day = map(int, parts)
            else:
                day, month, year = map(int, parts)
                # Простая проверка: если день > 31 или месяц > 12, возможно, поменяны местами
                if day > 31 or month > 12:
                    day, month = month, day
            # Корректировка двузначного года
            if year < 100:
                year += 2000 if year < 50 else 1900  # для дат после 1950
            try:
                return datetime(year, month, day)
            except ValueError:
                pass
    
    logger.warning(f"Не удалось распознать дату: {date_str}")
    return None


def natural_sort_key(filename: str):
    """
    Ключ для натуральной сортировки (1, 2, 10, 23 вместо 1, 10, 2, 23)
    """
    return [int(text) if text.isdigit() else text.lower()
            for text in re.split(r'(\d+)', filename)]


def validate_ob_format(value: str) -> Tuple[bool, Optional[int], Optional[str]]:
    """
    Проверяет формат ОБ: число + латинская буква
    Возвращает: (валидно, число, буква)
    """
    pattern = r'^(\d+)([A-Za-z])$'
    match = re.match(pattern, value)
    if match:
        return True, int(match.group(1)), match.group(2).upper()
    return False, None, None


def get_current_date() -> Tuple[int, str]:
    """Возвращает текущий день (число) и месяц в родительном падеже (например, 'марта')"""
    now = datetime.now()
    day = now.day
    # Словарь для преобразования номера месяца в русское название (именительный падеж)
    month_names = {
        1: "январь", 2: "февраль", 3: "март", 4: "апрель",
        5: "май", 6: "июнь", 7: "июль", 8: "август",
        9: "сентябрь", 10: "октябрь", 11: "ноябрь", 12: "декабрь"
    }
    month_nom = month_names[now.month]
    # Преобразуем в родительный падеж с помощью pymorphy2
    parsed = morph.parse(month_nom)[0]
    month_gen = parsed.inflect({'gent'})
    if month_gen:
        month_name_gen = month_gen.word
    else:
        # fallback словарь
        month_gen_dict = {
            "январь": "января", "февраль": "февраля", "март": "марта",
            "апрель": "апреля", "май": "мая", "июнь": "июня",
            "июль": "июля", "август": "августа", "сентябрь": "сентября",
            "октябрь": "октября", "ноябрь": "ноября", "декабрь": "декабря"
        }
        month_name_gen = month_gen_dict.get(month_nom, month_nom)
    return day, month_name_gen


def find_template_excel(script_dir: Path) -> Optional[Path]:
    """
    Находит файл Excel-шаблона в указанной директории.
    Если найден один файл, возвращает его.
    Если несколько, выводит список и просит выбрать.
    Если ни одного, возвращает None.
    """
    excel_files = list(script_dir.glob("*.xlsx")) + list(script_dir.glob("*.xlsm"))
    # Исключаем временные файлы
    excel_files = [f for f in excel_files if not f.name.startswith('~$')]
    if not excel_files:
        logger.error("В папке программы не найдено файлов Excel (шаблонов карт)")
        return None
    if len(excel_files) == 1:
        logger.info(f"Найден шаблон карты: {excel_files[0].name}")
        return excel_files[0]
    else:
        print("Найдено несколько файлов Excel. Выберите шаблон:")
        for i, f in enumerate(excel_files, 1):
            print(f"{i}. {f.name}")
        while True:
            try:
                choice = int(input("Введите номер: ").strip())
                if 1 <= choice <= len(excel_files):
                    return excel_files[choice-1]
                else:
                    print("Неверный номер.")
            except ValueError:
                print("Введите число.")


def find_cell_by_code(ws, code: str):
    """
    Ищет ячейку в первом столбце с указанным кодом.
    Сравнение производится после приведения к строке и удаления пробелов.
    Возвращает ячейку в третьем столбце (столбец значений).
    """
    target = str(code).strip()
    for row in ws.iter_rows(min_row=1, max_col=3):
        cell_code = row[0].value
        if cell_code is not None:
            # Приводим значение из ячейки к строке и удаляем пробелы
            if str(cell_code).strip() == target:
                return row[2]
    return None


def month_num_to_genitive(month_num: int) -> str:
    """Преобразование номера месяца в русское название в родительном падеже"""
    months = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля",
        5: "мая", 6: "июня", 7: "июля", 8: "августа",
        9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
    }
    return months.get(month_num, "")


def sanitize_filename(name: str) -> str:
    """Удаляет недопустимые для имени файла символы."""
    invalid_chars = r'[<>:"/\\|?*]'
    name = re.sub(invalid_chars, '_', name)
    name = name.strip().strip('.')
    if not name:
        name = "unknown"
    return name


def normalize_degree(degree: str) -> str:
    """Приводит строку степени родства к каноническому виду."""
    degree = degree.lower().strip().replace('.', '').replace(',', '')
    if degree in ('личный генотип', 'личный', 'прямая идентификация', 'прямая', 'прямой'):
        return 'личный генотип'
    if degree in ('отец', 'папа'):
        return 'отец'
    if degree in ('мать', 'мама'):
        return 'мать'
    if degree in ('брат', 'братан'):
        return 'брат'
    if degree in ('сестра', 'сестрёнка'):
        return 'сестра'
    if degree in ('сын', 'сыночек'):
        return 'сын'
    if degree in ('дочь', 'дочка'):
        return 'дочь'
    if degree in ('дядя', 'дядюшка'):
        return 'дядя'
    if degree in ('тётя', 'тетя', 'тётка'):
        return 'тётя'
    if degree in ('бабушка', 'бабка'):
        return 'бабушка'
    if degree in ('дедушка', 'дед'):
        return 'дедушка'
    if degree in ('неполнородная сестра', 'неполнородная сестра'):
        return 'неполнородная сестра'
    if degree in ('неполнородный брат', 'неполнородный брат'):
        return 'неполнородный брат'
    if degree in ('родство не указано', 'не указано', 'неизвестно', 'н/у', 'н\\у'):
        return 'не указано'
    return degree


def parse_full_name(full_name: str) -> Tuple[str, str, str]:
    """
    Разбирает строку ФИО на фамилию, имя, отчество.
    Поддерживает сложные отчества (оглы, кызы и т.д.) и скобки с пробелами внутри.
    Возвращает кортеж (фамилия, имя, отчество).
    """
    if not full_name:
        return "", "", ""

    full_name = full_name.strip()
    full_name = re.sub(r'\s+', ' ', full_name)

    # Обработка инициалов (например, "Иванов И.И.")
    if '.' in full_name:
        match = re.match(r'^([А-ЯЁа-яё]+)\s+([А-ЯЁ]\.?)\s*([А-ЯЁ]\.?)?$', full_name)
        if match:
            surname = match.group(1)
            first_initial = match.group(2).replace('.', '')
            patronymic_initial = match.group(3).replace('.', '') if match.group(3) else ""
            return surname, first_initial, patronymic_initial

    # Токенизация: скобочные группы (включая пробелы) – как один токен, остальные слова – по пробелам
    tokens = re.findall(r'\([^()]+\)|[^\s]+', full_name)
    word_count = len(tokens)
    complex_keywords = {'оглы', 'кызы', 'угли', 'гызы', 'заде', 'бек', 'хан'}

    # Удаляем завершающий "?" (неизвестное отчество)
    if word_count >= 3 and tokens[-1] == '?':
        tokens = tokens[:-1]
        word_count = len(tokens)

    if word_count == 0:
        return "", "", ""
    if word_count == 1:
        return tokens[0], "", ""
    if word_count == 2:
        return tokens[0], tokens[1], ""

    # Для 3 и более токенов
    surname = tokens[0]

    def has_keyword(token: str) -> bool:
        """Проверяет, содержит ли токен ключевое слово (игнорируя скобки)."""
        clean = token[1:-1] if token.startswith('(') and token.endswith(')') else token
        return any(kw in clean.lower() for kw in complex_keywords)

    # Ищем позицию, с которой начинается отчество.
    # Отчество начинается с токена, который либо сам является ключевым словом,
    # либо предшествует ключевому слову
    start_idx = None
    for i in range(2, word_count):
        if has_keyword(tokens[i]):
            # Ключевое слово найдено на позиции i.
            # Отчество начинается с i-1, если i-1 >= 2, иначе с i.
            start_idx = i - 1 if i - 1 >= 2 else i
            break

    if start_idx is not None:
        first_name = ' '.join(tokens[1:start_idx])
        patronymic = ' '.join(tokens[start_idx:])
    else:
        # Не нашли ключевого слова – берём имя = второй токен, отчество = остальное
        first_name = tokens[1] if word_count > 1 else ""
        patronymic = ' '.join(tokens[2:]) if word_count > 2 else ""
        if word_count > 3:
            logger.warning(f"Не найдено ключевое слово в отчестве, ФИО: {full_name}")

    return surname, first_name, patronymic


def parse_txt_file(file_path: Path) -> List[Dict[str, str]]:
    """
    Парсинг txt файла с данными (пробуем разные кодировки)
    Возвращает список словарей с ключами 'order', 'corpse_number', 'date'
    """
    data = []
    encodings = ['utf-8', 'cp1251', 'windows-1251', 'ansi']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                lines = f.readlines()

            # Пропускаем заголовок (первая строка)
            for line in lines[1:]:
                line = line.strip()
                if line:
                    # Разделяем по табуляции или по нескольким пробелам
                    parts = re.split(r'\t|\s{2,}', line)
                    if len(parts) >= 3:
                        # Удаляем возможные лишние пробелы
                        parts = [p.strip() for p in parts]
                        data.append({
                            'order': parts[0],
                            'corpse_number': parts[1],
                            'date': parts[2]
                        })
            # Если удалось прочитать, прерываем цикл
            if data:
                logger.info(f"Успешно прочитан файл {file_path.name} с кодировкой {encoding}")
                break

        except UnicodeDecodeError:
            continue  # Пробуем следующую кодировку
        except Exception as e:
            logger.error(f"Ошибка при чтении файла {file_path} с кодировкой {encoding}: {e}")
            continue

    if not data:
        logger.warning(f"Не удалось прочитать данные из файла {file_path} ни с одной кодировкой")

    return data


def fill_excel_from_template(template_path: Path, output_path: Path, expertise_data: Dict[str, str]) -> bool:
    """Заполнение шаблона Excel данными и сохранение в output_path"""
    try:
        wb = openpyxl.load_workbook(template_path)
        ws = wb.active

        # Заполняем ячейки по кодам
        codes_map = {
            "НОМ": expertise_data.get('expertise_number'),
            "ДП": expertise_data.get('day'),
            "МП": expertise_data.get('month'),
            "ДН": expertise_data.get('start_day'),
            "МН": expertise_data.get('start_month'),
            "ТР": expertise_data.get('corpse_number')
        }

        for code, value in codes_map.items():
            if value is not None:
                cell = find_cell_by_code(ws, code)
                if cell:
                    cell.value = value
                    logger.debug(f"Записан код {code} = {value}")
                else:
                    logger.warning(f"Код {code} не найден в шаблоне")

        wb.save(output_path)
        logger.info(f"Создан файл: {output_path}")
        return True

    except Exception as e:
        logger.error(f"Ошибка при заполнении Excel файла {output_path}: {e}")
        return False


def find_investigator(text: str) -> Optional[Dict[str, str]]:
    """
    Ищет в тексте ФИО следователя из предопределённого списка.
    Возвращает словарь с ключами 'naz', 'naz3', 'naz4' или None, если не найдено.
    """
    if not text:
        logger.warning("Текст постановления пуст, поиск следователя невозможен")
        return None

    text_lower = text.lower()
    for inv in INVESTIGATORS_CONFIG:
        for pattern in inv["fio_patterns"]:
            if pattern.lower() in text_lower:
                logger.info(f"Найден следователь по шаблону '{pattern}': {inv}")
                return {
                    "НАЗ": inv["naz"],
                    "НАЗ3": inv["naz3"],
                    "НАЗ4": inv["naz4"]
                }
    logger.info("Следователь из предопределённого списка не найден")
    return None


def split_name_part_with_brackets(part: str) -> Tuple[str, str]:
    """
    Разделяет часть ФИО (фамилию, имя или отчество) на основную часть и альтернативу в скобках.
    Возвращает (main_part, alternative_part). Если скобок нет, alternative_part = ''.
    Пример: "Суберджон (Собиржон)" -> ("Суберджон", "Собиржон")
    """
    if not part:
        return "", ""
    part = part.strip()
    # Ищем шаблон: основная часть, затем пробел, затем (альтернатива) в конце
    pattern = r'^(.*?)\s*\(([^)]+)\)\s*$'
    match = re.match(pattern, part)
    if match:
        main_part = match.group(1).strip()
        alt_part = match.group(2).strip()
        logger.debug(f"split_name_part_with_brackets: '{part}' -> main='{main_part}', alt='{alt_part}'")
        return main_part, alt_part
    else:
        logger.debug(f"split_name_part_with_brackets: '{part}' не содержит скобок")
        return part, ""


def format_name_part_with_alternative(main: str, alt: str) -> str:
    """
    Формирует строку из основной части и альтернативы в скобках.
    Если alt пуста, возвращает main.
    """
    if not alt:
        return main
    return f"{main} ({alt})"


# ========== КЛАСС ДЛЯ ИЗВЛЕЧЕНИЯ ДАННЫХ ИЗ ПОСТАНОВЛЕНИЯ ==========
class PostanovlenieExtractor:
    """Извлекает данные из Word-документа постановления (поддерживает .docx и .doc)"""

    # Регулярные выражения для поиска (старые, жёсткие, для fallback)
    RE_FIO_BIRTH = re.compile(
        r"у\s+свидетеля\s+([А-ЯЁа-яё]+)\s+([А-ЯЁа-яё]+)\s+([А-ЯЁа-яё]+)\s+(\d{2}\.\d{2}\.\d{4})",
        re.IGNORECASE
    )
    RE_FIO_BIRTH_NOMINATIVE = re.compile(
        r"Свидетелем\s+по\s+уголовному\s+делу[,\s]*является\s+([А-ЯЁа-яё]+)\s+([А-ЯЁа-яё]+)\s+([А-ЯЁа-яё]+)\s+(\d{2}\.\d{2}\.\d{4})",
        re.IGNORECASE
    )
    RE_RESOLUTION_DATE = re.compile(
        r"«(\d{1,2})»\s+([а-яё]+)\s+\d{4}\s+года",
        re.IGNORECASE
    )

    def __init__(self, file_path: Path):
        self.file_path = file_path
        self.text = self._extract_text()

    def _extract_text(self) -> str:
        """Извлекает весь текст из документа .docx или .doc"""
        try:
            if self.file_path.suffix.lower() == '.docx':
                doc = Document(self.file_path)
                full_text = [para.text for para in doc.paragraphs]
                return "\n".join(full_text)
            elif self.file_path.suffix.lower() == '.doc':
                word = None
                doc = None
                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(str(self.file_path.absolute()))
                    text = doc.Content.Text
                    logger.info(f"Текст из .doc извлечён через win32com")
                    return text
                except Exception as e:
                    logger.error(f"Ошибка при чтении .doc файла {self.file_path}: {e}")
                    raise
                finally:
                    if doc:
                        try:
                            doc.Close()
                            logger.debug(f"Документ {self.file_path.name} закрыт")
                        except Exception as e:
                            logger.warning(f"Не удалось закрыть документ {self.file_path.name}: {e}")
                    if word:
                        try:
                            word.Quit()
                            logger.debug("Приложение Word завершено")
                        except Exception as e:
                            logger.warning(f"Не удалось завершить приложение Word: {e}")
            else:
                logger.warning(f"Неподдерживаемый формат файла: {self.file_path.suffix}")
                return ""
        except Exception as e:
            logger.error(f"Ошибка при чтении файла {self.file_path}: {e}")
            return ""

    def extract_fio_birth(self) -> Optional[Dict[str, str]]:
        """
        Извлекает ФИО и дату рождения.
        Возвращает словарь с ключами: 'surname', 'first_name', 'patronymic', 'birth_date', 'case'
        case может быть 'genitive' (родительный падеж) или 'nominative' (именительный).
        Поддерживает 2, 3, 4 и 5 слов в ФИО (для сложных отчеств типа "угли").
        """
        logger.info(f"Начинаем извлечение ФИО из файла {self.file_path.name}")
        
        # Гибкие регулярные выражения (захватывают до 5 слов)
        RE_FIO_BIRTH_FLEXIBLE = re.compile(
            r"у\s+свидетеля\s+((?:[А-ЯЁа-яё]+\s+){1,4}[А-ЯЁа-яё]+)\s+(\d{1,2}\.\d{1,2}\.\d{4})",
            re.IGNORECASE
        )

        RE_FIO_BIRTH_NOMINATIVE_FLEXIBLE = re.compile(
            r"Свидетелем\s+по\s+уголовному\s+делу[,\s]*является\s+((?:[А-ЯЁа-яё]+\s+){1,4}[А-ЯЁа-яё]+)\s+(\d{1,2}\.\d{1,2}\.\d{4})",
            re.IGNORECASE
        )
        
        COMPLEX_PATRONYMIC_KEYWORDS = [
            'угли', 'углы', 'оглы', 'оглу', 'кызы', 'гызы', 'заде', 'бек', 'хан'
        ]
        
        def parse_fio_words(fio_string, case_type):
            """
            Парсит строку ФИО, корректно обрабатывая скобки с пробелами.
            Возвращает dict с ключами 'surname', 'first_name', 'patronymic', 'case'.
            """
            logger.info(f"Парсинг ФИО: '{fio_string}' (падеж: {case_type})")
            
            # Токенизация: скобочные группы как единое целое, остальные слова по пробелам
            tokens = re.findall(r'\([^()]+\)|[^\s]+', fio_string.strip())
            word_count = len(tokens)
            logger.info(f"Токены ({word_count}): {tokens}")
            
            COMPLEX_KEYWORDS = ['оглы', 'кызы', 'угли', 'углы', 'оглу', 'гызы', 'заде', 'бек', 'хан']
            
            def is_complex(token):
                # Убираем скобки для проверки
                clean = token[1:-1] if token.startswith('(') and token.endswith(')') else token
                return any(kw in clean.lower() for kw in COMPLEX_KEYWORDS)
            
            # Ищем, с какого токена начинается отчество
            # Реализуем адаптивный алгоритм
            if word_count < 2:
                logger.error(f"Слишком мало токенов: {tokens}")
                return None
            
            # Начальное присваивание
            surname = tokens[0]
            # Определяем, где заканчивается имя
            # Имя может состоять из одного или двух токенов, если второй токен – скобка
            if word_count >= 2 and tokens[1].startswith('(') and tokens[1].endswith(')'):
                first_name = f"{tokens[0]} {tokens[1]}"
                if word_count >= 3 and tokens[2].startswith('(') and tokens[2].endswith(')'):
                    first_name = f"{tokens[1]} {tokens[2]}"
                    rest_start = 3
                else:
                    first_name = tokens[1]
                    rest_start = 2
            else:
                # Обычный случай: имя – второй токен
                first_name = tokens[1] if word_count > 1 else ""
                rest_start = 2
            
            # Оставшиеся токены – отчество (может быть пустым)
            if rest_start < word_count:
                # Отчество может включать несколько токенов, в том числе скобки
                patronymic = ' '.join(tokens[rest_start:])
            else:
                patronymic = ""
            
            # Дополнительная эвристика
            if patronymic and patronymic.strip().startswith('(') and '(' not in first_name:
                # Перемещаем первый токен отчества в имя
                first_token_patr = tokens[rest_start]
                if first_token_patr.startswith('(') and first_token_patr.endswith(')'):
                    first_name = f"{first_name} {first_token_patr}"
                    patronymic = ' '.join(tokens[rest_start+1:]) if rest_start+1 < word_count else ""
                    logger.info(f"Корректировка: альтернатива имени перемещена из отчества в имя. Имя='{first_name}', Отчество='{patronymic}'")
            
            logger.info(f"Результат: фамилия='{surname}', имя='{first_name}', отчество='{patronymic}'")
            return {
                'surname': surname,
                'first_name': first_name,
                'patronymic': patronymic,
                'case': case_type
            }
        
        # 1. Гибкий родительный падеж
        logger.debug("Поиск гибким шаблоном родительного падежа (до 5 слов)...")
        match = RE_FIO_BIRTH_FLEXIBLE.search(self.text)
        if match:
            fio_string, birth_date = match.groups()
            logger.info(f"НАЙДЕН гибкий родительный: ФИО='{fio_string}', дата='{birth_date}'")
            parsed = parse_fio_words(fio_string, 'genitive')
            if parsed:
                parsed['birth_date'] = birth_date
                logger.info(f"Успешно: {parsed}")
                return parsed
        
        # 2. Гибкий именительный падеж
        logger.debug("Поиск гибким шаблоном именительного падежа (до 5 слов)...")
        match = RE_FIO_BIRTH_NOMINATIVE_FLEXIBLE.search(self.text)
        if match:
            fio_string, birth_date = match.groups()
            logger.info(f"НАЙДЕН гибкий именительный: ФИО='{fio_string}', дата='{birth_date}'")
            parsed = parse_fio_words(fio_string, 'nominative')
            if parsed:
                parsed['birth_date'] = birth_date
                logger.info(f"Успешно: {parsed}")
                return parsed
        
        # 3. Fallback на старые жёсткие шаблоны (3 слова, дата строго две цифры)
        logger.info("Гибкие шаблоны не сработали, пробуем старые жёсткие...")
        match = self.RE_FIO_BIRTH.search(self.text)
        if match:
            surname, first_name, patronymic, birth_date = match.groups()
            logger.info(f"Старый родительный: {surname} {first_name} {patronymic}, дата {birth_date}")
            return {
                'surname': surname,
                'first_name': first_name,
                'patronymic': patronymic,
                'birth_date': birth_date,
                'case': 'genitive'
            }
        match = self.RE_FIO_BIRTH_NOMINATIVE.search(self.text)
        if match:
            surname, first_name, patronymic, birth_date = match.groups()
            logger.info(f"Старый именительный: {surname} {first_name} {patronymic}, дата {birth_date}")
            return {
                'surname': surname,
                'first_name': first_name,
                'patronymic': patronymic,
                'birth_date': birth_date,
                'case': 'nominative'
            }
        
        logger.warning(f"ФИО и дата не найдены в {self.file_path.name}")
        return None

    def extract_resolution_date(self) -> Optional[Dict[str, str]]:
        """Извлекает дату постановления (день и месяц в текстовом виде)."""
        # Шаблон 1: с кавычками «день» месяц год
        pattern1 = re.compile(
            r'«(\d{1,2})»\s+([а-яё]+)\s+\d{4}\s+года',
            re.IGNORECASE
        )
        match = pattern1.search(self.text)
        if match:
            day, month = match.groups()
            logger.info(f"Извлечена дата постановления (с кавычками): {day} {month}")
            return {'day': day, 'month': month}

        # Шаблон 2: без кавычек (день месяц год)
        pattern2 = re.compile(
            r'(\d{1,2})\s+([а-яё]+)\s+\d{4}\s+года',
            re.IGNORECASE
        )
        match = pattern2.search(self.text)
        if match:
            day, month = match.groups()
            logger.info(f"Извлечена дата постановления (без кавычек): {day} {month}")
            return {'day': day, 'month': month}

        # Шаблон 3: дата в формате DD.MM.YYYY года
        pattern3 = re.compile(
            r'(\d{1,2})\.(\d{1,2})\.(\d{4})\s+года',
            re.IGNORECASE
        )
        match = pattern3.search(self.text)
        if match:
            day, month_num, year = match.groups()
            # Преобразуем номер месяца в русское название
            month_names = {
                1: "января", 2: "февраля", 3: "марта", 4: "апреля",
                5: "мая", 6: "июня", 7: "июля", 8: "августа",
                9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
            }
            month = month_names.get(int(month_num), "")
            if month:
                logger.info(f"Извлечена дата постановления (с точками): {day} {month}")
                return {'day': day, 'month': month}
            else:
                logger.warning(f"Неверный номер месяца в дате: {month_num}")

        logger.warning(f"Не найдена дата постановления в файле {self.file_path}")
        return None

    def extract_ekc_data(self) -> Optional[Dict[str, str]]:
        """
        Извлекает ФИО и дату рождения из постановления ЭКЦ по делу Щегловой.
        Ожидает фразу: "Прошу получить генетический профиль Иванова И.И., 10.10.1980 г.р."
        Возвращает словарь с ключами surname, first_name, patronymic, birth_date, case='genitive'.
        """
        logger.info(f"Извлечение данных ЭКЦ из файла {self.file_path.name}")
        text = self.text

        # Гибкий шаблон для поиска фразы и даты (учитывает возможные опечатки)
        phrase_pattern = re.compile(
            r"Прошу\s+получить\s+генетический\s+профиль\s+(.+?)(\d{1,2}\.\d{1,2}\.\d{4})",
            re.IGNORECASE | re.DOTALL
        )
        match = phrase_pattern.search(text)
        if not match:
            logger.warning(f"В файле {self.file_path.name} не найдена фраза 'Прошу получить генетический профиль' с датой рождения")
            return None

        fio_part = match.group(1).strip()
        birth_date_raw = match.group(2).strip()

        # Убираем возможные суффиксы типа "г.р.", "г.р", "гр." после даты
        after_date = text[match.end(2):].lstrip()
        gr_match = re.match(r'\s*(?:г\.р\.|г\.р|гр\.|гр|г\.\s*р\.)?', after_date, re.IGNORECASE)
        if gr_match:
            # Суффикс есть, но для ДР используем только чистую дату
            logger.debug(f"Обнаружен суффикс даты: '{gr_match.group().strip()}'")
        # Оставляем только дату в формате DD.MM.YYYY
        birth_date = birth_date_raw

        # Очищаем ФИО от висящих знаков препинания в конце (запятые, точки)
        fio_part = re.sub(r'[,.\s]+$', '', fio_part)
        logger.info(f"Извлечено: ФИО = '{fio_part}', дата рождения = '{birth_date}'")

        # Разбираем ФИО (ожидается родительный падеж)
        surname, first_name, patronymic = parse_full_name(fio_part)
        if not surname:
            logger.error(f"Не удалось разобрать ФИО: '{fio_part}' в файле {self.file_path.name}")
            return None

        logger.info(f"Разобрано: фамилия='{surname}', имя='{first_name}', отчество='{patronymic}' (падеж родительный)")
        return {
            'surname': surname,
            'first_name': first_name,
            'patronymic': patronymic,
            'birth_date': birth_date,
            'case': 'genitive'
        }

    def extract_dative_fio(self) -> Optional[str]:
        """
        Извлекает ФИО в дательном падеже из фразы:
        "Принадлежит ли сперма, обнаруженная в прямой кишке потерпевшей Щегловой Н.А., Иванову И.И.?"
        Возвращает строку с ФИО (например "Иванову И.И.") или None.
        """
        logger.info(f"Поиск ФИО в дательном падеже в файле {self.file_path.name}")
        text = self.text

        # Шаблон с учётом возможных вариаций пробелов, запятых и точек
        pattern = re.compile(
            r"Принадлежит\s+ли\s+сперма\s*,\s*обнаруженная\s+в\s+прямой\s+кишке\s+потерпевшей\s+[А-ЯЁа-яё]+\s+[А-ЯЁ]\s*\.\s*[А-ЯЁ]\s*\.\s*,\s*([А-ЯЁа-яё]+\s+[А-ЯЁ]\s*\.\s*[А-ЯЁ]\s*\.\s*\?)",
            re.IGNORECASE | re.DOTALL
        )
        match = pattern.search(text)
        if match:
            dative_fio = match.group(1).strip()
            # Убираем возможный вопросительный знак в конце
            dative_fio = re.sub(r'\?+$', '', dative_fio).strip()
            logger.info(f"Найдено ФИО в дательном падеже: '{dative_fio}'")
            return dative_fio

        # Если не нашли по строгому шаблону, пробуем более гибкий вариант
        # (ищем после последней запятой перед вопросом)
        alternative_pattern = re.compile(
            r"прямой\s+кишке\s+потерпевшей\s+[^,]+,?\s*([А-ЯЁа-яё]+\s+[А-ЯЁ]\s*\.\s*[А-ЯЁ]\s*\.\s*\?)",
            re.IGNORECASE
        )
        alt_match = alternative_pattern.search(text)
        if alt_match:
            dative_fio = alt_match.group(1).strip()
            dative_fio = re.sub(r'\?+$', '', dative_fio).strip()
            logger.info(f"Найдено ФИО в дательном падеже (альтернативный поиск): '{dative_fio}'")
            return dative_fio

        logger.warning(f"ФИО в дательном падеже не найдено в файле {self.file_path.name}")
        return None

# ========== КЛАСС ДЛЯ РАБОТЫ С КАРТОЙ EXCEL ==========
class ExcelCardFiller:
    """Заполняет карту Excel данными"""

    def __init__(self, template_path: Path):
        if not template_path.exists():
            raise FileNotFoundError(f"Шаблон карты не найден: {template_path}")
        self.template_path = template_path

    def create_copy(self, new_name: str, output_dir: Optional[Path] = None) -> Path:
        """
        Создаёт копию шаблона с новым именем.
        Если указан output_dir, копия сохраняется в эту папку, иначе рядом с шаблоном.
        Возвращает путь к созданной копии.
        """
        # Определяем целевую папку
        if output_dir is None:
            target_dir = self.template_path.parent
        else:
            target_dir = Path(output_dir)
            # Создаём папку, если её нет
            try:
                target_dir.mkdir(parents=True, exist_ok=True)
                logger.info(f"Папка для сохранения карт: {target_dir}")
            except Exception as e:
                logger.error(f"Не удалось создать папку {target_dir}: {e}")
                raise

        new_path = target_dir / new_name
        if new_path.exists():
            logger.warning(f"Файл {new_path} уже существует, он будет перезаписан")
        shutil.copy2(self.template_path, new_path)
        logger.info(f"Создана копия шаблона: {new_path}")
        return new_path

    def fill_card(self, workbook: openpyxl.Workbook, data: Dict[str, str]) -> None:
        """
        Заполняет ячейки карты значениями из словаря data.
        Ключи словаря – коды, значения – текст для вставки.
        """
        worksheet = workbook.active
        for code, value in data.items():
            cell = find_cell_by_code(worksheet, code)
            if cell:
                cell.value = value
                logger.debug(f"Записано {code} = {value}")
            else:
                logger.warning(f"Код {code} не найден в карте, значение не записано")

    def convert_to_nominative(self, word: str) -> str:
        """
        Преобразует слово из родительного падежа в именительный.
        Возвращает исходное слово, если преобразование невозможно.
        """
        if not word or not isinstance(word, str):
            return word
        word = word.strip()
        if not word:
            return word
        try:
            parsed = morph.parse(word)[0]
            nom = parsed.inflect({'nomn'})
            if nom:
                return nom.word.capitalize()
            else:
                logger.warning(f"Не удалось преобразовать слово '{word}' в именительный падеж")
                return word
        except Exception as e:
            logger.error(f"Ошибка при преобразовании слова '{word}': {e}")
            return word

    def convert_to_genitive(self, word: str) -> str:
        """
        Преобразует слово из именительного падежа в родительный.
        Возвращает исходное слово, если преобразование невозможно.
        """
        if not word or not isinstance(word, str):
            return word
        word = word.strip()
        if not word:
            return word
        try:
            parsed = morph.parse(word)[0]
            gen = parsed.inflect({'gent'})
            if gen:
                return gen.word.capitalize()
            else:
                logger.warning(f"Не удалось преобразовать слово '{word}' в родительный падеж")
                return word
        except Exception as e:
            logger.error(f"Ошибка при преобразовании слова '{word}': {e}")
            return word

    def fill_nominative_names(self, workbook: openpyxl.Workbook) -> None:
        """
        Находит в карте строки с кодами ФР, ИР, ОР (родительный падеж),
        преобразует их в именительный и записывает в ФИ, ИИ, ОИ.
        """
        worksheet = workbook.active
        fr_cell = find_cell_by_code(worksheet, "ФР")
        ir_cell = find_cell_by_code(worksheet, "ИР")
        or_cell = find_cell_by_code(worksheet, "ОР")

        fr = normalize_text(fr_cell.value if fr_cell else None)
        ir = normalize_text(ir_cell.value if ir_cell else None)
        or_ = normalize_text(or_cell.value if or_cell else None)

        fi = self.convert_to_nominative(fr) if fr else ""
        ii = self.convert_to_nominative(ir) if ir else ""
        oi = self.convert_to_nominative(or_) if or_ else ""

        data = {}
        if fi:
            data["ФИ"] = fi
        if ii:
            data["ИИ"] = ii
        if oi:
            data["ОИ"] = oi

        if data:
            self.fill_card(workbook, data)
            logger.info(f"Записаны именительные формы: ФИ={fi}, ИИ={ii}, ОИ={oi}")
        else:
            logger.info("Нет данных для преобразования в именительный падеж")


# ========== ФУНКЦИИ ОБРАБОТКИ РАЗНЫХ ТИПОВ РАБОТ ==========
def process_tobolsk_newborn(script_dir: Path, template_path: Path, start_number: int) -> None:
    """
    Обработка постановлений из Тобольска по новорожденным.
    """
    # Запрос пути к папке с постановлениями
    user_input = input("Введите путь к папке с постановлениями (или нажмите Enter для значения по умолчанию): ").strip()
    if user_input:
        postanovleniya_dir = user_input
    else:
        postanovleniya_dir = POSTANOVLENIYA_DIR_DEFAULT
        logger.info(f"Используется путь по умолчанию: {postanovleniya_dir}")

    postanovleniya_path = Path(postanovleniya_dir)
    if not postanovleniya_path.exists():
        logger.error(f"Папка с постановлениями не найдена: {postanovleniya_path}")
        logger.info("Операция прервана из-за отсутствия папки с постановлениями.")
        return
    logger.info(f"Папка с постановлениями найдена: {postanovleniya_path}")

    # Запрос пути для сохранения карт
    user_input_output = input("Введите путь для сохранения карт (или нажмите Enter для значения по умолчанию): ").strip()
    if user_input_output:
        output_cards_dir = user_input_output
    else:
        output_cards_dir = OUTPUT_CARDS_DIR_DEFAULT
        if not output_cards_dir:
            # Если константа пуста, сохраняем в папку программы (прежнее поведение)
            output_cards_dir = script_dir
            logger.info(f"Константа OUTPUT_CARDS_DIR_DEFAULT пуста, карты будут сохранены в папку программы: {output_cards_dir}")
        else:
            logger.info(f"Используется путь по умолчанию: {output_cards_dir}")

    output_cards_path = Path(output_cards_dir)
    try:
        # Создаём папку, если её нет (даже если это script_dir, ничего страшного)
        output_cards_path.mkdir(parents=True, exist_ok=True)
        logger.info(f"Папка для сохранения карт: {output_cards_path}")
    except Exception as e:
        logger.error(f"Не удалось создать/проверить папку {output_cards_path}: {e}")
        logger.info("Операция прервана из-за проблем с папкой сохранения.")
        return

    # Проверка существования шаблона карты
    if not template_path.exists():
        logger.error(f"Шаблон карты не найден: {template_path}")
        logger.info("Операция прервана из-за отсутствия шаблона.")
        return
    logger.info(f"Шаблон карты найден: {template_path}")

    # Поиск файлов .docx и .doc
    all_files = list(postanovleniya_path.glob("*"))
    doc_files = [f for f in all_files if f.suffix.lower() in ('.docx', '.doc')]
    if not doc_files:
        logger.warning(f"В папке {postanovleniya_path} не найдено файлов .docx или .doc")
        logger.info("Операция прервана из-за отсутствия постановлений.")
        return
    logger.info(f"Найдено {len(doc_files)} файлов постановлений")

    current_number = start_number
    for i, doc_file in enumerate(doc_files, 1):
        logger.info(f"Обработка файла {i}/{len(doc_files)}: {doc_file.name}")
        try:
            extractor = PostanovlenieExtractor(doc_file)
            fio_data = extractor.extract_fio_birth()
            if not fio_data:
                logger.warning(f"Пропуск файла {doc_file.name} (не найдены ФИО/дата рождения)")
                continue

            # Получаем извлечённые данные
            surname = fio_data['surname']
            first_name = fio_data['first_name']
            patronymic = fio_data['patronymic']
            birth_date = fio_data['birth_date']
            case = fio_data['case']

            # Создаём экземпляр заполнителя карты
            card_filler = ExcelCardFiller(template_path)

            # Разбиваем каждую часть на основную и альтернативу
            main_surname, alt_surname = split_name_part_with_brackets(surname)
            main_first, alt_first = split_name_part_with_brackets(first_name)
            main_patr, alt_patr = split_name_part_with_brackets(patronymic)

            logger.info(f"Разбор скобок: Фамилия: осн='{main_surname}', alt='{alt_surname}'")
            logger.info(f"Разбор скобок: Имя: осн='{main_first}', alt='{alt_first}'")
            logger.info(f"Разбор скобок: Отчество: осн='{main_patr}', alt='{alt_patr}'")

            # Преобразование падежей (работаем только с основной частью)
            if case == 'genitive':
                # Родительный падеж – исходные строки уже в родительном
                surname_gen_final = format_name_part_with_alternative(main_surname, alt_surname)
                first_name_gen_final = format_name_part_with_alternative(main_first, alt_first)
                patronymic_gen_final = format_name_part_with_alternative(main_patr, alt_patr)

                # Преобразуем основную часть в именительный
                main_surname_nom = card_filler.convert_to_nominative(main_surname)
                main_first_nom = card_filler.convert_to_nominative(main_first)
                main_patr_nom = card_filler.convert_to_nominative(main_patr)

                surname_nom_final = format_name_part_with_alternative(main_surname_nom, alt_surname)
                first_name_nom_final = format_name_part_with_alternative(main_first_nom, alt_first)
                patronymic_nom_final = format_name_part_with_alternative(main_patr_nom, alt_patr)

                logger.info(f"Родительный (с альт.): {surname_gen_final} {first_name_gen_final} {patronymic_gen_final}")
                logger.info(f"Именительный (с альт.): {surname_nom_final} {first_name_nom_final} {patronymic_nom_final}")

            elif case == 'nominative':
                # Именительный падеж – исходные строки уже в именительном
                surname_nom_final = format_name_part_with_alternative(main_surname, alt_surname)
                first_name_nom_final = format_name_part_with_alternative(main_first, alt_first)
                patronymic_nom_final = format_name_part_with_alternative(main_patr, alt_patr)

                # Преобразуем основную часть в родительный
                main_surname_gen = card_filler.convert_to_genitive(main_surname)
                main_first_gen = card_filler.convert_to_genitive(main_first)
                main_patr_gen = card_filler.convert_to_genitive(main_patr)

                surname_gen_final = format_name_part_with_alternative(main_surname_gen, alt_surname)
                first_name_gen_final = format_name_part_with_alternative(main_first_gen, alt_first)
                patronymic_gen_final = format_name_part_with_alternative(main_patr_gen, alt_patr)

                logger.info(f"Именительный (с альт.): {surname_nom_final} {first_name_nom_final} {patronymic_nom_final}")
                logger.info(f"Родительный (с альт.): {surname_gen_final} {first_name_gen_final} {patronymic_gen_final}")

            else:
                logger.error(f"Неизвестный падеж '{case}' в файле {doc_file.name}, пропуск")
                continue

            # Формируем имя файла карты (всегда в именительном падеже, с альтернативами)
            new_card_name = f"{current_number}-26 Тобольск {surname_nom_final} {first_name_nom_final} {patronymic_nom_final}.xlsx"
            # Передаём путь для сохранения карты
            new_card_path = card_filler.create_copy(new_card_name, output_cards_path)

            # Загружаем книгу
            wb = openpyxl.load_workbook(new_card_path)

            # Получаем текущую дату для ДН и МН
            current_day, current_month_gen = get_current_date()
            current_day = format_two_digits(current_day)  # форматируем

            # Извлекаем дату постановления
            resolution_date = extractor.extract_resolution_date()
            if not resolution_date:
                resolution_date = {'day': '', 'month': ''}
            resolution_day = format_two_digits(resolution_date.get('day', ''))
            resolution_month = resolution_date.get('month', '')

            # Формируем словарь для заполнения (включаем оба набора падежей)
            fill_data = {
                "НОМ": str(current_number),
                "ФР": surname_gen_final,
                "ИР": first_name_gen_final,
                "ОР": patronymic_gen_final,
                "ФИ": surname_nom_final,
                "ИИ": first_name_nom_final,
                "ОИ": patronymic_nom_final,
                "ДР": format_date_ddmmyyyy(birth_date),
                "ДП": resolution_day,
                "МП": resolution_month,
                "ДН": current_day,
                "МН": current_month_gen,
            }

            investigator_data = find_investigator(extractor.text)
            if investigator_data:
                fill_data.update(investigator_data)
                logger.info(f"Добавлены данные следователя: {investigator_data}")
            else:
                logger.info("Данные следователя не добавлены")

            # Заполняем карту
            card_filler.fill_card(wb, fill_data)
            wb.save(new_card_path)
            logger.info(f"Карта успешно создана и заполнена: {new_card_path}")

            current_number += 1

        except Exception as e:
            logger.error(f"Ошибка при обработке файла {doc_file.name}: {e}", exc_info=True)
            continue


def process_bone_cards(script_dir: Path) -> None:
    """
    Заполнение карт Excel значениями ДО, МО и инкрементом ОБ.
    """
    logger.info("Запущен режим заполнения Карт с костями")

    # Запрос значений у пользователя
    print("\n--- Заполнение Карт с костями ---")
    do_value_raw = input("Число окончания экспертизы (например, 01): ").strip()
    mo_value = input("Месяц окончания экспертизы в р.п. (например, января): ").strip()

    # Применяем форматирование с ведущим нулём для дня (как в других режимах)
    do_value = format_two_digits(do_value_raw)
    logger.info(f"День окончания экспертизы: введено '{do_value_raw}', после форматирования: '{do_value}'")
    if do_value_raw != do_value:
        logger.info("Добавлен ведущий ноль для дня окончания.")

    while True:
        ob_start = input("Начальный объект (например, 1A): ").strip()
        is_valid, _, _ = validate_ob_format(ob_start)
        if is_valid:
            break
        print("Неверный формат! Введите в формате 'числоБуква' (например, 1A)")

    logger.info(f"Введённые значения: ДО='{do_value}', МО='{mo_value}', ОБ начало='{ob_start}'")

    # Проверка формата ОБ (повторно, чтобы получить start_num и letter)
    is_valid, start_num, letter = validate_ob_format(ob_start)
    if not is_valid:
        logger.error(f"Неверный формат ОБ: '{ob_start}'. Обработка прервана.")
        return

    # Получение списка Excel-файлов в текущей папке программы
    excel_files = list(script_dir.glob("*.xlsx")) + list(script_dir.glob("*.xlsm"))
    # Исключаем временные файлы
    excel_files = [f for f in excel_files if not f.name.startswith('~$')]

    if not excel_files:
        logger.warning("В текущей папке не найдено Excel-файлов")
        return

    # Натуральная сортировка
    excel_files.sort(key=lambda x: natural_sort_key(x.name))

    logger.info(f"Найдено файлов для обработки: {len(excel_files)}")
    logger.info(f"Порядок обработки: {[f.name for f in excel_files]}")

    # Обработка каждого файла
    for i, file_path in enumerate(excel_files):
        try:
            logger.info(f"Обработка файла: {file_path.name}")

            # Загрузка книги
            wb = openpyxl.load_workbook(file_path)
            ws = wb.active

            # 1. ДО
            do_cell = find_cell_by_code(ws, "ДО")
            if do_cell:
                do_cell.value = do_value
                logger.info(f"  ДО: установлено '{do_value}'")
            else:
                logger.warning(f"  Код 'ДО' не найден в файле {file_path.name}")

            # 2. МО
            mo_cell = find_cell_by_code(ws, "МО")
            if mo_cell:
                mo_cell.value = mo_value
                logger.info(f"  МО: установлено '{mo_value}'")
            else:
                logger.warning(f"  Код 'МО' не найден в файле {file_path.name}")

            # 3. ОБ с инкрементом
            current_ob = f"{start_num + i}{letter}"
            ob_cell = find_cell_by_code(ws, "ОБ")
            if ob_cell:
                ob_cell.value = current_ob
                logger.info(f"  ОБ: установлено '{current_ob}'")
            else:
                logger.warning(f"  Код 'ОБ' не найден в файле {file_path.name}")

            # 4. ОБРТ = НОМ + пробел + ОБ
            obrt_cell = find_cell_by_code(ws, "ОБРТ")
            if obrt_cell:
                nom_cell = find_cell_by_code(ws, "НОМ")
                ob_cell = find_cell_by_code(ws, "ОБ")
                if nom_cell and nom_cell.value and ob_cell and ob_cell.value:
                    obrt_value = f"{nom_cell.value} {ob_cell.value}"
                    obrt_cell.value = obrt_value
                    logger.info(f"  ОБРТ: установлено '{obrt_value}'")
                else:
                    logger.warning(f"  Не удалось получить НОМ или ОБ для ОБРТ в файле {file_path.name}")
            else:
                logger.warning(f"  Код 'ОБРТ' не найден в файле {file_path.name}")

            wb.save(file_path)
            logger.info(f"Файл {file_path.name} успешно сохранён")

        except Exception as e:
            logger.error(f"Ошибка при обработке {file_path.name}: {e}", exc_info=True)
            continue

    logger.info("Обработка всех файлов завершена")


def process_registry_cards(script_dir: Path) -> None:
    """
    Режим «Регистратура: заполнение Карт с образцами в случае прямой идентификации».
    Заполняет все Excel-файлы в папке программы заданными значениями,
    выполняет копирование (ФПИ->ФИ и т.д.), преобразование в родительный падеж
    (ФПИ->ФР и т.д.) и удаляет значение из строки РОД.
    """
    logger.info("Запущен режим «Регистратура: заполнение Карт с образцами»")

    # 1. Выбор региона
    print("\nС какими образцами работаем?")
    print("1. Тюмень")
    print("2. Ростов")
    region_choice = input("Введите номер варианта (1 или 2): ").strip()
    while region_choice not in ('1', '2'):
        print("Неверный выбор. Введите 1 или 2.")
        region_choice = input("Введите номер варианта (1 или 2): ").strip()

    if region_choice == '1':
        code_1_value = "СВО_Молов_образец_прямая идентификация"
        logger.info("Выбран регион: Тюмень")
    else:
        code_1_value = "СВО_Ростов_образец_прямая идентификация"
        logger.info("Выбран регион: Ростов")

    data_to_fill = {
        "1": code_1_value,
        "ТО": "крови",
        "НО": "марле",
        "НАД": "с отпечатанной на принтере надписью",
        "МАТ": "бурого пятна на марле"
    }

    # 2. Поиск Excel-файлов
    excel_files = list(script_dir.glob("*.xlsx")) + list(script_dir.glob("*.xlsm"))
    excel_files = [f for f in excel_files if not f.name.startswith('~$')]
    if not excel_files:
        logger.warning("В текущей папке не найдено Excel-файлов (карт)")
        return

    excel_files.sort(key=lambda x: natural_sort_key(x.name))
    logger.info(f"Найдено файлов для обработки: {len(excel_files)}")
    logger.info(f"Порядок обработки: {[f.name for f in excel_files]}")

    # --- СОЗДАНИЕ ЭКЗЕМПЛЯРА ДЛЯ ПРЕОБРАЗОВАНИЯ ПАДЕЖЕЙ ---
    # Берём первый существующий Excel-файл, чтобы создать экземпляр класса
    # (сам файл не используется, только методы преобразования)
    template_path = excel_files[0]
    card_filler = ExcelCardFiller(template_path)
    logger.info(f"Создан экземпляр ExcelCardFiller на основе файла {template_path.name}")

    for file_path in excel_files:
        try:
            logger.info(f"Обработка файла: {file_path.name}")
            wb = openpyxl.load_workbook(file_path)
            ws = wb.active

            # 2.1 Заполнение основными значениями (используем улучшенную глобальную функцию)
            for code, value in data_to_fill.items():
                cell = find_cell_by_code(ws, code)
                if cell:
                    cell.value = value
                    logger.info(f"  Код '{code}': установлено значение '{value}'")
                else:
                    logger.warning(f"  Код '{code}' не найден в файле {file_path.name}")

            # Собираем все ячейки с кодами для быстрого доступа (ключи – строки без пробелов)
            cells = {}
            for row in ws.iter_rows(min_row=1, max_col=3):
                code_cell = row[0]
                if code_cell and code_cell.value is not None:
                    code_str = str(code_cell.value).strip()
                    if code_str:
                        cells[code_str] = row[2]  # ячейка значения

            # 2.2 Копирование
            copy_pairs = [('ФПИ', 'ФИ'), ('ИПИ', 'ИИ'), ('ОПИ', 'ОИ'), ('ДРП', 'ДР')]
            for src, dst in copy_pairs:
                if src in cells and dst in cells:
                    src_val = cells[src].value
                    if src_val is not None:
                        cells[dst].value = src_val
                        logger.info(f"      Скопировано {src} -> {dst}: '{src_val}'")
                    else:
                        logger.info(f"      Значение {src} пустое, копирование пропущено")
                else:
                    missing = [x for x in (src, dst) if x not in cells]
                    logger.warning(f"      Коды {', '.join(missing)} не найдены, копирование пропущено")

            # 2.3 Преобразование в родительный падеж (используем метод экземпляра)
            genitive_pairs = [('ФПИ', 'ФР'), ('ИПИ', 'ИР'), ('ОПИ', 'ОР')]
            for src, dst in genitive_pairs:
                if src in cells and dst in cells:
                    src_val = cells[src].value
                    if src_val and isinstance(src_val, str):
                        gen = card_filler.convert_to_genitive(src_val.strip())
                        cells[dst].value = gen
                        logger.info(f"      Преобразовано {src} -> {dst}: '{src_val}' -> '{gen}'")
                    else:
                        logger.info(f"      Значение {src} пустое или не строка, преобразование пропущено")
                else:
                    missing = [x for x in (src, dst) if x not in cells]
                    logger.warning(f"      Коды {', '.join(missing)} не найдены, преобразование пропущено")

            # 2.4 Удаление значения из РОД
            if 'РОД' in cells:
                rod_cell = cells['РОД']
                if rod_cell.value is not None:
                    logger.info(f"      Удалено значение из РОД: '{rod_cell.value}'")
                    rod_cell.value = None
                else:
                    logger.info("      Значение РОД уже пустое")
            else:
                logger.warning("      Код РОД не найден, удаление пропущено")

            wb.save(file_path)
            logger.info(f"Файл {file_path.name} успешно сохранён")

        except PermissionError:
            logger.error(f"Не удалось сохранить {file_path.name}: файл открыт в другой программе")
        except Exception as e:
            logger.error(f"Ошибка при обработке {file_path.name}: {e}", exc_info=True)
            continue

    logger.info("Обработка всех файлов завершена")


def process_distribution_cards(script_dir: Path) -> None:
    """
    Режим «Регистратура: Заполнение Карт с костями для РАСПРЕДЕЛЕНИЯ между экспертами».
    Поиск папок экспертов, в каждой папке чтение .txt-файла и создание Excel-карт.
    """
    logger.info("Запущен режим распределения карт (на основе .txt-файлов в папках экспертов)")

    # 1. Находим шаблон Excel
    template_path = find_template_excel(script_dir)
    if not template_path:
        logger.error("Не удалось найти шаблон карты")
        input("Нажмите Enter для завершения...")
        return

    # 2. Запрос начального номера экспертизы
    while True:
        try:
            start_num_str = input("Введите номер, с которого начать нумерацию экспертиз: ").strip()
            start_number = int(start_num_str)
            if start_number <= 0:
                raise ValueError
            break
        except ValueError:
            print("Ошибка: пожалуйста, введите целое положительное число")
            logger.warning("Пользователь ввёл неверный номер экспертизы")

    # 3. Запрос даты начала экспертизы
    while True:
        date_str = input("Введите дату начала экспертизы в формате ДД.ММ.ГГГГ: ").strip()
        try:
            if not re.match(r'^\d{2}\.\d{2}\.\d{4}$', date_str):
                raise ValueError
            day, month, year = map(int, date_str.split('.'))
            # Проверка валидности даты
            datetime(year, month, day)
            start_day = format_two_digits(day)          # форматируем
            start_month_gen = month_num_to_genitive(month)
            break
        except (ValueError, TypeError):
            print("Ошибка: неверный формат даты. Используйте формат ДД.ММ.ГГГГ")
            logger.warning("Пользователь ввёл неверную дату начала")

    logger.info(f"Параметры: начальный номер={start_number}, дата начала={start_day}.{month:02d}.{year}")

    # 4. Получаем список папок (экспертов) в корневой папке
    expert_folders = [d for d in script_dir.iterdir() if d.is_dir()]
    if not expert_folders:
        logger.warning("В корневой папке не найдено папок экспертов")
        input("Нажмите Enter для завершения...")
        return

    expert_folders.sort(key=lambda x: x.name)  # сортируем по имени
    logger.info(f"Найдено папок экспертов: {len(expert_folders)}")
    logger.info(f"Папки: {', '.join(f.name for f in expert_folders)}")

    # Подтверждение
    confirm = input(f"\nБудет создано несколько Excel файлов, начиная с номера {start_number}.\n"
                    f"Дата начала экспертизы: {start_day}.{month:02d}.{year}\n"
                    f"Папки для обработки: {len(expert_folders)}\n\n"
                    "Продолжить? (да/нет): ")
    if confirm.lower() not in ['да', 'д', 'yes', 'y']:
        logger.info("Операция отменена пользователем")
        return

    current_number = start_number
    total_files_created = 0
    results = []

    for folder in expert_folders:
        # Ищем первый .txt файл в папке
        txt_files = list(folder.glob("*.txt"))
        if not txt_files:
            logger.warning(f"В папке '{folder.name}' не найден .txt файл, пропускаем")
            results.append(f"Папка '{folder.name}': нет .txt файла")
            continue

        txt_file = txt_files[0]
        rows = parse_txt_file(txt_file)
        if not rows:
            logger.warning(f"В файле '{txt_file.name}' (папка '{folder.name}') нет данных")
            results.append(f"Папка '{folder.name}': файл '{txt_file.name}' не содержит данных")
            continue

        folder_files_created = 0
        for row in rows:
            # Проверка формата даты в строке
            try:
                date_parts = row['date'].split('.')
                if len(date_parts) != 3:
                    raise ValueError
                day_from_row = int(date_parts[0])
                month_from_row = int(date_parts[1])
                year_from_row = int(date_parts[2])
                # Проверка валидности
                datetime(year_from_row, month_from_row, day_from_row)
                month_text = month_num_to_genitive(month_from_row)
                day_from_row = format_two_digits(day_from_row)    # форматируем
            except (ValueError, TypeError):
                logger.error(f"Неверный формат даты '{row['date']}' в файле {txt_file.name}, строка пропущена")
                continue

            # Данные для заполнения
            expertise_data = {
                'expertise_number': str(current_number),
                'corpse_number': row['corpse_number'],
                'day': day_from_row,
                'month': month_text,
                'start_day': start_day,
                'start_month': start_month_gen
            }

            output_filename = f"{current_number}.xlsx"
            output_path = folder / output_filename

            if fill_excel_from_template(template_path, output_path, expertise_data):
                folder_files_created += 1
                total_files_created += 1
                logger.info(f"Создан файл: {output_path} (номер={current_number}, труп={row['corpse_number']})")
            else:
                logger.error(f"Не удалось создать файл {output_path}")

            current_number += 1

        results.append(f"Папка '{folder.name}': создано {folder_files_created} файлов")

    # Вывод статистики
    print("\n" + "=" * 50)
    print("РЕЗУЛЬТАТЫ:")
    print("=" * 50)
    print(f"Всего создано файлов: {total_files_created}")
    print(f"Последний использованный номер: {current_number - 1}")
    print("\nДетали по папкам:")
    for detail in results:
        print(f"  • {detail}")

    # Сохранение лога
    log_file = script_dir / "distribution_log.txt"
    try:
        with open(log_file, 'w', encoding='utf-8') as f:
            f.write(f"Дата обработки: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}\n")
            f.write(f"Начальный номер: {start_number}\n")
            f.write(f"Дата начала экспертизы: {start_day}.{month:02d}.{year}\n")
            f.write(f"Всего создано файлов: {total_files_created}\n")
            f.write(f"Последний номер: {current_number - 1}\n\n")
            f.write("Детали:\n")
            for detail in results:
                f.write(f"{detail}\n")
        logger.info(f"Лог сохранён в {log_file}")
    except Exception as e:
        logger.error(f"Не удалось сохранить лог: {e}")

    input("\nНажмите Enter для завершения...")


def process_mass_fill_cards_with_samples(script_dir: Path) -> None:
    """
    Режим «Регистратура: Массовое заполнение Карт с ОБРАЗЦАМИ».
    Реализует логику первого кода: удаление пробелов в строках с заданными кодами,
    поиск ФР, ИР, ОР, преобразование их в именительный падеж и запись в ФИ, ИИ, ОИ.
    """
    logger.info("Запущен режим массового заполнения карт с образцами (преобразование ФИО из родительного в именительный)")

    # Коды, для которых нужно удалить лишние пробелы
    codes_for_space_removal = ['ФР', 'ИР', 'ОР', 'ФОР', 'ИОР', 'ООР', 'ФПИ', 'ИПИ', 'ОПИ']

    # Поиск Excel-файлов (только .xlsx и .xlsm)
    excel_files = list(script_dir.glob("*.xlsx")) + list(script_dir.glob("*.xlsm"))
    excel_files = [f for f in excel_files if not f.name.startswith('~$')]

    if not excel_files:
        logger.warning("В текущей папке не найдено Excel-файлов (карт)")
        return

    excel_files.sort(key=lambda x: natural_sort_key(x.name))
    logger.info(f"Найдено файлов для обработки: {len(excel_files)}")
    logger.info(f"Порядок обработки: {[f.name for f in excel_files]}")

    # --- СОЗДАНИЕ ЭКЗЕМПЛЯРА ДЛЯ ПРЕОБРАЗОВАНИЯ ПАДЕЖЕЙ ---
    # Берём первый существующий Excel-файл, чтобы создать экземпляр класса
    template_path = excel_files[0]
    card_filler = ExcelCardFiller(template_path)
    logger.info(f"Создан экземпляр ExcelCardFiller на основе файла {template_path.name}")

    processed_count = 0
    error_count = 0
    skipped_count = 0

    for file_path in excel_files:
        logger.info(f"Обработка файла: {file_path.name}")
        try:
            wb = openpyxl.load_workbook(file_path)
            ws = wb.active
            changes_made = False
            space_removal_changes_made = False

            # ---------- 1. Удаление лишних пробелов ----------
            logger.debug("Удаление лишних пробелов в строках с кодами: %s", codes_for_space_removal)
            for row in ws.iter_rows(min_row=1, max_col=3):
                code_cell = row[0]
                value_cell = row[2] if len(row) > 2 else None
                if code_cell is None or code_cell.value is None:
                    continue
                code = normalize_text(str(code_cell.value))
                if code in codes_for_space_removal and value_cell is not None:
                    original_value = value_cell.value
                    stripped_value = normalize_text(original_value)
                    if original_value != stripped_value:
                        value_cell.value = stripped_value
                        logger.info(f"Удалены лишние пробелы в строке с кодом '{code}': '{original_value}' -> '{stripped_value}'")
                        space_removal_changes_made = True
                    else:
                        logger.debug(f"Лишние пробелы не найдены в строке с кодом '{code}': '{original_value}'")

            # ---------- 2. Поиск ФР, ИР, ОР ----------
            fr_value = None
            ir_value = None
            or_value = None
            for row in ws.iter_rows(min_row=1, max_col=3):
                code_cell = row[0]
                value_cell = row[2] if len(row) > 2 else None
                if code_cell is None or code_cell.value is None:
                    continue
                code = normalize_text(str(code_cell.value))
                if code == 'ФР' and value_cell is not None:
                    fr_value = value_cell.value
                    logger.info(f"Найдено ФР: '{fr_value}'")
                elif code == 'ИР' and value_cell is not None:
                    ir_value = value_cell.value
                    logger.info(f"Найдено ИР: '{ir_value}'")
                elif code == 'ОР' and value_cell is not None:
                    or_value = value_cell.value
                    logger.info(f"Найдено ОР: '{or_value}'")

            # ---------- 3. Преобразование в именительный падеж (используем метод экземпляра) ----------
            nominative_values = {}
            if fr_value is not None:
                fi = card_filler.convert_to_nominative(fr_value)
                nominative_values['ФИ'] = fi
                if fi != fr_value:
                    logger.info(f"Фамилия преобразована: '{fr_value}' -> '{fi}'")
                else:
                    logger.info(f"Фамилия не требует преобразования: '{fr_value}'")
            if ir_value is not None:
                ii = card_filler.convert_to_nominative(ir_value)
                nominative_values['ИИ'] = ii
                if ii != ir_value:
                    logger.info(f"Имя преобразовано: '{ir_value}' -> '{ii}'")
                else:
                    logger.info(f"Имя не требует преобразования: '{ir_value}'")
            if or_value is not None:
                oi = card_filler.convert_to_nominative(or_value)
                nominative_values['ОИ'] = oi
                if oi != or_value:
                    logger.info(f"Отчество преобразовано: '{or_value}' -> '{oi}'")
                else:
                    logger.info(f"Отчество не требует преобразования: '{or_value}'")

            # ---------- 4. Запись в ячейки с кодами ФИ, ИИ, ОИ ----------
            if nominative_values:
                for code, value in nominative_values.items():
                    cell = find_cell_by_code(ws, code)
                    if cell:
                        cell.value = value
                        logger.info(f"Записано {code}: '{value}'")
                        changes_made = True
                    else:
                        logger.warning(f"Код {code} не найден в файле {file_path.name}, значение не записано")
            else:
                logger.info("Не найдено ни одной части ФИО в родительном падеже для обработки")

            # ---------- 5. Сохранение, если были изменения ----------
            if changes_made or space_removal_changes_made:
                wb.save(file_path)
                logger.info(f"Файл {file_path.name} сохранён")
                processed_count += 1
            else:
                logger.info(f"Файл {file_path.name} не требует обработки (нет изменений)")
                skipped_count += 1

        except Exception as e:
            logger.error(f"Ошибка при обработке файла {file_path.name}: {e}", exc_info=True)
            error_count += 1

    # Итоговая статистика
    logger.info("=" * 50)
    logger.info("ИТОГИ ОБРАБОТКИ (массовое заполнение карт с образцами):")
    logger.info(f"Всего файлов найдено: {len(excel_files)}")
    logger.info(f"Успешно обработано (с изменениями): {processed_count}")
    logger.info(f"Пропущено (без изменений): {skipped_count}")
    logger.info(f"Ошибок: {error_count}")
    logger.info("=" * 50)


def process_mass_print_postanovleniy():
    """Режим «Регистратура: Массовая печать постановлений»."""
    logger.info("Запущен режим массовой печати постановлений")
    logger.info(f"Папка для сохранения постановлений: {MASSPRINT_SAVEPATH_DEFAULT}")
    logger.info(f"Папка с шаблонами: {TEMPLATES_FOLDER_DEFAULT}")

    # ---- Вспомогательные функции ----
    def get_excel_files(folder_path):
        """Получает список файлов Excel в указанной папке, отсортированных естественным образом."""
        excel_files = [f for f in os.listdir(folder_path) if f.endswith(('.xlsx', '.xls')) and not f.startswith('~$')]
        excel_files.sort(key=natural_sort_key)
        logger.info(f"Найдено файлов Excel: {len(excel_files)}")
        for i, f in enumerate(excel_files, 1):
            logger.info(f"  {i}. {f}")
        return [os.path.join(folder_path, f) for f in excel_files]

    def read_card(file_path):
        """Читает данные из карты Excel, используя поиск по кодам (аналог find_cell_by_code)."""
        data = {}
        try:
            wb = openpyxl.load_workbook(file_path)
            ws = wb.active
            for row in ws.iter_rows(min_row=1, max_col=3):
                code_cell = row[0]
                value_cell = row[2] if len(row) > 2 else None
                if code_cell is None or code_cell.value is None:
                    continue
                code = str(code_cell.value).strip()
                if not code:
                    continue
                value = value_cell.value if value_cell is not None else None
                # Специальная обработка для кода "ДБИ"
                if code == "ДБИ" and value is not None and str(value).strip():
                    value = f"с {value.strip()} г."
                data[code] = value
            logger.debug(f"Из файла {file_path} прочитаны коды: {list(data.keys())}")
        except Exception as e:
            logger.error(f"Ошибка при чтении файла {file_path}: {e}", exc_info=True)
            raise
        return data

    def check_card_data(card_data):
        """Проверяет наличие обязательных полей."""
        required_fields = ["2", "3", "4"]
        for field in required_fields:
            if field not in card_data or card_data[field] is None:
                raise ValueError(f"Ошибка: в Карте отсутствует значение для кода '{field}'.")

    def copy_template(card_data, template_name, save_path, template_folder):
        """Копирует шаблон Word в указанное место."""
        for ext in [".doc", ".docx"]:
            template_path = template_folder / f"{template_name}{ext}"
            if template_path.exists():
                output_path = Path(save_path) / f"{card_data['ФИ']}-26{ext}"
                shutil.copy(str(template_path), str(output_path))
                return str(output_path)
        raise FileNotFoundError(f"Файл шаблона '{template_name}' не найден в папке {template_folder}.")

    def replace_in_doc(doc_path, replacements):
        """Заменяет метки вида {код} в Word-документе."""
        doc = Document(doc_path)
        found_keys = set()

        def replace_text(paragraph):
            nonlocal found_keys
            text = paragraph.text
            if not paragraph.runs:
                return
            first_run = paragraph.runs[0]
            base_formatting = {
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline,
                'font': first_run.font.name,
                'size': first_run.font.size
            }
            modified_text = text
            any_replacements = False
            for key, value in replacements.items():
                patterns = [f"{{{key}}}", f"{{ {key} }}", f"{{{key} }}", f"{{ {key}}}"]
                for pattern in patterns:
                    if pattern in modified_text:
                        modified_text = modified_text.replace(pattern, str(value))
                        found_keys.add(key)
                        any_replacements = True
            if any_replacements:
                for run in paragraph.runs:
                    run.text = ""
                new_run = paragraph.add_run(modified_text)
                new_run.bold = base_formatting['bold']
                new_run.italic = base_formatting['italic']
                new_run.underline = base_formatting['underline']
                new_run.font.name = base_formatting['font']
                new_run.font.size = base_formatting['size']

        for p in doc.paragraphs:
            replace_text(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text(p)
        for section in doc.sections:
            for p in section.header.paragraphs:
                replace_text(p)
            for p in section.footer.paragraphs:
                replace_text(p)

        missing_keys = set(replacements.keys()) - found_keys
        if missing_keys:
            logger.warning(f"Предупреждение: следующие метки не найдены в документе: {missing_keys}")

        doc.save(doc_path)
        logger.info(f"Все метки заменены в документе: {doc_path}")

    def print_document(doc_path):
        """Печатает документ Word."""
        try:
            printer_name = win32print.GetDefaultPrinter()
            logger.info(f"ОТПРАВКА НА ПЕЧАТЬ: {os.path.basename(doc_path)}")
            time.sleep(3)
            win32api.ShellExecute(0, "print", doc_path, None, ".", 0)
            logger.info(f"Документ {doc_path} отправлен на печать на принтере {printer_name}.")
            time.sleep(3)
        except Exception as e:
            logger.error(f"Ошибка при печати документа {doc_path}: {e}", exc_info=True)

    def process_multiple_cards(folder_path, template_choice):
        """Обрабатывает все файлы Excel в указанной папке по очереди."""
        excel_files = get_excel_files(folder_path)
        logger.info(f"=== НАЧАЛО ОБРАБОТКИ {len(excel_files)} ФАЙЛОВ ===")

        for index, card_path in enumerate(excel_files, 1):
            output_path = None
            try:
                logger.info(f"=== ОБРАБОТКА ФАЙЛА {index}/{len(excel_files)}: {os.path.basename(card_path)} ===")

                card_data = read_card(card_path)
                check_card_data(card_data)

                # Replacements создаётся для всех кодов, даже пустых
                empty_keys = [code for code, value in card_data.items() if value is None]
                if empty_keys:
                    logger.info(f"Следующие коды имеют пустые значения и будут заменены на пустую строку: {empty_keys}")

                replacements = {code: str(value) if value is not None else "" for code, value in card_data.items()}

                # Выбор шаблона и пути сохранения
                if template_choice == 1:
                    template_name = "Постановление экспертиза ФОНД"
                elif template_choice == 2:
                    template_name = "Постановление экспертиза ФОНД прямая идентификация"
                elif template_choice == 3:
                    template_name = "Постановление экспертиза ФОНД БВП прямая идентификация"
                    logger.info("Выбран шаблон для БВП (без вести пропавших)")
                else:
                    raise ValueError("Неверный выбор шаблона.")

                save_path = MASSPRINT_SAVEPATH_DEFAULT
                save_path_obj = Path(save_path)
                if not save_path_obj.exists():
                    save_path_obj.mkdir(parents=True, exist_ok=True)
                    logger.info(f"Создана папка для сохранения: {save_path_obj}")

                template_folder = Path(TEMPLATES_FOLDER_DEFAULT)
                if not template_folder.exists():
                    logger.error(f"Папка с шаблонами не найдена: {template_folder}")
                    continue

                # Проверка существования файла шаблона
                template_exists = False
                found_ext = None
                for ext in [".doc", ".docx"]:
                    template_path = template_folder / f"{template_name}{ext}"
                    if template_path.exists():
                        template_exists = True
                        found_ext = ext
                        break
                if not template_exists:
                    logger.error(f"Файл шаблона '{template_name}' (с расширением .doc или .docx) не найден в папке {template_folder}")
                    continue
                logger.info(f"Шаблон найден: {template_name}{found_ext}")

                output_path = copy_template(card_data, template_name, save_path, template_folder)
                logger.info(f"Документ создан: {output_path}")

                replace_in_doc(output_path, replacements)

                print_document(output_path)

                logger.info(f"ФАЙЛ {index}/{len(excel_files)} УСПЕШНО ОБРАБОТАН: {os.path.basename(card_path)}")

            except Exception as e:
                logger.error(f"ОШИБКА при обработке файла {card_path}: {e}", exc_info=True)
            finally:
                if output_path:
                    logger.info(f"Итоговый документ: {output_path}")

        logger.info(f"=== ЗАВЕРШЕНИЕ ОБРАБОТКИ ВСЕХ ФАЙЛОВ ===")

    def choose_template():
        """Выбор типа шаблона пользователем."""
        print("Распечатать постановления по:")
        print("1 - родству")
        print("2 - прямой идентификации")
        print("3 - БВП (без вести пропавшим)")
        choice = input("Введите номер шаблона: ")
        try:
            choice = int(choice)
            if choice not in [1, 2, 3]:
                raise ValueError
            logger.info(f"Выбран шаблон №{choice}")
            return choice
        except ValueError:
            print("Неверный выбор. Пожалуйста, введите 1, 2 или 3.")
            return choose_template()

    # ---- Основная логика режима ----
    try:
        folder_path = Path.cwd()
        excel_files = get_excel_files(str(folder_path))
        if not excel_files:
            raise FileNotFoundError(f"В папке {folder_path} не найдено файлов Excel (.xlsx или .xls).")

        card_path = excel_files[0]
        card_data = read_card(card_path)
        logger.info(f"Обновленные данные из файла: {card_data}")

        template_choice = choose_template()
        process_multiple_cards(str(folder_path), template_choice)

    except Exception as e:
        logger.error(f"Произошла ошибка в режиме массовой печати: {e}", exc_info=True)


def process_rostov_cards(script_dir: Path) -> None:
    """
    Режим «Заведующая: заполнение Карт по ОБРАЗЦАМ ИЗ РОСТОВА».
    """
    logger.info("Запущен режим «Заведующая: заполнение Карт по ОБРАЗЦАМ ИЗ РОСТОВА»")

    default_input_dir = script_dir / ROSTOV_DIR_DEFAULT
    user_input = input(f"Введите путь к папке с исходным Excel-файлом (или Enter для '{default_input_dir}'): ").strip()
    input_dir = Path(user_input) if user_input else default_input_dir

    if not input_dir.exists():
        logger.error(f"Папка не найдена: {input_dir}")
        input("Нажмите Enter для завершения...")
        return
    logger.info(f"Папка с исходными данными: {input_dir}")

    default_output_dir = input_dir
    user_output = input(f"Введите путь для сохранения карт (или Enter для '{default_output_dir}'): ").strip()
    output_dir = Path(user_output) if user_output else default_output_dir

    try:
        output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Папка для сохранения карт: {output_dir}")
    except Exception as e:
        logger.error(f"Не удалось создать папку {output_dir}: {e}")
        return

    excel_files = list(input_dir.glob("*.xlsx")) + list(input_dir.glob("*.xlsm"))
    excel_files = [f for f in excel_files if not f.name.startswith('~$')]
    if not excel_files:
        logger.error(f"В папке {input_dir} не найдено файлов Excel.")
        return

    if len(excel_files) == 1:
        source_file = excel_files[0]
        logger.info(f"Найден файл: {source_file.name}")
    else:
        print("Найдено несколько файлов Excel. Выберите нужный:")
        for i, f in enumerate(excel_files, 1):
            print(f"{i}. {f.name}")
        while True:
            try:
                choice = int(input("Введите номер: ").strip())
                if 1 <= choice <= len(excel_files):
                    source_file = excel_files[choice-1]
                    break
                else:
                    print("Неверный номер.")
            except ValueError:
                print("Введите число.")
        logger.info(f"Выбран файл: {source_file.name}")

    template_path = find_template_excel(script_dir)
    if not template_path:
        logger.error("Не удалось найти шаблон карты (Excel-файл в папке программы).")
        return

    try:
        wb_source = openpyxl.load_workbook(source_file)
        ws_source = wb_source.active
    except Exception as e:
        logger.error(f"Не удалось открыть файл {source_file}: {e}")
        return

    # Определяем заголовки
    headers = {}
    for col in range(1, ws_source.max_column + 1):
        cell_value = ws_source.cell(row=1, column=col).value
        if cell_value:
            headers[cell_value.strip()] = col

    required_columns = [
        "Фамилия, Имя, Отчество родственника",
        "Дата рождения родственника",
        "Степень родства",
        "Фамилия, Имя, Отчество пропавшего",
        "Дата рождения пропавшего",
        "Дата отправки"
    ]
    for col_name in required_columns:
        if col_name not in headers:
            logger.error(f"В файле отсутствует обязательный столбец '{col_name}'. Обработка прервана.")
            return

    col_rel_fio = headers["Фамилия, Имя, Отчество родственника"]
    col_rel_birth = headers["Дата рождения родственника"]
    col_degree = headers["Степень родства"]
    col_missing_fio = headers["Фамилия, Имя, Отчество пропавшего"]
    col_missing_birth = headers["Дата рождения пропавшего"]
    col_send_date = headers["Дата отправки"]

    total_rows = 0
    success_rows = 0
    error_rows = 0

    card_filler = ExcelCardFiller(template_path)

    for row_idx in range(2, ws_source.max_row + 1):
        try:
            total_rows += 1
            logger.info(f"Обработка строки {row_idx}")

            rel_fio = normalize_text(ws_source.cell(row=row_idx, column=col_rel_fio).value)
            rel_birth = normalize_text(ws_source.cell(row=row_idx, column=col_rel_birth).value)
            degree = normalize_text(ws_source.cell(row=row_idx, column=col_degree).value)
            missing_fio = normalize_text(ws_source.cell(row=row_idx, column=col_missing_fio).value)
            missing_birth = normalize_text(ws_source.cell(row=row_idx, column=col_missing_birth).value)
            send_date_str = normalize_text(ws_source.cell(row=row_idx, column=col_send_date).value)

            if not degree:
                logger.warning(f"Строка {row_idx} не содержит степени родства, пропущена")
                continue

            send_day = ""
            send_month = ""
            if send_date_str is not None and send_date_str != "":
                try:
                    send_dt = parse_date_flexible(send_date_str)
                    if send_dt:
                        send_day = format_two_digits(send_dt.day)
                        send_month = month_num_to_genitive(send_dt.month)
                        logger.debug(f"Дата отправки распознана: {send_dt.strftime('%Y-%m-%d')} -> день={send_day}, месяц={send_month}")
                    else:
                        logger.warning(f"Не удалось распознать дату отправки: {send_date_str}")
                except Exception as e:
                    logger.warning(f"Ошибка при разборе даты отправки: {e}")

            current_day, current_month = get_current_date()
            current_day = format_two_digits(current_day)

            normalized_degree = normalize_degree(degree)
            is_direct = normalized_degree == 'личный генотип'

            if is_direct:
                code1_value = "СВО_Ростов_образец_прямая идентификация"
                fio_to_parse = missing_fio
                birth_to_parse = missing_birth
            else:
                code1_value = "СВО_Ростов_образец_родственники"
                fio_to_parse = rel_fio
                birth_to_parse = rel_birth

            surname, first_name, patronymic = parse_full_name(fio_to_parse)

            surname = add_dot_to_initials(surname)
            first_name = add_dot_to_initials(first_name)
            patronymic = add_dot_to_initials(patronymic)

            surname_gen = card_filler.convert_to_genitive(surname) if surname else ""
            first_name_gen = card_filler.convert_to_genitive(first_name) if first_name else ""
            patronymic_gen = card_filler.convert_to_genitive(patronymic) if patronymic else ""

            file_name_raw = rel_fio if rel_fio else missing_fio
            safe_name = sanitize_filename(file_name_raw)
            card_name = f"{safe_name}.xlsx"
            card_path = output_dir / card_name

            try:
                shutil.copy2(template_path, card_path)
                logger.info(f"Создана карта: {card_path}")
            except Exception as e:
                logger.error(f"Ошибка при копировании шаблона: {e}")
                error_rows += 1
                continue

            wb_card = openpyxl.load_workbook(card_path)
            ws_card = wb_card.active

            fill_data = {
                "1": code1_value,
                "ДП": send_day,
                "МП": send_month,
                "ДН": current_day,
                "МН": current_month,
                "ТЕР": "Ростова-на-Дону",
            }

            if is_direct:
                fill_data["ФИ"] = surname
                fill_data["ИИ"] = first_name
                fill_data["ОИ"] = patronymic
                fill_data["ФПИ"] = surname
                fill_data["ИПИ"] = first_name
                fill_data["ОПИ"] = patronymic
                fill_data["ФР"] = surname_gen
                fill_data["ИР"] = first_name_gen
                fill_data["ОР"] = patronymic_gen
                fill_data["ДР"] = format_date_ddmmyyyy(missing_birth) if missing_birth else ""
                fill_data["ДРП"] = format_date_ddmmyyyy(missing_birth) if missing_birth else ""
            else:
                fill_data["ФИ"] = surname
                fill_data["ИИ"] = first_name
                fill_data["ОИ"] = patronymic
                fill_data["ФР"] = surname_gen
                fill_data["ИР"] = first_name_gen
                fill_data["ОР"] = patronymic_gen
                fill_data["ДР"] = format_date_ddmmyyyy(birth_to_parse) if birth_to_parse else ""

                missing_surname, missing_first, missing_patr = parse_full_name(missing_fio)

                missing_surname = add_dot_to_initials(missing_surname)
                missing_first = add_dot_to_initials(missing_first)
                missing_patr = add_dot_to_initials(missing_patr)

                fill_data["ФПИ"] = missing_surname
                fill_data["ИПИ"] = missing_first
                fill_data["ОПИ"] = missing_patr
                fill_data["ДРП"] = format_date_ddmmyyyy(missing_birth) if missing_birth else ""

                degree_key = normalized_degree
                rod_value = ROD_MAPPING.get(degree_key, "родство не указано")
                fill_data["РОД"] = rod_value

            # Ячейки для очистки
            clear_codes = ["ФОР", "ИОР", "ООР", "ДРО", "РОД1", "ИНД", "АТО", "ДБИ"]
            if not is_direct:
                # Для родства РОД оставляем, остальные очищаем
                for code in clear_codes:
                    fill_data[code] = None
            else:
                # Для прямой идентификации очищаем и РОД тоже
                clear_codes.append("РОД")
                for code in clear_codes:
                    fill_data[code] = None

            card_filler.fill_card(wb_card, fill_data)
            wb_card.save(card_path)
            logger.info(f"Карта {card_path} успешно заполнена")
            success_rows += 1

        except Exception as e:
            logger.error(f"Ошибка при обработке строки {row_idx}: {e}", exc_info=True)
            error_rows += 1

    logger.info("=" * 50)
    logger.info(f"Обработка завершена. Всего строк: {total_rows}, успешно: {success_rows}, ошибок: {error_rows}")
    logger.info("=" * 50)


# ========== НОВАЯ ФУНКЦИЯ ОБРАБОТКИ ==========
def process_ekc_postanovleniya(script_dir: Path, template_path: Path, start_number: int) -> None:
    """
    Обработка постановлений из ЭКЦ (дело Щегловой).
    Извлекает ФИО и дату рождения из фразы "Прошу получить генетический профиль ...",
    заполняет карты и сохраняет в папку !ПРОВЕРЕНО.
    """
    logger.info("=== Запущен режим обработки постановлений ЭКЦ ===")

    # 1. Запрос пути к папке с постановлениями
    user_input = input(f"Введите путь к папке с постановлениями ЭКЦ\n(или нажмите Enter для '{EKC_POSTANOVLENIYA_DIR_DEFAULT}'): ").strip()
    if user_input:
        postanovleniya_dir = user_input
    else:
        postanovleniya_dir = EKC_POSTANOVLENIYA_DIR_DEFAULT
        logger.info(f"Используется путь по умолчанию для постановлений: {postanovleniya_dir}")

    postanovleniya_path = Path(postanovleniya_dir)
    if not postanovleniya_path.exists():
        logger.error(f"Папка с постановлениями не найдена: {postanovleniya_path}")
        logger.info("Операция прервана.")
        return
    logger.info(f"Папка с постановлениями: {postanovleniya_path}")

    # 2. Запрос пути для сохранения карт
    user_output = input(f"Введите путь для сохранения готовых карт\n(или нажмите Enter для '{EKC_OUTPUT_DIR_DEFAULT}'): ").strip()
    if user_output:
        output_cards_dir = user_output
    else:
        output_cards_dir = EKC_OUTPUT_DIR_DEFAULT
        logger.info(f"Используется путь по умолчанию для сохранения: {output_cards_dir}")

    output_cards_path = Path(output_cards_dir)
    try:
        output_cards_path.mkdir(parents=True, exist_ok=True)
        logger.info(f"Папка для сохранения карт готова: {output_cards_path}")
    except Exception as e:
        logger.error(f"Не удалось создать папку {output_cards_path}: {e}")
        return

    # 3. Проверка шаблона карты
    if not template_path.exists():
        logger.error(f"Шаблон карты не найден: {template_path}")
        return
    logger.info(f"Шаблон карты: {template_path}")

    # 4. Поиск файлов постановлений (.docx, .doc)
    all_files = list(postanovleniya_path.glob("*"))
    doc_files = [f for f in all_files if f.suffix.lower() in ('.docx', '.doc')]
    if not doc_files:
        logger.warning(f"В папке {postanovleniya_path} не найдено файлов .docx или .doc")
        return
    logger.info(f"Найдено файлов постановлений: {len(doc_files)}")

    # 5. Обработка каждого файла
    current_number = start_number
    processed_count = 0
    skipped_count = 0

    for i, doc_file in enumerate(doc_files, 1):
        logger.info(f"=== Обработка файла {i}/{len(doc_files)}: {doc_file.name} ===")
        try:
            extractor = PostanovlenieExtractor(doc_file)
            ekc_data = extractor.extract_ekc_data()
            if not ekc_data:
                logger.warning(f"Пропуск файла {doc_file.name}: не удалось извлечь данные ЭКЦ")
                skipped_count += 1
                continue

            # Исходные строки (родительный падеж, из постановления)
            orig_surname = ekc_data['surname']
            orig_first = ekc_data['first_name']
            orig_patr = ekc_data['patronymic']
            birth_date = ekc_data['birth_date']

            # Разбиваем каждую часть на основную и альтернативу (скобки)
            main_surname, alt_surname = split_name_part_with_brackets(orig_surname)
            main_first, alt_first = split_name_part_with_brackets(orig_first)
            main_patr, alt_patr = split_name_part_with_brackets(orig_patr)

            logger.info(f"Разбор скобок (ЭКЦ): Фамилия: осн='{main_surname}', alt='{alt_surname}'")
            logger.info(f"Разбор скобок (ЭКЦ): Имя: осн='{main_first}', alt='{alt_first}'")
            logger.info(f"Разбор скобок (ЭКЦ): Отчество: осн='{main_patr}', alt='{alt_patr}'")

            # Создаём заполнитель карты (нужен для преобразования падежей)
            card_filler = ExcelCardFiller(template_path)

            # Родительный падеж (с альтернативой) – основная часть уже в родительном
            surname_gen_final = format_name_part_with_alternative(main_surname, alt_surname)
            first_name_gen_final = format_name_part_with_alternative(main_first, alt_first)
            patronymic_gen_final = format_name_part_with_alternative(main_patr, alt_patr)

            # Преобразуем основную часть в именительный
            main_surname_nom = card_filler.convert_to_nominative(main_surname)
            main_first_nom = card_filler.convert_to_nominative(main_first)
            main_patr_nom = card_filler.convert_to_nominative(main_patr)

            # Именительный падеж (с альтернативой)
            surname_nom_final = format_name_part_with_alternative(main_surname_nom, alt_surname)
            first_name_nom_final = format_name_part_with_alternative(main_first_nom, alt_first)
            patronymic_nom_final = format_name_part_with_alternative(main_patr_nom, alt_patr)

            # Добавляем точки к инициалам (ко всей итоговой строке)
            surname_gen_final = add_dot_to_initials(surname_gen_final)
            first_name_gen_final = add_dot_to_initials(first_name_gen_final)
            patronymic_gen_final = add_dot_to_initials(patronymic_gen_final)
            surname_nom_final = add_dot_to_initials(surname_nom_final)
            first_name_nom_final = add_dot_to_initials(first_name_nom_final)
            patronymic_nom_final = add_dot_to_initials(patronymic_nom_final)

            logger.info(f"Родительный (с альт. и точками): {surname_gen_final} {first_name_gen_final} {patronymic_gen_final}")
            logger.info(f"Именительный (с альт. и точками): {surname_nom_final} {first_name_nom_final} {patronymic_nom_final}")

            # Формируем имя файла карты (используем значения с точками)
            safe_fio = f"{surname_nom_final} {first_name_nom_final} {patronymic_nom_final}".strip()
            new_card_name = f"{current_number}-26 ЭКЦ {safe_fio}.xlsx"
            new_card_path = card_filler.create_copy(new_card_name, output_cards_path)

            wb = openpyxl.load_workbook(new_card_path)

            # Текущая дата
            current_day, current_month_gen = get_current_date()
            current_day = format_two_digits(current_day)

            # Дата постановления
            resolution_date = extractor.extract_resolution_date()
            if not resolution_date:
                resolution_date = {'day': '', 'month': ''}
            resolution_day = format_two_digits(resolution_date.get('day', ''))
            resolution_month = resolution_date.get('month', '')

            # Извлечение ФИО в дательном падеже
            dative_fio = extractor.extract_dative_fio()
            if dative_fio:
                logger.info(f"ФИО в дательном падеже для ФДП: '{dative_fio}'")
            else:
                logger.warning("ФИО в дательном падеже не найдено, поле ФДП останется пустым")

            # Данные для заполнения
            fill_data = {
                "НОМ": str(current_number),
                "ФР": surname_gen_final,
                "ИР": first_name_gen_final,
                "ОР": patronymic_gen_final,
                "ФИ": surname_nom_final,
                "ИИ": first_name_nom_final,
                "ОИ": patronymic_nom_final,
                "ДР": format_date_ddmmyyyy(birth_date),
                "ДП": resolution_day,
                "МП": resolution_month,
                "ДН": current_day,
                "МН": current_month_gen,
            }

            if dative_fio:
                fill_data["ФДП"] = dative_fio

            # Следователь
            investigator_data = find_investigator(extractor.text)
            if investigator_data:
                fill_data.update(investigator_data)
                logger.info(f"Добавлены данные следователя: {investigator_data}")
            else:
                logger.info("Следователь не найден в тексте постановления")

            card_filler.fill_card(wb, fill_data)
            wb.save(new_card_path)
            logger.info(f"Карта успешно создана: {new_card_path}")
            processed_count += 1
            current_number += 1

        except Exception as e:
            logger.error(f"Критическая ошибка при обработке файла {doc_file.name}: {e}", exc_info=True)
            skipped_count += 1

    # Итоги
    logger.info("=" * 60)
    logger.info(f"Обработка завершена. Всего файлов: {len(doc_files)}")
    logger.info(f"Успешно создано карт: {processed_count}")
    logger.info(f"Пропущено (ошибки или нет данных): {skipped_count}")
    if processed_count > 0:
        logger.info(f"Следующий свободный номер заключения: {current_number}")
    logger.info("=" * 60)


def process_bvp_cards(script_dir: Path) -> None:
    """
    Режим «Заведующая: Заполнение Карт по СПИСКУ ОБРАЗЦОВ БВП».
    """
    logger.info("=== Запущен режим заполнения карт по списку образцов БВП ===")

    # ---------- 1. Ввод пути к папке с исходным Excel ----------
    user_input = input(
        f"Введите путь к папке с файлом 'Список_БВП_совпавшие'\n"
        f"(или нажмите Enter для '{BVP_DIR_DEFAULT}'): "
    ).strip()
    input_dir = Path(user_input) if user_input else Path(BVP_DIR_DEFAULT)
    if not input_dir.exists():
        logger.error(f"Папка с исходными данными не найдена: {input_dir}")
        input("Нажмите Enter для завершения...")
        return
    logger.info(f"Папка с исходным файлом: {input_dir}")

    # ---------- 2. Поиск файла «Список_БВП_совпавшие» ----------
    source_candidates = [
        f for f in input_dir.iterdir()
        if f.suffix.lower() in ('.xlsx', '.xlsm', '.xls')
        and not f.name.startswith('~$')
        and f.stem == "Список_БВП_совпавшие"
    ]
    if not source_candidates:
        logger.error(
            f"В папке {input_dir} не найден файл с именем 'Список_БВП_совпавшие' "
            f"(допустимые расширения: .xlsx, .xlsm, .xls)."
        )
        input("Нажмите Enter для завершения...")
        return
    if len(source_candidates) > 1:
        logger.warning(
            f"Найдено несколько файлов с именем 'Список_БВП_совпавшие': "
            f"{[f.name for f in source_candidates]}. Будет использован первый: "
            f"{source_candidates[0].name}"
        )
    source_file = source_candidates[0]
    logger.info(f"Исходный файл: {source_file}")

    # ---------- 3. Ввод папки для сохранения карт ----------
    default_output_dir = input_dir
    user_output = input(
        f"Введите путь для сохранения создаваемых карт\n"
        f"(или нажмите Enter для '{default_output_dir}'): "
    ).strip()
    output_dir = Path(user_output) if user_output else default_output_dir
    try:
        output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Папка для сохранения карт: {output_dir}")
    except Exception as e:
        logger.error(f"Не удалось создать папку {output_dir}: {e}")
        input("Нажмите Enter для завершения...")
        return

    # ---------- 4. Поиск шаблона карты в папке программы ----------
    template_path = find_template_excel(script_dir)
    if not template_path:
        logger.error("Не удалось найти шаблон карты (Excel-файл) в папке программы.")
        input("Нажмите Enter для завершения...")
        return
    logger.info(f"Шаблон карты: {template_path}")

    # ---------- 5. Чтение исходного файла ----------
    try:
        wb_source = openpyxl.load_workbook(source_file, data_only=True)
        ws_source = wb_source.active
    except Exception as e:
        logger.error(f"Не удалось открыть файл {source_file}: {e}")
        input("Нажмите Enter для завершения...")
        return

    # ---------- 6. Определение столбцов (заголовки + fallback на индексы) ----------
    def find_column_index(ws, header_candidates: list, fallback_col: int) -> int:
        """
        Ищет столбец по одному из вариантов заголовка (без учёта регистра,
        с обрезкой пробелов). Если не найден — возвращает fallback_col.
        """
        for col_idx in range(1, ws.max_column + 1):
            cell_val = ws.cell(row=1, column=col_idx).value
            if cell_val is not None:
                cell_str = str(cell_val).strip().lower()
                for h in header_candidates:
                    if h.lower() == cell_str:
                        logger.info(
                            f"Столбец '{h}' найден как колонка {col_idx} "
                            f"(заголовок: '{cell_val}')"
                        )
                        return col_idx
        logger.warning(
            f"Не найден ни один из заголовков {header_candidates}. "
            f"Используется fallback-колонка {fallback_col}."
        )
        return fallback_col

    # Варианты заголовков для каждого поля
    col_fio_nom = find_column_index(
        ws_source,
        [
            "ФИО (полностью) БВП в именительном падеже",
            "ФИО (полностью) БВП в именительном падеже",
            "ФИО БВП в именительном падеже",
            "ФИО БВП",
            "ФИО полностью БВП"
        ],
        fallback_col=2
    )
    col_date_birth = find_column_index(
        ws_source,
        ["дата рождения БВП", "Дата рождения", "дата рождения"],
        fallback_col=3
    )
    col_military_unit = find_column_index(
        ws_source,
        ["Войсковая часть", "Войсковая часть", "ВЧ"],
        fallback_col=4
    )
    col_personal_num = find_column_index(
        ws_source,
        ["Личный номер", "Личный номер", "ЛН"],
        fallback_col=5
    )
    col_disappearance = find_column_index(
        ws_source,
        [
            "Дата безвестного исчезновения",
            "Дата исчезновения",
            "Дата БИ",
            "ДБИ"
        ],
        fallback_col=6
    )

    # ---------- 7. Обработка строк ----------
    card_filler = ExcelCardFiller(template_path)
    processed = 0
    skipped = 0

    for row_idx in range(2, ws_source.max_row + 1):
        try:
            # Считываем данные из строки
            fio_raw = ws_source.cell(row=row_idx, column=col_fio_nom).value
            if fio_raw is None or str(fio_raw).strip() == "":
                logger.debug(f"Строка {row_idx}: пустое ФИО, пропускаем")
                continue

            birth_date_raw = ws_source.cell(row=row_idx, column=col_date_birth).value
            military_unit = ws_source.cell(row=row_idx, column=col_military_unit).value
            personal_number = ws_source.cell(row=row_idx, column=col_personal_num).value
            disappearance_raw = ws_source.cell(row=row_idx, column=col_disappearance).value

            # Парсинг ФИО (именительный падеж)
            surname, first_name, patronymic = parse_full_name(str(fio_raw).strip())
            if not surname:
                logger.warning(
                    f"Строка {row_idx}: не удалось разобрать ФИО '{fio_raw}', пропуск"
                )
                skipped += 1
                continue

            # Форматирование дат (приводим к DD.MM.YYYY)
            def format_date_safe(val):
                if val is None:
                    return ""
                val_str = str(val).strip()
                if not val_str:
                    return ""
                # Сначала пробуем прямой формат DD.MM.YYYY
                fmt = format_date_ddmmyyyy(val_str)
                if re.match(r'\d{2}\.\d{2}\.\d{4}', fmt):
                    return fmt
                # Если не вышло – парсим через parse_date_flexible и форматируем
                dt = parse_date_flexible(val_str)
                if dt:
                    return f"{dt.day:02d}.{dt.month:02d}.{dt.year}"
                # Если и это не удалось – возвращаем как есть
                logger.warning(
                    f"Строка {row_idx}: не удалось привести дату '{val_str}' "
                    f"к формату DD.MM.YYYY, используется исходное значение"
                )
                return val_str

            birth_date_fmt = format_date_safe(birth_date_raw)
            disappearance_fmt = format_date_safe(disappearance_raw)

            # Преобразование в родительный падеж
            surname_gen = card_filler.convert_to_genitive(surname)
            first_name_gen = card_filler.convert_to_genitive(first_name)
            patronymic_gen = card_filler.convert_to_genitive(patronymic)

            # Текущая дата
            current_day, current_month_gen = get_current_date()
            current_day = format_two_digits(current_day)

            # Создание копии шаблона
            safe_surname = sanitize_filename(surname)
            safe_first = sanitize_filename(first_name)
            safe_patr = sanitize_filename(patronymic)
            card_filename = (
                f"БВП {safe_surname} {safe_first} {safe_patr}.xlsx"
            )
            new_card_path = card_filler.create_copy(card_filename, output_dir)

            # Заполнение карты
            wb_card = openpyxl.load_workbook(new_card_path)

            fill_data = {
                "1": "СВО_Молов БВП_образец_прямая идентификация",
                "ТО": "крови",
                "ФИ": surname,
                "ФПИ": surname,
                "ИИ": first_name,
                "ИПИ": first_name,
                "ОИ": patronymic,
                "ОПИ": patronymic,
                "ФР": surname_gen,
                "ИР": first_name_gen,
                "ОР": patronymic_gen,
                "ДР": birth_date_fmt,
                "ДРП": birth_date_fmt,
                "ВЧ": military_unit if military_unit else "",
                "ЛН": personal_number if personal_number else "",
                "ДБИ": disappearance_fmt,
                "ТЕР": "Тюменской области",
                "ДП": current_day,
                "ДН": current_day,
                "МП": current_month_gen,
                "МН": current_month_gen,
                "РОД": None,
                "РОД1": None,
                "НО": "марле",
                "МАТ": "бурого пятна на марле"
            }

            card_filler.fill_card(wb_card, fill_data)
            wb_card.save(new_card_path)
            logger.info(
                f"Строка {row_idx}: создана карта '{new_card_path.name}' "
                f"(ФИО: {surname} {first_name} {patronymic})"
            )
            processed += 1

        except Exception as e:
            logger.error(
                f"Ошибка при обработке строки {row_idx}: {e}",
                exc_info=True
            )
            skipped += 1

    # ---------- 8. Итоговая статистика ----------
    logger.info("=" * 60)
    logger.info(
        f"Обработка завершена. Всего строк (с данными): {processed}, "
        f"пропущено: {skipped}"
    )
    logger.info("=" * 60)
    input("Нажмите Enter для завершения...")

def process_fzo_cards(script_dir: Path) -> None:
    """
    Режим 10: Заполнение Карт по СПИСКУ ОБРАЗЦОВ БВП (ФЗО).
    Гибкое определение столбцов, поддержка ФИО в одной ячейке и разнесённого по трём.
    """
    logger.info("=== Запущен режим заполнения карт по списку ФЗО ===")

    # ------------------- Вспомогательные функции (внутренние) -------------------
    def extract_birth_date_from_fio(fio_str: str) -> Tuple[str, str]:
        """
        Пытается отделить дату рождения, приклеенную к концу строки ФИО.
        Возвращает (очищенное_ФИО, дата_в_формате_DD.MM.YYYY или пустая строка).
        """
        if not fio_str:
            return "", ""
        match = re.search(r'(\d{1,2})[.,](\d{1,2})[.,](\d{4})\s*$', fio_str)
        if match:
            day, month, year = match.group(1), match.group(2), match.group(3)
            birth = f"{int(day):02d}.{int(month):02d}.{year}"
            fio_clean = fio_str[:match.start()].strip().rstrip(',').strip()
            logger.debug(f"Из ФИО '{fio_str}' выделена дата '{birth}', осталось '{fio_clean}'")
            return fio_clean, birth
        return fio_str.strip(), ""

    def find_column_by_keywords(ws, keywords: List[str], exclude_keywords: Optional[List[str]] = None) -> Optional[int]:
        """
        Ищет столбец, заголовок которого содержит ВСЕ keywords (без учёта регистра)
        и НЕ содержит ни одного из exclude_keywords (если заданы).
        Возвращает индекс столбца (начиная с 1) или None.
        """
        for col_idx in range(1, ws.max_column + 1):
            cell_val = ws.cell(row=1, column=col_idx).value
            if cell_val is None:
                continue
            header = str(cell_val).strip().lower()
            if all(kw.lower() in header for kw in keywords):
                if exclude_keywords and any(excl.lower() in header for excl in exclude_keywords):
                    continue
                return col_idx
        return None

    def guess_columns(ws) -> Dict[str, Optional[int]]:
        """
        Пытается автоматически определить столбцы по заголовкам.
        """
        col_appl_fio = find_column_by_keywords(ws, ['заявител', 'фио']) or \
                       find_column_by_keywords(ws, ['фио', 'заявител']) or \
                       find_column_by_keywords(ws, ['ф.и.о.', 'заявител'])
        col_appl_sur = find_column_by_keywords(ws, ['фамил', 'заявител'])
        col_appl_name = find_column_by_keywords(ws, ['имя', 'заявител'], exclude_keywords=['отчество'])
        col_appl_patr = find_column_by_keywords(ws, ['отчество', 'заявител'])
        col_appl_birth = find_column_by_keywords(ws, ['дата', 'рожден', 'заявител']) or \
                         find_column_by_keywords(ws, ['д.р.', 'заявител']) or \
                         find_column_by_keywords(ws, ['др', 'заявител'])
        col_rel = find_column_by_keywords(ws, ['степень', 'родств']) or \
                  find_column_by_keywords(ws, ['родство'])
        col_miss_fio = find_column_by_keywords(ws, ['разыскиваем', 'фио']) or \
                       find_column_by_keywords(ws, ['пропавш', 'фио']) or \
                       find_column_by_keywords(ws, ['бвп', 'фио'])
        col_miss_sur = find_column_by_keywords(ws, ['фамил', 'разыскиваем']) or \
                       find_column_by_keywords(ws, ['фамил', 'пропавш'])
        col_miss_name = find_column_by_keywords(ws, ['имя', 'разыскиваем'], exclude_keywords=['отчество']) or \
                        find_column_by_keywords(ws, ['имя', 'пропавш'], exclude_keywords=['отчество'])
        col_miss_patr = find_column_by_keywords(ws, ['отчество', 'разыскиваем']) or \
                        find_column_by_keywords(ws, ['отчество', 'пропавш'])
        col_miss_birth = find_column_by_keywords(ws, ['дата', 'рожден', 'разыскиваем']) or \
                         find_column_by_keywords(ws, ['дата', 'рожден', 'пропавш']) or \
                         find_column_by_keywords(ws, ['д.р.', 'разыскиваем']) or \
                         find_column_by_keywords(ws, ['д.р.', 'пропавш'])

        return {
            'applicant_fio_single': col_appl_fio,
            'applicant_surname': col_appl_sur,
            'applicant_firstname': col_appl_name,
            'applicant_patronymic': col_appl_patr,
            'applicant_birthdate': col_appl_birth,
            'relationship': col_rel,
            'missing_fio_single': col_miss_fio,
            'missing_surname': col_miss_sur,
            'missing_firstname': col_miss_name,
            'missing_patronymic': col_miss_patr,
            'missing_birthdate': col_miss_birth,
        }

    def manual_column_input(ws, max_col: int) -> Dict[str, Optional[int]]:
        """
        Интерактивный ручной ввод номеров колонок.
        """
        print("\nНе удалось автоматически определить столбцы. Укажите номера колонок вручную.")
        print("(0 — колонка отсутствует, Enter — оставить пустым)\n")
        result = {}
        fields = [
            ('applicant_fio_single', 'ФИО заявителя (одной ячейкой)'),
            ('applicant_surname', 'Фамилия заявителя'),
            ('applicant_firstname', 'Имя заявителя'),
            ('applicant_patronymic', 'Отчество заявителя'),
            ('applicant_birthdate', 'Дата рождения заявителя'),
            ('relationship', 'Степень родства'),
            ('missing_fio_single', 'ФИО разыскиваемого (одной ячейкой)'),
            ('missing_surname', 'Фамилия разыскиваемого'),
            ('missing_firstname', 'Имя разыскиваемого'),
            ('missing_patronymic', 'Отчество разыскиваемого'),
            ('missing_birthdate', 'Дата рождения разыскиваемого'),
        ]
        for key, desc in fields:
            while True:
                try:
                    inp = input(f"  {desc} (номер столбца 1-{max_col}, 0=нет): ").strip()
                    if not inp:
                        result[key] = None
                        break
                    val = int(inp)
                    if 0 <= val <= max_col:
                        result[key] = val if val > 0 else None
                        break
                    print(f"Введите число от 0 до {max_col}")
                except ValueError:
                    print("Введите целое число.")
        return result

    # ------------------- Основное тело функции -------------------
    # 1. Папка с исходным файлом
    user_input = input(
        f"Введите путь к папке с файлом 'Список_ФЗО'\n"
        f"(или нажмите Enter для '{FZO_DIR_DEFAULT}'): "
    ).strip()
    input_dir = Path(user_input) if user_input else Path(FZO_DIR_DEFAULT)
    if not input_dir.exists():
        logger.error(f"Папка с исходными данными не найдена: {input_dir}")
        input("Нажмите Enter для завершения...")
        return
    logger.info(f"Папка с исходным файлом: {input_dir}")

    # 2. Поиск файла Список_ФЗО
    source_candidates = [
        f for f in input_dir.iterdir()
        if f.suffix.lower() in ('.xlsx', '.xlsm')
        and not f.name.startswith('~$')
        and f.stem == "Список_ФЗО"
    ]
    if not source_candidates:
        logger.error(
            f"В папке {input_dir} не найден файл 'Список_ФЗО.xlsx/.xlsm'. "
            f"Проверьте имя файла."
        )
        input("Нажмите Enter для завершения...")
        return
    if len(source_candidates) > 1:
        logger.warning(
            f"Найдено несколько файлов 'Список_ФЗО': {[f.name for f in source_candidates]}. "
            f"Будет использован первый: {source_candidates[0].name}"
        )
    source_file = source_candidates[0]
    logger.info(f"Исходный файл: {source_file}")

    # 3. Шаблон карты в папке программы
    template_path = find_template_excel(script_dir)
    if not template_path:
        logger.error("Не удалось найти шаблон карты (Excel-файл) в папке программы.")
        input("Нажмите Enter для завершения...")
        return

    # 4. Папка для сохранения карт
    default_output_dir = script_dir
    user_output = input(
        f"Введите путь для сохранения создаваемых карт\n"
        f"(или нажмите Enter для '{default_output_dir}'): "
    ).strip()
    output_dir = Path(user_output) if user_output else default_output_dir
    try:
        output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Папка для сохранения карт: {output_dir}")
    except Exception as e:
        logger.error(f"Не удалось создать папку {output_dir}: {e}")
        input("Нажмите Enter для завершения...")
        return

    # 5. Загрузка исходной книги
    try:
        wb_source = openpyxl.load_workbook(source_file, data_only=True)
        ws_source = wb_source.active
    except Exception as e:
        logger.error(f"Не удалось открыть файл {source_file}: {e}")
        input("Нажмите Enter для завершения...")
        return

    # 6. Определение столбцов
    col_map = guess_columns(ws_source)
    print("\nАвтоматически определённые столбцы:")
    desc = {
        'applicant_fio_single': 'ФИО заявителя (одна ячейка)',
        'applicant_surname': 'Фамилия заявителя',
        'applicant_firstname': 'Имя заявителя',
        'applicant_patronymic': 'Отчество заявителя',
        'applicant_birthdate': 'Дата рождения заявителя',
        'relationship': 'Степень родства',
        'missing_fio_single': 'ФИО разыскиваемого (одна ячейка)',
        'missing_surname': 'Фамилия разыскиваемого',
        'missing_firstname': 'Имя разыскиваемого',
        'missing_patronymic': 'Отчество разыскиваемого',
        'missing_birthdate': 'Дата рождения разыскиваемого',
    }
    for key, d in desc.items():
        idx = col_map.get(key)
        if idx:
            header_text = ws_source.cell(row=1, column=idx).value
            print(f"  {d}: колонка {idx} (заголовок: {header_text})")
        else:
            print(f"  {d}: не определён")

    confirm = input("\nПринять автоматическое определение? (д/н/вручную): ").strip().lower()
    if confirm not in ('д', 'y', 'yes', 'да'):
        if confirm in ('вручную', 'manual', 'm'):
            col_map = manual_column_input(ws_source, ws_source.max_column)
        else:
            logger.info("Пользователь отказался от автоматического определения, переходим к ручному вводу.")
            col_map = manual_column_input(ws_source, ws_source.max_column)

    # 7. Определяем режим извлечения ФИО
    use_single_applicant = bool(col_map.get('applicant_fio_single'))
    use_split_applicant = all([
        col_map.get('applicant_surname'),
        col_map.get('applicant_firstname')
    ])

    use_single_missing = bool(col_map.get('missing_fio_single'))
    use_split_missing = all([
        col_map.get('missing_surname'),
        col_map.get('missing_firstname')
    ])

    if not (use_single_applicant or use_split_applicant):
        logger.error("Не удалось определить колонки для ФИО заявителя (ни одна ячейка, ни разнесённые).")
        input("Нажмите Enter для завершения...")
        return
    if not (use_single_missing or use_split_missing):
        logger.error("Не удалось определить колонки для ФИО разыскиваемого.")
        input("Нажмите Enter для завершения...")
        return

    logger.info(
        f"Режим заявителя: {'одинарный' if use_single_applicant else 'разнесённый'}, "
        f"Режим разыскиваемого: {'одинарный' if use_single_missing else 'разнесённый'}"
    )

    # 8. Основной цикл обработки строк
    card_filler = ExcelCardFiller(template_path)
    processed = 0
    skipped = 0

    for row_idx in range(2, ws_source.max_row + 1):
        try:
            # ----- Заявитель -----
            applicant_surname = applicant_first = applicant_patr = ""
            applicant_birthdate = ""
            missing_birthdate = ""

            if use_single_applicant:
                fio_cell = ws_source.cell(row=row_idx, column=col_map['applicant_fio_single'])
                fio_raw = fio_cell.value
                if not fio_raw or not str(fio_raw).strip():
                    logger.debug(f"Строка {row_idx}: пустое ФИО заявителя, пропускаем")
                    continue
                fio_clean, extracted_date = extract_birth_date_from_fio(str(fio_raw))
                if extracted_date and not col_map.get('applicant_birthdate'):
                    applicant_birthdate = extracted_date
                applicant_surname, applicant_first, applicant_patr = parse_full_name(fio_clean)
                if not applicant_surname:
                    logger.warning(f"Строка {row_idx}: не удалось разобрать ФИО заявителя из '{fio_clean}', пропуск")
                    skipped += 1
                    continue
            else:  # разнесённый
                surname_cell = ws_source.cell(row=row_idx, column=col_map['applicant_surname'])
                first_cell = ws_source.cell(row=row_idx, column=col_map['applicant_firstname'])
                applicant_surname = normalize_text(surname_cell.value) if surname_cell.value else ""
                applicant_first = normalize_text(first_cell.value) if first_cell.value else ""

                if not applicant_surname:
                    logger.debug(f"Строка {row_idx}: фамилия заявителя пуста, пропускаем")
                    continue

                # Отчество (может отсутствовать)
                patr_idx = col_map.get('applicant_patronymic')
                if patr_idx:
                    patr_cell = ws_source.cell(row=row_idx, column=patr_idx)
                    applicant_patr = normalize_text(patr_cell.value) if patr_cell.value else ""

            # Дата рождения заявителя из отдельной колонки (если есть) имеет приоритет
            birth_idx = col_map.get('applicant_birthdate')
            if birth_idx:
                birth_cell = ws_source.cell(row=row_idx, column=birth_idx)
                if birth_cell.value:
                    applicant_birthdate = str(birth_cell.value).strip()

            # Приводим дату к формату DD.MM.YYYY
            if applicant_birthdate:
                applicant_birthdate = format_date_ddmmyyyy(applicant_birthdate)
                if not re.match(r'\d{2}\.\d{2}\.\d{4}', applicant_birthdate):
                    dt = parse_date_flexible(applicant_birthdate)
                    if dt:
                        applicant_birthdate = f"{dt.day:02d}.{dt.month:02d}.{dt.year}"
                    else:
                        logger.warning(f"Строка {row_idx}: не удалось распознать дату рождения '{applicant_birthdate}', оставлена как есть")

            # ----- Степень родства -----
            relationship = ""
            rel_idx = col_map.get('relationship')
            if rel_idx:
                rel_cell = ws_source.cell(row=row_idx, column=rel_idx)
                if rel_cell.value:
                    relationship = normalize_degree(str(rel_cell.value))
                    logger.debug(f"Строка {row_idx}: степень родства '{rel_cell.value}' -> '{relationship}'")

            # ----- Разыскиваемый -----
            missing_surname = missing_first = missing_patr = ""

            if use_single_missing:
                miss_cell = ws_source.cell(row=row_idx, column=col_map['missing_fio_single'])
                if miss_cell.value:
                    miss_fio_clean, miss_extracted_date = extract_birth_date_from_fio(str(miss_cell.value))
                    # Если дата извлечена и отдельная колонка даты не указана – используем её
                    if miss_extracted_date and not col_map.get('missing_birthdate'):
                        missing_birthdate = miss_extracted_date
                    missing_surname, missing_first, missing_patr = parse_full_name(miss_fio_clean)
                    logger.debug(f"Разыскиваемый из одной ячейки: {missing_surname} {missing_first} {missing_patr}")
            else:
                miss_sur_cell = ws_source.cell(row=row_idx, column=col_map['missing_surname'])
                miss_name_cell = ws_source.cell(row=row_idx, column=col_map['missing_firstname'])
                missing_surname = normalize_text(miss_sur_cell.value) if miss_sur_cell.value else ""
                missing_first = normalize_text(miss_name_cell.value) if miss_name_cell.value else ""

                miss_patr_idx = col_map.get('missing_patronymic')
                if miss_patr_idx:
                    miss_patr_cell = ws_source.cell(row=row_idx, column=miss_patr_idx)
                    missing_patr = normalize_text(miss_patr_cell.value) if miss_patr_cell.value else ""

                # Чтение даты рождения разыскиваемого из отдельного столбца
                miss_birth_idx = col_map.get('missing_birthdate')
                if miss_birth_idx:
                    miss_birth_cell = ws_source.cell(row=row_idx, column=miss_birth_idx)
                    if miss_birth_cell.value:
                        missing_birthdate = str(miss_birth_cell.value).strip()

            if not missing_surname:
                logger.warning(f"Строка {row_idx}: фамилия разыскиваемого не определена")

            # Приводим дату рождения разыскиваемого к формату DD.MM.YYYY
            if missing_birthdate:
                missing_birthdate = format_date_ddmmyyyy(missing_birthdate)
                if not re.match(r'\d{2}\.\d{2}\.\d{4}', missing_birthdate):
                    dt = parse_date_flexible(missing_birthdate)
                    if dt:
                        missing_birthdate = f"{dt.day:02d}.{dt.month:02d}.{dt.year}"
                    else:
                        logger.warning(f"Строка {row_idx}: не удалось распознать дату рождения разыскиваемого '{missing_birthdate}', оставлена как есть")

            # ----- Заполнение карты -----
            applicant_surname_gen = card_filler.convert_to_genitive(applicant_surname)
            applicant_first_gen = card_filler.convert_to_genitive(applicant_first)
            applicant_patr_gen = card_filler.convert_to_genitive(applicant_patr)

            rod_value = ROD_MAPPING.get(relationship, "") if relationship else ""

            safe_name = sanitize_filename(f"{applicant_surname} {applicant_first} {applicant_patr}".strip())
            card_filename = f"ФЗО {safe_name}.xlsx"
            new_card_path = card_filler.create_copy(card_filename, output_dir)

            wb_card = openpyxl.load_workbook(new_card_path)

            # Получаем текущую дату (день и месяц в родительном падеже)
            current_day, current_month_gen = get_current_date()
            current_day_str = format_two_digits(current_day)

            fill_data = {
                "1": "СВО_Молов_образец_прямая идентификация",
                "ТО": "крови",
                "ДП": current_day_str,
                "МП": current_month_gen,
                "ДН": current_day_str,
                "МН": current_month_gen,
                "ФИ": applicant_surname,
                "ИИ": applicant_first,
                "ОИ": applicant_patr,
                "ФР": applicant_surname_gen,
                "ИР": applicant_first_gen,
                "ОР": applicant_patr_gen,
                "ДР": applicant_birthdate,
                "ДРП": missing_birthdate,
                "РОД": rod_value,
                "ФПИ": missing_surname,
                "ИПИ": missing_first,
                "ОПИ": missing_patr,
                "КОН": "бумажном конверте",
                "НАД": "с отпечатанной на принтере надписью",
                "НО": "марле",
                "МАТ": "бурого пятна на марле",
            }
            # Очищаем неиспользуемые ячейки
            for code in ["ВЧ", "ЛН", "ДБИ", "РОД1"]:
                fill_data[code] = None

            card_filler.fill_card(wb_card, fill_data)
            wb_card.save(new_card_path)
            logger.info(f"Строка {row_idx}: создана карта '{new_card_path.name}'")
            processed += 1

        except Exception as e:
            logger.error(f"Ошибка при обработке строки {row_idx}: {e}", exc_info=True)
            skipped += 1

    logger.info("=" * 60)
    logger.info(f"Обработка завершена. Создано карт: {processed}, пропущено: {skipped}")
    logger.info("=" * 60)
    input("Нажмите Enter для завершения...")

# ========== ОСНОВНАЯ ФУНКЦИЯ ==========
def main():
    logger.info("Программа запущена")
    try:
        script_dir = Path(__file__).parent
        work_type = select_work_type()

        if work_type == "tobolsk_newborn":
            template_path = find_template_excel(script_dir)
            if not template_path:
                logger.error("Не удалось найти шаблон карты")
                input("Нажмите Enter для завершения...")
                return
            try:
                start_number = int(input("Введите начальный номер заключения: ").strip())
            except ValueError:
                logger.error("Номер заключения должен быть целым числом")
                input("Нажмите Enter для завершения...")
                return
            process_tobolsk_newborn(script_dir, template_path, start_number)

        elif work_type == "bone_cards":
            process_bone_cards(script_dir)

        elif work_type == "registry_cards":
            process_registry_cards(script_dir)

        elif work_type == "distribution_cards":
            process_distribution_cards(script_dir)

        elif work_type == "mass_fill_samples":
            process_mass_fill_cards_with_samples(script_dir)

        elif work_type == "mass_print_postanovleniy":
            process_mass_print_postanovleniy()

        elif work_type == "rostov_cards":
            process_rostov_cards(script_dir)

        elif work_type == "bvp_cards":
            process_bvp_cards(script_dir)

        elif work_type == "fzo_cards":
            process_fzo_cards(script_dir)

        elif work_type == "ekc_postanovleniya":
            template_path = find_template_excel(script_dir)
            if not template_path:
                logger.error("Не удалось найти шаблон карты")
                input("Нажмите Enter для завершения...")
                return
            try:
                start_number = int(input("Введите начальный номер заключения: ").strip())
            except ValueError:
                logger.error("Номер заключения должен быть целым числом")
                input("Нажмите Enter для завершения...")
                return
            process_ekc_postanovleniya(script_dir, template_path, start_number)

        else:
            logger.error(f"Неизвестный тип работы: {work_type}")

        logger.info("Обработка завершена")
        input("Нажмите Enter для завершения...")

    except KeyboardInterrupt:
        logger.info("Программа прервана пользователем")
        input("Нажмите Enter для завершения...")
    except Exception as e:
        logger.error(f"Критическая ошибка: {e}", exc_info=True)
        input("Нажмите Enter для завершения...")


if __name__ == "__main__":
    main()