#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Программа для автоматического заполнения карт Excel (заключений эксперта)
на основе данных из постановлений (Word-документы .docx или .doc).
"""

import os
import re
import shutil
import logging
import sys
import win32com.client
import win32print
import win32api
import time
import openpyxl
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional, Tuple, List
from docx import Document
from pymorphy2 import MorphAnalyzer
# from docx.oxml import OxmlElement, qn


# ========== НАСТРОЙКИ ==========
# Путь к папке с постановлениями
POSTANOVLENIYA_DIR_DEFAULT = r"F:\Работа\Работа с Картами\POSTANOVENIA_TOBOLSK"

# Путь для сохранения создаваемых карт (если оставить пустым, то будет использоваться папка с программой)
OUTPUT_CARDS_DIR_DEFAULT = r"F:\Работа\Работа с Картами\OUTPUT_CARDS"

# Путь к папке с исходными данными для режима "Ростов"
ROSTOV_DIR_DEFAULT = r"F:\Работа\Регистратура\4_Ростов"

# Путь для сохранения сгенерированных постановлений при массовой печати
MASSPRINT_SAVEPATH_DEFAULT = r"F:\Работа\Работа с Картами\MASSPRINT_SAVEPATH"

# Путь к папке с шаблонами Word для постановлений
TEMPLATES_FOLDER_DEFAULT = r"U:\ШАБЛОНЫ\Заключения\СВО\Образцы"

# Имя файла лога
LOG_FILENAME = "card_filler.log"


# ========== НАСТРОЙКА ЛОГИРОВАНИЯ ==========
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOG_FILENAME, encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Инициализация морфологического анализатора
morph = MorphAnalyzer()


# ========== КОНФИГУРАЦИЯ СЛЕДОВАТЕЛЕЙ ==========
# Добавляйте новых следователей в этот список.
# Каждый элемент – словарь с ключами:
#   "fio_patterns" – список строк (вариантов написания ФИО),
#   "naz", "naz3", "naz4" – значения, которые нужно записать в карту.
INVESTIGATORS_CONFIG = [
    {
        "fio_patterns": ["Чекан П.Д.", "Чекан Павел Дмитриевич"],
        "naz": "следователя-криминалиста отдела криминалистики СУ СК России по Тюменской области Чекана П.Д.",
        "naz3": "следователю-криминалисту отдела криминалистики СУ СК России по Тюменской области",
        "naz4": "Чекану П.Д."
    },
    #{
    #    "fio_patterns": ["Петров П.П.", "Петров Петр Петрович"],
    #    "naz": "Следователя-криминалиста отдела криминалистики СУ СК России по Тюменской области Петрова П.П.",
    #    "naz3": "Следователю-криминалисту отдела криминалистики СУ СК России по Тюменской области",
    #    "naz4": "Петрову П.П."
    #},
    # Добавляйте следующих следователей аналогично
]


# ========== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ==========
# Маппинг степени родства -> значение для ячейки РОД (родительный падеж)
ROD_MAPPING = {
    'отец': 'отца',
    'мать': 'матери',
    'брат': 'брата',
    'сестра': 'сестры',
    'сын': 'сына',
    'дочь': 'дочери',
    'дядя': 'дяди',
    'тётя': 'тёти',
    'бабушка': 'бабушки',
    'дедушка': 'дедушки',
    'неполнородная сестра': 'неполнородной сестры',
    'неполнородный брат': 'неполнородного брата',
    'не указано': 'родство не указано'
}


def select_work_type() -> str:
    """
    Предлагает пользователю выбрать тип работы.
    Возвращает строку-идентификатор выбранного варианта.
    """
    print("\nВыберите тип работы:")
    print("\n1. Заведующая: Постановления из ТОБОЛЬСКА по новорожденному\n(поместить в любую папку с образцом нужной Карты)")
    print("\n2. Регистратура: Заполнение Карт с КОСТЯМИ для РАСПРЕДЕЛЕНИЯ между экспертами\n(поместить в корневую папку '1_Кости' с папками экспертов)")
    print("\n3. Регистратура: Заполнение Карт с ОБРАЗЦАМИ в случае ПРЯМОЙ идентификации\n(поместить в папку с Картами, предназначенными для преобразования)")
    print("\n4. Регистратура: Массовое заполнение Карт с ОБРАЗЦАМИ (преобразование ФИО из родительного в именительный)\n(поместить в папку с Картами образцов)")
    print("\n5. Регистратура: Массовая печать постановлений\n(поместить в папку с распечатываемыми Картами)")
    print("\n6. Эксперты: Заполнение Карт с КОСТЯМИ при ЗАВЕРШЕНИИ работы с заключением (дата окончания и объект)\n(поместить в папку с Картами костей)")
    print("\n7. Заведующая: Заполнение Карт по ОБРАЗЦАМ ИЗ РОСТОВА\n(поместить в любую папку с образцом нужной Карты, обрабатываемый Excel - в папку 4_Ростов)")

    while True:
        choice = input("\nВведите номер: ").strip()
        if choice == "1":
            return "tobolsk_newborn"
        elif choice == "2":
            return "distribution_cards"
        elif choice == "3":
            return "registry_cards"
        elif choice == "4":
            return "mass_fill_samples"
        elif choice == "5":
            return "mass_print_postanovleniy"
        elif choice == "6":
            return "bone_cards"
        elif choice == "7":
            return "rostov_cards"
        else:
            print("Неверный выбор, попробуйте снова.")


def normalize_text(text: Optional[str]) -> str:
    """Удаляет лишние пробелы и приводит к строке"""
    if text is None:
        return ""
    return str(text).strip()


def add_dot_to_initials(text: str) -> str:
    """Если строка состоит из одной буквы (не содержит точки), добавляет точку."""
    if not text:
        return text
    text = text.strip()
    if len(text) == 1 and text.isalpha():
        return text + "."
    return text


def parse_date_flexible(date_value) -> Optional[datetime]:
    """
    Пытается преобразовать входное значение (строка или datetime) в объект datetime.
    Поддерживает различные форматы, включая наличие времени.
    Возвращает datetime или None, если не удалось.
    """
    if date_value is None:
        return None
    
    # Если уже datetime, возвращаем как есть
    if isinstance(date_value, datetime):
        return date_value
    
    # Преобразуем в строку
    date_str = str(date_value).strip()
    if not date_str:
        return None
    
    # Список возможных форматов (без времени и с временем)
    formats = [
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M:%S.%f",
        "%d.%m.%Y %H:%M:%S",
        "%d.%m.%Y %H:%M",
        "%d.%m.%Y",
        "%d.%m.%y",
        "%Y-%m-%d",
        "%d/%m/%Y",
        "%d/%m/%y",
        "%m/%d/%Y",
        "%m/%d/%y",
        "%Y/%m/%d",
    ]
    
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    
    # Если стандартные форматы не помогли, пробуем извлечь компоненты с помощью регулярного выражения
    patterns = [
        r'(\d{1,2})[./-](\d{1,2})[./-](\d{2,4})',   # DD.MM.YYYY или MM.DD.YYYY
        r'(\d{4})[./-](\d{1,2})[./-](\d{1,2})',     # YYYY-MM-DD
    ]
    
    for pattern in patterns:
        match = re.search(pattern, date_str)
        if match:
            parts = match.groups()
            # Если первый компонент длиной 4, то это год
            if len(parts[0]) == 4:
                year, month, day = map(int, parts)
            else:
                day, month, year = map(int, parts)
                # Простая проверка: если день > 31 или месяц > 12, возможно, поменяны местами
                if day > 31 or month > 12:
                    day, month = month, day
            # Корректировка двузначного года
            if year < 100:
                year += 2000 if year < 50 else 1900  # для дат после 1950
            try:
                return datetime(year, month, day)
            except ValueError:
                pass
    
    logger.warning(f"Не удалось распознать дату: {date_str}")
    return None


def natural_sort_key(filename: str):
    """
    Ключ для натуральной сортировки (1, 2, 10, 23 вместо 1, 10, 2, 23)
    """
    return [int(text) if text.isdigit() else text.lower()
            for text in re.split(r'(\d+)', filename)]


def validate_ob_format(value: str) -> Tuple[bool, Optional[int], Optional[str]]:
    """
    Проверяет формат ОБ: число + латинская буква
    Возвращает: (валидно, число, буква)
    """
    pattern = r'^(\d+)([A-Za-z])$'
    match = re.match(pattern, value)
    if match:
        return True, int(match.group(1)), match.group(2).upper()
    return False, None, None


def get_current_date() -> Tuple[int, str]:
    """Возвращает текущий день (число) и месяц в родительном падеже (например, 'марта')"""
    now = datetime.now()
    day = now.day
    # Словарь для преобразования номера месяца в русское название (именительный падеж)
    month_names = {
        1: "январь", 2: "февраль", 3: "март", 4: "апрель",
        5: "май", 6: "июнь", 7: "июль", 8: "август",
        9: "сентябрь", 10: "октябрь", 11: "ноябрь", 12: "декабрь"
    }
    month_nom = month_names[now.month]
    # Преобразуем в родительный падеж с помощью pymorphy2
    parsed = morph.parse(month_nom)[0]
    month_gen = parsed.inflect({'gent'})
    if month_gen:
        month_name_gen = month_gen.word
    else:
        # fallback словарь
        month_gen_dict = {
            "январь": "января", "февраль": "февраля", "март": "марта",
            "апрель": "апреля", "май": "мая", "июнь": "июня",
            "июль": "июля", "август": "августа", "сентябрь": "сентября",
            "октябрь": "октября", "ноябрь": "ноября", "декабрь": "декабря"
        }
        month_name_gen = month_gen_dict.get(month_nom, month_nom)
    return day, month_name_gen


def find_template_excel(script_dir: Path) -> Optional[Path]:
    """
    Находит файл Excel-шаблона в указанной директории.
    Если найден один файл, возвращает его.
    Если несколько, выводит список и просит выбрать.
    Если ни одного, возвращает None.
    """
    excel_files = list(script_dir.glob("*.xlsx")) + list(script_dir.glob("*.xlsm"))
    # Исключаем временные файлы
    excel_files = [f for f in excel_files if not f.name.startswith('~$')]
    if not excel_files:
        logger.error("В папке программы не найдено файлов Excel (шаблонов карт)")
        return None
    if len(excel_files) == 1:
        logger.info(f"Найден шаблон карты: {excel_files[0].name}")
        return excel_files[0]
    else:
        print("Найдено несколько файлов Excel. Выберите шаблон:")
        for i, f in enumerate(excel_files, 1):
            print(f"{i}. {f.name}")
        while True:
            try:
                choice = int(input("Введите номер: ").strip())
                if 1 <= choice <= len(excel_files):
                    return excel_files[choice-1]
                else:
                    print("Неверный номер.")
            except ValueError:
                print("Введите число.")


def find_cell_by_code(ws, code: str):
    """
    Ищет ячейку в первом столбце с указанным кодом.
    Сравнение производится после приведения к строке и удаления пробелов.
    Возвращает ячейку в третьем столбце (столбец значений).
    """
    target = str(code).strip()
    for row in ws.iter_rows(min_row=1, max_col=3):
        cell_code = row[0].value
        if cell_code is not None:
            # Приводим значение из ячейки к строке и удаляем пробелы
            if str(cell_code).strip() == target:
                return row[2]
    return None


def month_num_to_genitive(month_num: int) -> str:
    """Преобразование номера месяца в русское название в родительном падеже"""
    months = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля",
        5: "мая", 6: "июня", 7: "июля", 8: "августа",
        9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
    }
    return months.get(month_num, "")


def sanitize_filename(name: str) -> str:
    """Удаляет недопустимые для имени файла символы."""
    invalid_chars = r'[<>:"/\\|?*]'
    name = re.sub(invalid_chars, '_', name)
    name = name.strip().strip('.')
    if not name:
        name = "unknown"
    return name


def normalize_degree(degree: str) -> str:
    """Приводит строку степени родства к каноническому виду."""
    degree = degree.lower().strip().replace('.', '').replace(',', '')
    if degree in ('личный генотип', 'личный', 'прямая идентификация', 'прямая', 'прямой'):
        return 'личный генотип'
    if degree in ('отец', 'папа'):
        return 'отец'
    if degree in ('мать', 'мама'):
        return 'мать'
    if degree in ('брат', 'братан'):
        return 'брат'
    if degree in ('сестра', 'сестрёнка'):
        return 'сестра'
    if degree in ('сын', 'сыночек'):
        return 'сын'
    if degree in ('дочь', 'дочка'):
        return 'дочь'
    if degree in ('дядя', 'дядюшка'):
        return 'дядя'
    if degree in ('тётя', 'тетя', 'тётка'):
        return 'тётя'
    if degree in ('бабушка', 'бабка'):
        return 'бабушка'
    if degree in ('дедушка', 'дед'):
        return 'дедушка'
    if degree in ('неполнородная сестра', 'неполнородная сестра'):
        return 'неполнородная сестра'
    if degree in ('неполнородный брат', 'неполнородный брат'):
        return 'неполнородный брат'
    if degree in ('родство не указано', 'не указано', 'неизвестно', 'н/у', 'н\\у'):
        return 'не указано'
    return degree


def parse_full_name(full_name: str) -> Tuple[str, str, str]:
    """
    Разбирает строку ФИО на фамилию, имя, отчество.
    Поддерживает 2-4 слова, сложные отчества (оглы, кызы и т.д.).
    Если отчество неизвестно (встречается "?" как отдельное слово), оно игнорируется.
    Возвращает кортеж (фамилия, имя, отчество).
    """
    if not full_name:
        return "", "", ""

    full_name = full_name.strip()
    full_name = re.sub(r'\s+', ' ', full_name)

    # Инициалы
    if '.' in full_name:
        match = re.match(r'^([А-ЯЁа-яё]+)\s+([А-ЯЁ]\.?)\s*([А-ЯЁ]\.?)?$', full_name)
        if match:
            surname = match.group(1)
            first_initial = match.group(2).replace('.', '')
            patronymic_initial = match.group(3).replace('.', '') if match.group(3) else ""
            return surname, first_initial, patronymic_initial

    words = full_name.split()
    word_count = len(words)
    complex_keywords = {'оглы', 'кызы', 'угли', 'гызы', 'заде', 'бек', 'хан'}

    # Обработка вопросительного знака (неизвестное отчество)
    if word_count >= 3 and words[-1] == '?':
        # Последнее слово - "?", отчество неизвестно, убираем его
        words = words[:-1]
        word_count = len(words)

    if word_count == 2:
        return words[0], words[1], ""
    elif word_count == 3:
        return words[0], words[1], words[2]
    elif word_count >= 4:
        last_word = words[-1].lower()
        if any(kw in last_word for kw in complex_keywords):
            return words[0], words[1], ' '.join(words[2:])
        else:
            logger.warning(f"Обнаружено более 4 слов без признаков сложного отчества: {full_name}. Берём первые 3.")
            return words[0], words[1], words[2]
    else:
        logger.error(f"Невозможно разобрать ФИО: {full_name}")
        return "", "", ""


def parse_txt_file(file_path: Path) -> List[Dict[str, str]]:
    """
    Парсинг txt файла с данными (пробуем разные кодировки)
    Возвращает список словарей с ключами 'order', 'corpse_number', 'date'
    """
    data = []
    encodings = ['utf-8', 'cp1251', 'windows-1251', 'ansi']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                lines = f.readlines()

            # Пропускаем заголовок (первая строка)
            for line in lines[1:]:
                line = line.strip()
                if line:
                    # Разделяем по табуляции или по нескольким пробелам
                    parts = re.split(r'\t|\s{2,}', line)
                    if len(parts) >= 3:
                        # Удаляем возможные лишние пробелы
                        parts = [p.strip() for p in parts]
                        data.append({
                            'order': parts[0],
                            'corpse_number': parts[1],
                            'date': parts[2]
                        })
            # Если удалось прочитать, прерываем цикл
            if data:
                logger.info(f"Успешно прочитан файл {file_path.name} с кодировкой {encoding}")
                break

        except UnicodeDecodeError:
            continue  # Пробуем следующую кодировку
        except Exception as e:
            logger.error(f"Ошибка при чтении файла {file_path} с кодировкой {encoding}: {e}")
            continue

    if not data:
        logger.warning(f"Не удалось прочитать данные из файла {file_path} ни с одной кодировкой")

    return data


def fill_excel_from_template(template_path: Path, output_path: Path, expertise_data: Dict[str, str]) -> bool:
    """Заполнение шаблона Excel данными и сохранение в output_path"""
    try:
        wb = openpyxl.load_workbook(template_path)
        ws = wb.active

        # Заполняем ячейки по кодам
        codes_map = {
            "НОМ": expertise_data.get('expertise_number'),
            "ДП": expertise_data.get('day'),
            "МП": expertise_data.get('month'),
            "ДН": expertise_data.get('start_day'),
            "МН": expertise_data.get('start_month'),
            "ТР": expertise_data.get('corpse_number')
        }

        for code, value in codes_map.items():
            if value is not None:
                cell = find_cell_by_code(ws, code)
                if cell:
                    cell.value = value
                    logger.debug(f"Записан код {code} = {value}")
                else:
                    logger.warning(f"Код {code} не найден в шаблоне")

        wb.save(output_path)
        logger.info(f"Создан файл: {output_path}")
        return True

    except Exception as e:
        logger.error(f"Ошибка при заполнении Excel файла {output_path}: {e}")
        return False


def find_investigator(text: str) -> Optional[Dict[str, str]]:
    """
    Ищет в тексте ФИО следователя из предопределённого списка.
    Возвращает словарь с ключами 'naz', 'naz3', 'naz4' или None, если не найдено.
    """
    if not text:
        logger.warning("Текст постановления пуст, поиск следователя невозможен")
        return None

    text_lower = text.lower()
    for inv in INVESTIGATORS_CONFIG:
        for pattern in inv["fio_patterns"]:
            if pattern.lower() in text_lower:
                logger.info(f"Найден следователь по шаблону '{pattern}': {inv}")
                return {
                    "НАЗ": inv["naz"],
                    "НАЗ3": inv["naz3"],
                    "НАЗ4": inv["naz4"]
                }
    logger.info("Следователь из предопределённого списка не найден")
    return None


# ========== КЛАСС ДЛЯ ИЗВЛЕЧЕНИЯ ДАННЫХ ИЗ ПОСТАНОВЛЕНИЯ ==========
class PostanovlenieExtractor:
    """Извлекает данные из Word-документа постановления (поддерживает .docx и .doc)"""

    # Регулярные выражения для поиска
    RE_FIO_BIRTH = re.compile(
        r"у\s+свидетеля\s+([А-ЯЁа-яё]+)\s+([А-ЯЁа-яё]+)\s+([А-ЯЁа-яё]+)\s+(\d{2}\.\d{2}\.\d{4})",
        re.IGNORECASE
    )
    RE_FIO_BIRTH_NOMINATIVE = re.compile(
        r"Свидетелем\s+по\s+уголовному\s+делу[,\s]*является\s+([А-ЯЁа-яё]+)\s+([А-ЯЁа-яё]+)\s+([А-ЯЁа-яё]+)\s+(\d{2}\.\d{2}\.\d{4})",
        re.IGNORECASE
    )
    RE_RESOLUTION_DATE = re.compile(
        r"«(\d{1,2})»\s+([а-яё]+)\s+\d{4}\s+года",
        re.IGNORECASE
    )

    def __init__(self, file_path: Path):
        self.file_path = file_path
        self.text = self._extract_text()

    def _extract_text(self) -> str:
        """Извлекает весь текст из документа .docx или .doc"""
        try:
            # Для .docx используем python-docx
            if self.file_path.suffix.lower() == '.docx':
                doc = Document(self.file_path)
                full_text = [para.text for para in doc.paragraphs]
                return "\n".join(full_text)

            # Для .doc используем win32com (Word Automation)
            elif self.file_path.suffix.lower() == '.doc':
                word = None
                doc = None
                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(str(self.file_path.absolute()))
                    text = doc.Content.Text
                    logger.info(f"Текст из .doc извлечён через win32com")
                    return text
                except Exception as e:
                    logger.error(f"Ошибка при чтении .doc файла {self.file_path}: {e}")
                    raise
                finally:
                    # Закрываем документ и выходим из Word, даже если произошла ошибка
                    if doc:
                        try:
                            doc.Close()
                            logger.debug(f"Документ {self.file_path.name} закрыт")
                        except Exception as e:
                            logger.warning(f"Не удалось закрыть документ {self.file_path.name}: {e}")
                    if word:
                        try:
                            word.Quit()
                            logger.debug("Приложение Word завершено")
                        except Exception as e:
                            logger.warning(f"Не удалось завершить приложение Word: {e}")
            else:
                logger.warning(f"Неподдерживаемый формат файла: {self.file_path.suffix}")
                return ""
        except Exception as e:
            logger.error(f"Ошибка при чтении файла {self.file_path}: {e}")
            return ""

    def extract_fio_birth(self) -> Optional[Dict[str, str]]:
        """
        Извлекает ФИО и дату рождения.
        Возвращает словарь с ключами: 'surname', 'first_name', 'patronymic', 'birth_date', 'case'
        case может быть 'genitive' (родительный падеж) или 'nominative' (именительный).
        Поддерживает 2, 3 и 4 слова в ФИО.
        """
        logger.info(f"Начинаем извлечение ФИО из файла {self.file_path.name}")
        
        # Расширенные регулярные выражения для разного количества слов в ФИО
        
        # Паттерн для родительного падежа (у свидетеля...)
        # Захватываем от 2 до 4 слов ФИО + дату
        RE_FIO_BIRTH_FLEXIBLE = re.compile(
            r"у\s+свидетеля\s+((?:[А-ЯЁа-яё]+\s+){1,3}[А-ЯЁа-яё]+)\s+(\d{2}\.\d{2}\.\d{4})",
            re.IGNORECASE
        )
        
        # Паттерн для именительного падежа (Свидетелем по уголовному делу, является...)
        # Захватываем от 2 до 4 слов ФИО + дату
        RE_FIO_BIRTH_NOMINATIVE_FLEXIBLE = re.compile(
            r"Свидетелем\s+по\s+уголовному\s+делу[,\s]*является\s+((?:[А-ЯЁа-яё]+\s+){1,3}[А-ЯЁа-яё]+)\s+(\d{2}\.\d{2}\.\d{4})",
            re.IGNORECASE
        )
        
        # Ключевые слова для определения сложного отчества
        COMPLEX_PATRONYMIC_KEYWORDS = [
            'угли',   # тюркский: сын (основной вариант)
            'углы',   # альтернативное написание "угли"
            'оглы',   # азербайджанский/тюркский: сын 
            'оглу',   # вариант "оглы"
            'кызы',   # тюркский: дочь
            'гызы',   # азербайджанский: дочь
            'заде',   # персидский: потомок/рождённый
            'бек',    # тюркский титул (иногда часть отчества)
            'хан'     # тюркский/монгольский титул (иногда часть отчества)
        ]
        
        def parse_fio_words(fio_string, case_type):
            """
            Парсит строку ФИО и определяет фамилию, имя и отчество
            """
            logger.info(f"Парсинг ФИО: '{fio_string}' (падеж: {case_type})")
            
            words = fio_string.strip().split()
            word_count = len(words)
            
            logger.info(f"Количество слов в ФИО: {word_count} - {words}")
            
            if word_count < 2:
                logger.error(f"Недостаточно слов в ФИО: {words}")
                return None
            elif word_count == 2:
                # Фамилия + имя (без отчества)
                surname, first_name = words
                patronymic = ""
                logger.info(f"Обнаружено 2 слова: фамилия='{surname}', имя='{first_name}', отчество отсутствует")
            elif word_count == 3:
                # Стандартный случай: фамилия + имя + отчество
                surname, first_name, patronymic = words
                logger.info(f"Обнаружено 3 слова: фамилия='{surname}', имя='{first_name}', отчество='{patronymic}'")
            elif word_count == 4:
                # Возможно сложное отчество
                surname = words[0]
                first_name = words[1]
                
                # Проверяем, есть ли ключевые слова сложного отчества в последнем слове
                last_word_lower = words[3].lower()
                has_complex_keyword = any(keyword.lower() in last_word_lower for keyword in COMPLEX_PATRONYMIC_KEYWORDS)
                
                if has_complex_keyword:
                    # Сложное отчество: слово3 + слово4
                    patronymic = f"{words[2]} {words[3]}"
                    matched_keyword = next((kw for kw in COMPLEX_PATRONYMIC_KEYWORDS if kw.lower() in last_word_lower), "неизвестное")
                    logger.info(f"Обнаружено 4 слова со сложным отчеством (ключевое слово: '{matched_keyword}'): фамилия='{surname}', имя='{first_name}', отчество='{patronymic}'")
                else:
                    # Предполагаем, что это ошибка распознавания или нетипичный случай
                    # Берём первые 3 слова как стандартное ФИО
                    surname, first_name, patronymic = words[:3]
                    logger.warning(f"Обнаружено 4 слова без признаков сложного отчества. Используем первые 3: фамилия='{surname}', имя='{first_name}', отчество='{patronymic}'. Лишнее слово: '{words[3]}'")
            else:
                # Более 4 слов - берём первые 4 и применяем логику для 4 слов
                logger.warning(f"Обнаружено {word_count} слов в ФИО, что больше ожидаемого. Используем первые 4 слова: {words[:4]}")
                return parse_fio_words(' '.join(words[:4]), case_type)
            
            return {
                'surname': surname,
                'first_name': first_name,
                'patronymic': patronymic,
                'case': case_type
            }
        
        # Сначала пытаемся найти новым гибким методом
        
        # 1. Поиск родительного падежа с гибким количеством слов
        logger.debug("Попытка поиска гибким шаблоном родительного падежа...")
        match = RE_FIO_BIRTH_FLEXIBLE.search(self.text)
        if match:
            fio_string, birth_date = match.groups()
            logger.info(f"НАЙДЕН гибкий шаблон родительного падежа: ФИО='{fio_string}', дата='{birth_date}'")
            
            parsed_fio = parse_fio_words(fio_string, 'genitive')
            if parsed_fio:
                parsed_fio['birth_date'] = birth_date
                logger.info(f"Успешно извлечены данные (родительный падеж): {parsed_fio}")
                return parsed_fio
        
        # 2. Поиск именительного падежа с гибким количеством слов
        logger.debug("Попытка поиска гибким шаблоном именительного падежа...")
        match = RE_FIO_BIRTH_NOMINATIVE_FLEXIBLE.search(self.text)
        if match:
            fio_string, birth_date = match.groups()
            logger.info(f"НАЙДЕН гибкий шаблон именительного падежа: ФИО='{fio_string}', дата='{birth_date}'")
            
            parsed_fio = parse_fio_words(fio_string, 'nominative')
            if parsed_fio:
                parsed_fio['birth_date'] = birth_date
                logger.info(f"Успешно извлечены данные (именительный падеж): {parsed_fio}")
                return parsed_fio
        
        # 3. Fallback: используем старые регулярные выражения для совместимости
        logger.info("Гибкие шаблоны не сработали, пробуем старые фиксированные шаблоны для совместимости...")
        
        # Сначала старый шаблон родительного падежа (3 слова)
        logger.debug("Попытка поиска старым шаблоном родительного падежа (3 слова)...")
        match = self.RE_FIO_BIRTH.search(self.text)
        if match:
            surname, first_name, patronymic, birth_date = match.groups()
            logger.info(
                f"НАЙДЕН старый шаблон родительного падежа (3 слова): "
                f"{surname} {first_name} {patronymic}, дата {birth_date}"
            )
            return {
                'surname': surname,
                'first_name': first_name,
                'patronymic': patronymic,
                'birth_date': birth_date,
                'case': 'genitive'
            }

        # Потом старый шаблон именительного падежа (3 слова)
        logger.debug("Попытка поиска старым шаблоном именительного падежа (3 слова)...")
        match = self.RE_FIO_BIRTH_NOMINATIVE.search(self.text)
        if match:
            surname, first_name, patronymic, birth_date = match.groups()
            logger.info(
                f"НАЙДЕН старый шаблон именительного падежа (3 слова): "
                f"{surname} {first_name} {patronymic}, дата {birth_date}"
            )
            return {
                'surname': surname,
                'first_name': first_name,
                'patronymic': patronymic,
                'birth_date': birth_date,
                'case': 'nominative'
            }

        logger.warning(f"НЕ НАЙДЕНЫ ФИО и дата рождения ни одним из методов в файле {self.file_path.name}")
        return None

    def extract_resolution_date(self) -> Optional[Dict[str, str]]:
        """
        Извлекает дату постановления (день и месяц в текстовом виде).
        Возвращает словарь: {'day': str, 'month': str}
        """
        match = self.RE_RESOLUTION_DATE.search(self.text)
        if not match:
            logger.warning(f"Не найдена дата постановления в файле {self.file_path}")
            return None
        day, month = match.groups()
        logger.info(f"Извлечена дата постановления: {day} {month}")
        return {'day': day, 'month': month}


# ========== КЛАСС ДЛЯ РАБОТЫ С КАРТОЙ EXCEL ==========
class ExcelCardFiller:
    """Заполняет карту Excel данными"""

    def __init__(self, template_path: Path):
        if not template_path.exists():
            raise FileNotFoundError(f"Шаблон карты не найден: {template_path}")
        self.template_path = template_path

    def create_copy(self, new_name: str, output_dir: Optional[Path] = None) -> Path:
        """
        Создаёт копию шаблона с новым именем.
        Если указан output_dir, копия сохраняется в эту папку, иначе рядом с шаблоном.
        Возвращает путь к созданной копии.
        """
        # Определяем целевую папку
        if output_dir is None:
            target_dir = self.template_path.parent
        else:
            target_dir = Path(output_dir)
            # Создаём папку, если её нет
            try:
                target_dir.mkdir(parents=True, exist_ok=True)
                logger.info(f"Папка для сохранения карт: {target_dir}")
            except Exception as e:
                logger.error(f"Не удалось создать папку {target_dir}: {e}")
                raise

        new_path = target_dir / new_name
        if new_path.exists():
            logger.warning(f"Файл {new_path} уже существует, он будет перезаписан")
        shutil.copy2(self.template_path, new_path)
        logger.info(f"Создана копия шаблона: {new_path}")
        return new_path

    def fill_card(self, workbook: openpyxl.Workbook, data: Dict[str, str]) -> None:
        """
        Заполняет ячейки карты значениями из словаря data.
        Ключи словаря – коды, значения – текст для вставки.
        """
        worksheet = workbook.active
        for code, value in data.items():
            cell = find_cell_by_code(worksheet, code)
            if cell:
                cell.value = value
                logger.debug(f"Записано {code} = {value}")
            else:
                logger.warning(f"Код {code} не найден в карте, значение не записано")

    def convert_to_nominative(self, word: str) -> str:
        """
        Преобразует слово из родительного падежа в именительный.
        Возвращает исходное слово, если преобразование невозможно.
        """
        if not word or not isinstance(word, str):
            return word
        word = word.strip()
        if not word:
            return word
        try:
            parsed = morph.parse(word)[0]
            nom = parsed.inflect({'nomn'})
            if nom:
                return nom.word.capitalize()
            else:
                logger.warning(f"Не удалось преобразовать слово '{word}' в именительный падеж")
                return word
        except Exception as e:
            logger.error(f"Ошибка при преобразовании слова '{word}': {e}")
            return word

    def convert_to_genitive(self, word: str) -> str:
        """
        Преобразует слово из именительного падежа в родительный.
        Возвращает исходное слово, если преобразование невозможно.
        """
        if not word or not isinstance(word, str):
            return word
        word = word.strip()
        if not word:
            return word
        try:
            parsed = morph.parse(word)[0]
            gen = parsed.inflect({'gent'})
            if gen:
                return gen.word.capitalize()
            else:
                logger.warning(f"Не удалось преобразовать слово '{word}' в родительный падеж")
                return word
        except Exception as e:
            logger.error(f"Ошибка при преобразовании слова '{word}': {e}")
            return word

    def fill_nominative_names(self, workbook: openpyxl.Workbook) -> None:
        """
        Находит в карте строки с кодами ФР, ИР, ОР (родительный падеж),
        преобразует их в именительный и записывает в ФИ, ИИ, ОИ.
        """
        worksheet = workbook.active
        fr_cell = find_cell_by_code(worksheet, "ФР")
        ir_cell = find_cell_by_code(worksheet, "ИР")
        or_cell = find_cell_by_code(worksheet, "ОР")

        fr = normalize_text(fr_cell.value if fr_cell else None)
        ir = normalize_text(ir_cell.value if ir_cell else None)
        or_ = normalize_text(or_cell.value if or_cell else None)

        fi = self.convert_to_nominative(fr) if fr else ""
        ii = self.convert_to_nominative(ir) if ir else ""
        oi = self.convert_to_nominative(or_) if or_ else ""

        data = {}
        if fi:
            data["ФИ"] = fi
        if ii:
            data["ИИ"] = ii
        if oi:
            data["ОИ"] = oi

        if data:
            self.fill_card(workbook, data)
            logger.info(f"Записаны именительные формы: ФИ={fi}, ИИ={ii}, ОИ={oi}")
        else:
            logger.info("Нет данных для преобразования в именительный падеж")


# ========== ФУНКЦИИ ОБРАБОТКИ РАЗНЫХ ТИПОВ РАБОТ ==========
def process_tobolsk_newborn(script_dir: Path, template_path: Path, start_number: int) -> None:
    """
    Обработка постановлений из Тобольска по новорожденным.
    """
    # Запрос пути к папке с постановлениями
    user_input = input("Введите путь к папке с постановлениями (или нажмите Enter для значения по умолчанию): ").strip()
    if user_input:
        postanovleniya_dir = user_input
    else:
        postanovleniya_dir = POSTANOVLENIYA_DIR_DEFAULT
        logger.info(f"Используется путь по умолчанию: {postanovleniya_dir}")

    postanovleniya_path = Path(postanovleniya_dir)
    if not postanovleniya_path.exists():
        logger.error(f"Папка с постановлениями не найдена: {postanovleniya_path}")
        logger.info("Операция прервана из-за отсутствия папки с постановлениями.")
        return
    logger.info(f"Папка с постановлениями найдена: {postanovleniya_path}")

    # НОВОЕ: запрос пути для сохранения карт
    user_input_output = input("Введите путь для сохранения карт (или нажмите Enter для значения по умолчанию): ").strip()
    if user_input_output:
        output_cards_dir = user_input_output
    else:
        output_cards_dir = OUTPUT_CARDS_DIR_DEFAULT
        if not output_cards_dir:
            # Если константа пуста, сохраняем в папку программы (прежнее поведение)
            output_cards_dir = script_dir
            logger.info(f"Константа OUTPUT_CARDS_DIR_DEFAULT пуста, карты будут сохранены в папку программы: {output_cards_dir}")
        else:
            logger.info(f"Используется путь по умолчанию: {output_cards_dir}")

    output_cards_path = Path(output_cards_dir)
    try:
        # Создаём папку, если её нет (даже если это script_dir, ничего страшного)
        output_cards_path.mkdir(parents=True, exist_ok=True)
        logger.info(f"Папка для сохранения карт: {output_cards_path}")
    except Exception as e:
        logger.error(f"Не удалось создать/проверить папку {output_cards_path}: {e}")
        logger.info("Операция прервана из-за проблем с папкой сохранения.")
        return

    # Проверка существования шаблона карты
    if not template_path.exists():
        logger.error(f"Шаблон карты не найден: {template_path}")
        logger.info("Операция прервана из-за отсутствия шаблона.")
        return
    logger.info(f"Шаблон карты найден: {template_path}")

    # Поиск файлов .docx и .doc
    all_files = list(postanovleniya_path.glob("*"))
    doc_files = [f for f in all_files if f.suffix.lower() in ('.docx', '.doc')]
    if not doc_files:
        logger.warning(f"В папке {postanovleniya_path} не найдено файлов .docx или .doc")
        logger.info("Операция прервана из-за отсутствия постановлений.")
        return
    logger.info(f"Найдено {len(doc_files)} файлов постановлений")

    current_number = start_number
    for i, doc_file in enumerate(doc_files, 1):
        logger.info(f"Обработка файла {i}/{len(doc_files)}: {doc_file.name}")
        try:
            extractor = PostanovlenieExtractor(doc_file)
            fio_data = extractor.extract_fio_birth()
            if not fio_data:
                logger.warning(f"Пропуск файла {doc_file.name} (не найдены ФИО/дата рождения)")
                continue

            # Получаем извлечённые данные
            surname = fio_data['surname']
            first_name = fio_data['first_name']
            patronymic = fio_data['patronymic']
            birth_date = fio_data['birth_date']
            case = fio_data['case']

            # Создаём экземпляр заполнителя карты (для доступа к методам преобразования)
            card_filler = ExcelCardFiller(template_path)

            # Определяем именительную и родительную формы
            if case == 'genitive':
                surname_gen = surname
                first_name_gen = first_name
                patronymic_gen = patronymic
                # Преобразуем в именительный для названия и кодов ФИ/ИИ/ОИ
                surname_nom = card_filler.convert_to_nominative(surname_gen)
                first_name_nom = card_filler.convert_to_nominative(first_name_gen)
                patronymic_nom = card_filler.convert_to_nominative(patronymic_gen)
                logger.info(f"Исходный падеж родительный -> имя: {surname_nom} {first_name_nom} {patronymic_nom}")
            elif case == 'nominative':
                surname_nom = surname
                first_name_nom = first_name
                patronymic_nom = patronymic
                # Преобразуем в родительный для кодов ФР/ИР/ОР
                surname_gen = card_filler.convert_to_genitive(surname_nom)
                first_name_gen = card_filler.convert_to_genitive(first_name_nom)
                patronymic_gen = card_filler.convert_to_genitive(patronymic_nom)
                logger.info(f"Исходный падеж именительный -> род.: {surname_gen} {first_name_gen} {patronymic_gen}")
            else:
                logger.error(f"Неизвестный падеж '{case}' в файле {doc_file.name}, пропуск")
                continue

            # Формируем имя файла карты (всегда в именительном падеже)
            new_card_name = f"{current_number}-26 Тобольск {surname_nom} {first_name_nom} {patronymic_nom}.xlsx"
            # НОВОЕ: передаём путь для сохранения карты
            new_card_path = card_filler.create_copy(new_card_name, output_cards_path)

            # Загружаем книгу
            wb = openpyxl.load_workbook(new_card_path)

            # Получаем текущую дату для ДН и МН
            current_day, current_month_gen = get_current_date()

            # Извлекаем дату постановления
            resolution_date = extractor.extract_resolution_date()
            if not resolution_date:
                resolution_date = {'day': '', 'month': ''}

            # Формируем словарь для заполнения (включаем оба набора падежей)
            fill_data = {
                "НОМ": str(current_number),
                "ФР": surname_gen,
                "ИР": first_name_gen,
                "ОР": patronymic_gen,
                "ФИ": surname_nom,
                "ИИ": first_name_nom,
                "ОИ": patronymic_nom,
                "ДР": birth_date,
                "ДП": resolution_date['day'],
                "МП": resolution_date['month'],
                "ДН": str(current_day),
                "МН": current_month_gen,
            }

            investigator_data = find_investigator(extractor.text)
            if investigator_data:
                fill_data.update(investigator_data)
                logger.info(f"Добавлены данные следователя: {investigator_data}")
            else:
                logger.info("Данные следователя не добавлены")

            # Заполняем карту
            card_filler.fill_card(wb, fill_data)
            wb.save(new_card_path)
            logger.info(f"Карта успешно создана и заполнена: {new_card_path}")

            current_number += 1

        except Exception as e:
            logger.error(f"Ошибка при обработке файла {doc_file.name}: {e}", exc_info=True)
            continue


def process_bone_cards(script_dir: Path) -> None:
    """
    Заполнение карт Excel значениями ДО, МО и инкрементом ОБ.
    """
    logger.info("Запущен режим заполнения Карт с костями")

    # Запрос значений у пользователя
    print("\n--- Заполнение Карт с костями ---")
    do_value = input("Число окончания экспертизы (например, 1): ").strip()
    mo_value = input("Месяц окончания экспертизы в р.п. (например, января): ").strip()

    while True:
        ob_start = input("Начальный объект (например, 1A): ").strip()
        is_valid, _, _ = validate_ob_format(ob_start)
        if is_valid:
            break
        print("Неверный формат! Введите в формате 'числоБуква' (например, 1A)")

    logger.info(f"Введённые значения: ДО='{do_value}', МО='{mo_value}', ОБ начало='{ob_start}'")

    # Проверка формата ОБ (повторно, чтобы получить start_num и letter)
    is_valid, start_num, letter = validate_ob_format(ob_start)
    if not is_valid:
        logger.error(f"Неверный формат ОБ: '{ob_start}'. Обработка прервана.")
        return

    # Получение списка Excel-файлов в текущей папке программы
    excel_files = list(script_dir.glob("*.xlsx")) + list(script_dir.glob("*.xlsm"))
    # Исключаем временные файлы
    excel_files = [f for f in excel_files if not f.name.startswith('~$')]

    if not excel_files:
        logger.warning("В текущей папке не найдено Excel-файлов")
        return

    # Натуральная сортировка
    excel_files.sort(key=lambda x: natural_sort_key(x.name))

    logger.info(f"Найдено файлов для обработки: {len(excel_files)}")
    logger.info(f"Порядок обработки: {[f.name for f in excel_files]}")

    # Обработка каждого файла
    for i, file_path in enumerate(excel_files):
        try:
            logger.info(f"Обработка файла: {file_path.name}")

            # Загрузка книги
            wb = openpyxl.load_workbook(file_path)
            ws = wb.active

            # 1. ДО
            do_cell = find_cell_by_code(ws, "ДО")
            if do_cell:
                do_cell.value = do_value
                logger.info(f"  ДО: установлено '{do_value}'")
            else:
                logger.warning(f"  Код 'ДО' не найден в файле {file_path.name}")

            # 2. МО
            mo_cell = find_cell_by_code(ws, "МО")
            if mo_cell:
                mo_cell.value = mo_value
                logger.info(f"  МО: установлено '{mo_value}'")
            else:
                logger.warning(f"  Код 'МО' не найден в файле {file_path.name}")

            # 3. ОБ с инкрементом
            current_ob = f"{start_num + i}{letter}"
            ob_cell = find_cell_by_code(ws, "ОБ")
            if ob_cell:
                ob_cell.value = current_ob
                logger.info(f"  ОБ: установлено '{current_ob}'")
            else:
                logger.warning(f"  Код 'ОБ' не найден в файле {file_path.name}")

            # 4. ОБРТ = НОМ + пробел + ОБ
            obrt_cell = find_cell_by_code(ws, "ОБРТ")
            if obrt_cell:
                nom_cell = find_cell_by_code(ws, "НОМ")
                ob_cell = find_cell_by_code(ws, "ОБ")
                if nom_cell and nom_cell.value and ob_cell and ob_cell.value:
                    obrt_value = f"{nom_cell.value} {ob_cell.value}"
                    obrt_cell.value = obrt_value
                    logger.info(f"  ОБРТ: установлено '{obrt_value}'")
                else:
                    logger.warning(f"  Не удалось получить НОМ или ОБ для ОБРТ в файле {file_path.name}")
            else:
                logger.warning(f"  Код 'ОБРТ' не найден в файле {file_path.name}")

            wb.save(file_path)
            logger.info(f"Файл {file_path.name} успешно сохранён")

        except Exception as e:
            logger.error(f"Ошибка при обработке {file_path.name}: {e}", exc_info=True)
            continue

    logger.info("Обработка всех файлов завершена")


def process_registry_cards(script_dir: Path) -> None:
    """
    Режим «Регистратура: заполнение Карт с образцами в случае прямой идентификации».
    Заполняет все Excel-файлы в папке программы заданными значениями,
    выполняет копирование (ФПИ->ФИ и т.д.), преобразование в родительный падеж
    (ФПИ->ФР и т.д.) и удаляет значение из строки РОД.
    """
    logger.info("Запущен режим «Регистратура: заполнение Карт с образцами»")

    # 1. Выбор региона
    print("\nС какими образцами работаем?")
    print("1. Тюмень")
    print("2. Ростов")
    region_choice = input("Введите номер варианта (1 или 2): ").strip()
    while region_choice not in ('1', '2'):
        print("Неверный выбор. Введите 1 или 2.")
        region_choice = input("Введите номер варианта (1 или 2): ").strip()

    if region_choice == '1':
        code_1_value = "СВО_Молов_образец_прямая идентификация"
        logger.info("Выбран регион: Тюмень")
    else:
        code_1_value = "СВО_Ростов_образец_прямая идентификация"
        logger.info("Выбран регион: Ростов")

    data_to_fill = {
        "1": code_1_value,
        "ТО": "крови",
        "НО": "марле",
        "НАД": "с отпечатанной на принтере надписью",
        "МАТ": "бурого пятна на марле"
    }

    # 2. Поиск Excel-файлов
    excel_files = list(script_dir.glob("*.xlsx")) + list(script_dir.glob("*.xlsm"))
    excel_files = [f for f in excel_files if not f.name.startswith('~$')]
    if not excel_files:
        logger.warning("В текущей папке не найдено Excel-файлов (карт)")
        return

    excel_files.sort(key=lambda x: natural_sort_key(x.name))
    logger.info(f"Найдено файлов для обработки: {len(excel_files)}")
    logger.info(f"Порядок обработки: {[f.name for f in excel_files]}")

    # --- СОЗДАНИЕ ЭКЗЕМПЛЯРА ДЛЯ ПРЕОБРАЗОВАНИЯ ПАДЕЖЕЙ ---
    # Берём первый существующий Excel-файл, чтобы создать экземпляр класса
    # (сам файл не используется, только методы преобразования)
    template_path = excel_files[0]
    card_filler = ExcelCardFiller(template_path)
    logger.info(f"Создан экземпляр ExcelCardFiller на основе файла {template_path.name}")

    for file_path in excel_files:
        try:
            logger.info(f"Обработка файла: {file_path.name}")
            wb = openpyxl.load_workbook(file_path)
            ws = wb.active

            # 2.1 Заполнение основными значениями (используем улучшенную глобальную функцию)
            for code, value in data_to_fill.items():
                cell = find_cell_by_code(ws, code)
                if cell:
                    cell.value = value
                    logger.info(f"  Код '{code}': установлено значение '{value}'")
                else:
                    logger.warning(f"  Код '{code}' не найден в файле {file_path.name}")

            # Собираем все ячейки с кодами для быстрого доступа (ключи – строки без пробелов)
            cells = {}
            for row in ws.iter_rows(min_row=1, max_col=3):
                code_cell = row[0]
                if code_cell and code_cell.value is not None:
                    code_str = str(code_cell.value).strip()
                    if code_str:
                        cells[code_str] = row[2]  # ячейка значения

            # 2.2 Копирование
            copy_pairs = [('ФПИ', 'ФИ'), ('ИПИ', 'ИИ'), ('ОПИ', 'ОИ'), ('ДРП', 'ДР')]
            for src, dst in copy_pairs:
                if src in cells and dst in cells:
                    src_val = cells[src].value
                    if src_val is not None:
                        cells[dst].value = src_val
                        logger.info(f"      Скопировано {src} -> {dst}: '{src_val}'")
                    else:
                        logger.info(f"      Значение {src} пустое, копирование пропущено")
                else:
                    missing = [x for x in (src, dst) if x not in cells]
                    logger.warning(f"      Коды {', '.join(missing)} не найдены, копирование пропущено")

            # 2.3 Преобразование в родительный падеж (используем метод экземпляра)
            genitive_pairs = [('ФПИ', 'ФР'), ('ИПИ', 'ИР'), ('ОПИ', 'ОР')]
            for src, dst in genitive_pairs:
                if src in cells and dst in cells:
                    src_val = cells[src].value
                    if src_val and isinstance(src_val, str):
                        gen = card_filler.convert_to_genitive(src_val.strip())
                        cells[dst].value = gen
                        logger.info(f"      Преобразовано {src} -> {dst}: '{src_val}' -> '{gen}'")
                    else:
                        logger.info(f"      Значение {src} пустое или не строка, преобразование пропущено")
                else:
                    missing = [x for x in (src, dst) if x not in cells]
                    logger.warning(f"      Коды {', '.join(missing)} не найдены, преобразование пропущено")

            # 2.4 Удаление значения из РОД
            if 'РОД' in cells:
                rod_cell = cells['РОД']
                if rod_cell.value is not None:
                    logger.info(f"      Удалено значение из РОД: '{rod_cell.value}'")
                    rod_cell.value = None
                else:
                    logger.info("      Значение РОД уже пустое")
            else:
                logger.warning("      Код РОД не найден, удаление пропущено")

            wb.save(file_path)
            logger.info(f"Файл {file_path.name} успешно сохранён")

        except PermissionError:
            logger.error(f"Не удалось сохранить {file_path.name}: файл открыт в другой программе")
        except Exception as e:
            logger.error(f"Ошибка при обработке {file_path.name}: {e}", exc_info=True)
            continue

    logger.info("Обработка всех файлов завершена")



def process_distribution_cards(script_dir: Path) -> None:
    """
    Режим «Регистратура: Заполнение Карт с костями для РАСПРЕДЕЛЕНИЯ между экспертами».
    Поиск папок экспертов, в каждой папке чтение .txt-файла и создание Excel-карт.
    """
    logger.info("Запущен режим распределения карт (на основе .txt-файлов в папках экспертов)")

    # 1. Находим шаблон Excel
    template_path = find_template_excel(script_dir)
    if not template_path:
        logger.error("Не удалось найти шаблон карты")
        input("Нажмите Enter для завершения...")
        return

    # 2. Запрос начального номера экспертизы
    while True:
        try:
            start_num_str = input("Введите номер, с которого начать нумерацию экспертиз: ").strip()
            start_number = int(start_num_str)
            if start_number <= 0:
                raise ValueError
            break
        except ValueError:
            print("Ошибка: пожалуйста, введите целое положительное число")
            logger.warning("Пользователь ввёл неверный номер экспертизы")

    # 3. Запрос даты начала экспертизы
    while True:
        date_str = input("Введите дату начала экспертизы в формате ДД.ММ.ГГГГ: ").strip()
        try:
            if not re.match(r'^\d{2}\.\d{2}\.\d{4}$', date_str):
                raise ValueError
            day, month, year = map(int, date_str.split('.'))
            # Проверка валидности даты
            datetime(year, month, day)
            start_day = day
            start_month_gen = month_num_to_genitive(month)
            break
        except (ValueError, TypeError):
            print("Ошибка: неверный формат даты. Используйте формат ДД.ММ.ГГГГ")
            logger.warning("Пользователь ввёл неверную дату начала")

    logger.info(f"Параметры: начальный номер={start_number}, дата начала={start_day:02d}.{month:02d}.{year}")

    # 4. Получаем список папок (экспертов) в корневой папке
    expert_folders = [d for d in script_dir.iterdir() if d.is_dir()]
    if not expert_folders:
        logger.warning("В корневой папке не найдено папок экспертов")
        input("Нажмите Enter для завершения...")
        return

    expert_folders.sort(key=lambda x: x.name)  # сортируем по имени
    logger.info(f"Найдено папок экспертов: {len(expert_folders)}")
    logger.info(f"Папки: {', '.join(f.name for f in expert_folders)}")

    # Подтверждение
    confirm = input(f"\nБудет создано несколько Excel файлов, начиная с номера {start_number}.\n"
                    f"Дата начала экспертизы: {start_day}.{month:02d}.{year}\n"
                    f"Папки для обработки: {len(expert_folders)}\n\n"
                    "Продолжить? (да/нет): ")
    if confirm.lower() not in ['да', 'д', 'yes', 'y']:
        logger.info("Операция отменена пользователем")
        return

    current_number = start_number
    total_files_created = 0
    results = []

    for folder in expert_folders:
        # Ищем первый .txt файл в папке
        txt_files = list(folder.glob("*.txt"))
        if not txt_files:
            logger.warning(f"В папке '{folder.name}' не найден .txt файл, пропускаем")
            results.append(f"Папка '{folder.name}': нет .txt файла")
            continue

        txt_file = txt_files[0]
        rows = parse_txt_file(txt_file)
        if not rows:
            logger.warning(f"В файле '{txt_file.name}' (папка '{folder.name}') нет данных")
            results.append(f"Папка '{folder.name}': файл '{txt_file.name}' не содержит данных")
            continue

        folder_files_created = 0
        for row in rows:
            # Проверка формата даты в строке
            try:
                date_parts = row['date'].split('.')
                if len(date_parts) != 3:
                    raise ValueError
                day_from_row = int(date_parts[0])
                month_from_row = int(date_parts[1])
                year_from_row = int(date_parts[2])
                # Проверка валидности
                datetime(year_from_row, month_from_row, day_from_row)
                month_text = month_num_to_genitive(month_from_row)
            except (ValueError, TypeError):
                logger.error(f"Неверный формат даты '{row['date']}' в файле {txt_file.name}, строка пропущена")
                continue

            # Данные для заполнения
            expertise_data = {
                'expertise_number': str(current_number),
                'corpse_number': row['corpse_number'],
                'day': str(day_from_row),
                'month': month_text,
                'start_day': str(start_day),
                'start_month': start_month_gen
            }

            output_filename = f"{current_number}.xlsx"
            output_path = folder / output_filename

            if fill_excel_from_template(template_path, output_path, expertise_data):
                folder_files_created += 1
                total_files_created += 1
                logger.info(f"Создан файл: {output_path} (номер={current_number}, труп={row['corpse_number']})")
            else:
                logger.error(f"Не удалось создать файл {output_path}")

            current_number += 1

        results.append(f"Папка '{folder.name}': создано {folder_files_created} файлов")

    # Вывод статистики
    print("\n" + "=" * 50)
    print("РЕЗУЛЬТАТЫ:")
    print("=" * 50)
    print(f"Всего создано файлов: {total_files_created}")
    print(f"Последний использованный номер: {current_number - 1}")
    print("\nДетали по папкам:")
    for detail in results:
        print(f"  • {detail}")

    # Сохранение лога
    log_file = script_dir / "distribution_log.txt"
    try:
        with open(log_file, 'w', encoding='utf-8') as f:
            f.write(f"Дата обработки: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}\n")
            f.write(f"Начальный номер: {start_number}\n")
            f.write(f"Дата начала экспертизы: {start_day}.{month:02d}.{year}\n")
            f.write(f"Всего создано файлов: {total_files_created}\n")
            f.write(f"Последний номер: {current_number - 1}\n\n")
            f.write("Детали:\n")
            for detail in results:
                f.write(f"{detail}\n")
        logger.info(f"Лог сохранён в {log_file}")
    except Exception as e:
        logger.error(f"Не удалось сохранить лог: {e}")

    input("\nНажмите Enter для завершения...")


def process_mass_fill_cards_with_samples(script_dir: Path) -> None:
    """
    Режим «Регистратура: Массовое заполнение Карт с ОБРАЗЦАМИ».
    Реализует логику первого кода: удаление пробелов в строках с заданными кодами,
    поиск ФР, ИР, ОР, преобразование их в именительный падеж и запись в ФИ, ИИ, ОИ.
    """
    logger.info("Запущен режим массового заполнения карт с образцами (преобразование ФИО из родительного в именительный)")

    # Коды, для которых нужно удалить лишние пробелы
    codes_for_space_removal = ['ФР', 'ИР', 'ОР', 'ФОР', 'ИОР', 'ООР', 'ФПИ', 'ИПИ', 'ОПИ']

    # Поиск Excel-файлов (только .xlsx и .xlsm)
    excel_files = list(script_dir.glob("*.xlsx")) + list(script_dir.glob("*.xlsm"))
    excel_files = [f for f in excel_files if not f.name.startswith('~$')]

    if not excel_files:
        logger.warning("В текущей папке не найдено Excel-файлов (карт)")
        return

    excel_files.sort(key=lambda x: natural_sort_key(x.name))
    logger.info(f"Найдено файлов для обработки: {len(excel_files)}")
    logger.info(f"Порядок обработки: {[f.name for f in excel_files]}")

    # --- СОЗДАНИЕ ЭКЗЕМПЛЯРА ДЛЯ ПРЕОБРАЗОВАНИЯ ПАДЕЖЕЙ ---
    # Берём первый существующий Excel-файл, чтобы создать экземпляр класса
    template_path = excel_files[0]
    card_filler = ExcelCardFiller(template_path)
    logger.info(f"Создан экземпляр ExcelCardFiller на основе файла {template_path.name}")

    processed_count = 0
    error_count = 0
    skipped_count = 0

    for file_path in excel_files:
        logger.info(f"Обработка файла: {file_path.name}")
        try:
            wb = openpyxl.load_workbook(file_path)
            ws = wb.active
            changes_made = False
            space_removal_changes_made = False

            # ---------- 1. Удаление лишних пробелов ----------
            logger.debug("Удаление лишних пробелов в строках с кодами: %s", codes_for_space_removal)
            for row in ws.iter_rows(min_row=1, max_col=3):
                code_cell = row[0]
                value_cell = row[2] if len(row) > 2 else None
                if code_cell is None or code_cell.value is None:
                    continue
                code = normalize_text(str(code_cell.value))
                if code in codes_for_space_removal and value_cell is not None:
                    original_value = value_cell.value
                    stripped_value = normalize_text(original_value)
                    if original_value != stripped_value:
                        value_cell.value = stripped_value
                        logger.info(f"Удалены лишние пробелы в строке с кодом '{code}': '{original_value}' -> '{stripped_value}'")
                        space_removal_changes_made = True
                    else:
                        logger.debug(f"Лишние пробелы не найдены в строке с кодом '{code}': '{original_value}'")

            # ---------- 2. Поиск ФР, ИР, ОР ----------
            fr_value = None
            ir_value = None
            or_value = None
            for row in ws.iter_rows(min_row=1, max_col=3):
                code_cell = row[0]
                value_cell = row[2] if len(row) > 2 else None
                if code_cell is None or code_cell.value is None:
                    continue
                code = normalize_text(str(code_cell.value))
                if code == 'ФР' and value_cell is not None:
                    fr_value = value_cell.value
                    logger.info(f"Найдено ФР: '{fr_value}'")
                elif code == 'ИР' and value_cell is not None:
                    ir_value = value_cell.value
                    logger.info(f"Найдено ИР: '{ir_value}'")
                elif code == 'ОР' and value_cell is not None:
                    or_value = value_cell.value
                    logger.info(f"Найдено ОР: '{or_value}'")

            # ---------- 3. Преобразование в именительный падеж (используем метод экземпляра) ----------
            nominative_values = {}
            if fr_value is not None:
                fi = card_filler.convert_to_nominative(fr_value)
                nominative_values['ФИ'] = fi
                if fi != fr_value:
                    logger.info(f"Фамилия преобразована: '{fr_value}' -> '{fi}'")
                else:
                    logger.info(f"Фамилия не требует преобразования: '{fr_value}'")
            if ir_value is not None:
                ii = card_filler.convert_to_nominative(ir_value)
                nominative_values['ИИ'] = ii
                if ii != ir_value:
                    logger.info(f"Имя преобразовано: '{ir_value}' -> '{ii}'")
                else:
                    logger.info(f"Имя не требует преобразования: '{ir_value}'")
            if or_value is not None:
                oi = card_filler.convert_to_nominative(or_value)
                nominative_values['ОИ'] = oi
                if oi != or_value:
                    logger.info(f"Отчество преобразовано: '{or_value}' -> '{oi}'")
                else:
                    logger.info(f"Отчество не требует преобразования: '{or_value}'")

            # ---------- 4. Запись в ячейки с кодами ФИ, ИИ, ОИ ----------
            if nominative_values:
                for code, value in nominative_values.items():
                    cell = find_cell_by_code(ws, code)
                    if cell:
                        cell.value = value
                        logger.info(f"Записано {code}: '{value}'")
                        changes_made = True
                    else:
                        logger.warning(f"Код {code} не найден в файле {file_path.name}, значение не записано")
            else:
                logger.info("Не найдено ни одной части ФИО в родительном падеже для обработки")

            # ---------- 5. Сохранение, если были изменения ----------
            if changes_made or space_removal_changes_made:
                wb.save(file_path)
                logger.info(f"Файл {file_path.name} сохранён")
                processed_count += 1
            else:
                logger.info(f"Файл {file_path.name} не требует обработки (нет изменений)")
                skipped_count += 1

        except Exception as e:
            logger.error(f"Ошибка при обработке файла {file_path.name}: {e}", exc_info=True)
            error_count += 1

    # Итоговая статистика
    logger.info("=" * 50)
    logger.info("ИТОГИ ОБРАБОТКИ (массовое заполнение карт с образцами):")
    logger.info(f"Всего файлов найдено: {len(excel_files)}")
    logger.info(f"Успешно обработано (с изменениями): {processed_count}")
    logger.info(f"Пропущено (без изменений): {skipped_count}")
    logger.info(f"Ошибок: {error_count}")
    logger.info("=" * 50)


def process_mass_print_postanovleniy():
    """Режим «Регистратура: Массовая печать постановлений»."""
    logger.info("Запущен режим массовой печати постановлений")
    logger.info(f"Папка для сохранения постановлений: {MASSPRINT_SAVEPATH_DEFAULT}")
    logger.info(f"Папка с шаблонами: {TEMPLATES_FOLDER_DEFAULT}")

    # ---- Вспомогательные функции ----
    def get_excel_files(folder_path):
        """Получает список файлов Excel в указанной папке, отсортированных естественным образом."""
        excel_files = [f for f in os.listdir(folder_path) if f.endswith(('.xlsx', '.xls')) and not f.startswith('~$')]
        excel_files.sort(key=natural_sort_key)
        logger.info(f"Найдено файлов Excel: {len(excel_files)}")
        for i, f in enumerate(excel_files, 1):
            logger.info(f"  {i}. {f}")
        return [os.path.join(folder_path, f) for f in excel_files]

    def read_card(file_path):
        """Читает данные из карты Excel, используя поиск по кодам (аналог find_cell_by_code)."""
        data = {}
        try:
            wb = openpyxl.load_workbook(file_path)
            ws = wb.active
            for row in ws.iter_rows(min_row=1, max_col=3):
                code_cell = row[0]
                value_cell = row[2] if len(row) > 2 else None
                if code_cell is None or code_cell.value is None:
                    continue
                code = str(code_cell.value).strip()
                if not code:
                    continue
                value = value_cell.value if value_cell is not None else None
                # Специальная обработка для кода "ДБИ"
                if code == "ДБИ" and value is not None and str(value).strip():
                    value = f"с {value.strip()} г."
                data[code] = value
            logger.debug(f"Из файла {file_path} прочитаны коды: {list(data.keys())}")
        except Exception as e:
            logger.error(f"Ошибка при чтении файла {file_path}: {e}", exc_info=True)
            raise
        return data

    def check_card_data(card_data):
        """Проверяет наличие обязательных полей."""
        required_fields = ["2", "3", "4"]
        for field in required_fields:
            if field not in card_data or card_data[field] is None:
                raise ValueError(f"Ошибка: в Карте отсутствует значение для кода '{field}'.")

    def copy_template(card_data, template_name, save_path, template_folder):
        """Копирует шаблон Word в указанное место."""
        for ext in [".doc", ".docx"]:
            template_path = template_folder / f"{template_name}{ext}"
            if template_path.exists():
                output_path = Path(save_path) / f"{card_data['ФИ']}-26{ext}"
                shutil.copy(str(template_path), str(output_path))
                return str(output_path)
        raise FileNotFoundError(f"Файл шаблона '{template_name}' не найден в папке {template_folder}.")

    def replace_in_doc(doc_path, replacements):
        """Заменяет метки вида {код} в Word-документе."""
        doc = Document(doc_path)
        found_keys = set()

        def replace_text(paragraph):
            nonlocal found_keys
            text = paragraph.text
            if not paragraph.runs:
                return
            first_run = paragraph.runs[0]
            base_formatting = {
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline,
                'font': first_run.font.name,
                'size': first_run.font.size
            }
            modified_text = text
            any_replacements = False
            for key, value in replacements.items():
                patterns = [f"{{{key}}}", f"{{ {key} }}", f"{{{key} }}", f"{{ {key}}}"]
                for pattern in patterns:
                    if pattern in modified_text:
                        modified_text = modified_text.replace(pattern, str(value))
                        found_keys.add(key)
                        any_replacements = True
            if any_replacements:
                for run in paragraph.runs:
                    run.text = ""
                new_run = paragraph.add_run(modified_text)
                new_run.bold = base_formatting['bold']
                new_run.italic = base_formatting['italic']
                new_run.underline = base_formatting['underline']
                new_run.font.name = base_formatting['font']
                new_run.font.size = base_formatting['size']

        for p in doc.paragraphs:
            replace_text(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text(p)
        for section in doc.sections:
            for p in section.header.paragraphs:
                replace_text(p)
            for p in section.footer.paragraphs:
                replace_text(p)

        missing_keys = set(replacements.keys()) - found_keys
        if missing_keys:
            logger.warning(f"Предупреждение: следующие метки не найдены в документе: {missing_keys}")

        doc.save(doc_path)
        logger.info(f"Все метки заменены в документе: {doc_path}")

    def print_document(doc_path):
        """Печатает документ Word."""
        try:
            printer_name = win32print.GetDefaultPrinter()
            logger.info(f"ОТПРАВКА НА ПЕЧАТЬ: {os.path.basename(doc_path)}")
            time.sleep(3)
            win32api.ShellExecute(0, "print", doc_path, None, ".", 0)
            logger.info(f"Документ {doc_path} отправлен на печать на принтере {printer_name}.")
            time.sleep(3)
        except Exception as e:
            logger.error(f"Ошибка при печати документа {doc_path}: {e}", exc_info=True)

    def process_multiple_cards(folder_path, template_choice):
        """Обрабатывает все файлы Excel в указанной папке по очереди."""
        excel_files = get_excel_files(folder_path)
        logger.info(f"=== НАЧАЛО ОБРАБОТКИ {len(excel_files)} ФАЙЛОВ ===")

        for index, card_path in enumerate(excel_files, 1):
            output_path = None
            try:
                logger.info(f"=== ОБРАБОТКА ФАЙЛА {index}/{len(excel_files)}: {os.path.basename(card_path)} ===")

                card_data = read_card(card_path)
                check_card_data(card_data)

                # Replacements создаётся для всех кодов, даже пустых
                empty_keys = [code for code, value in card_data.items() if value is None]
                if empty_keys:
                    logger.info(f"Следующие коды имеют пустые значения и будут заменены на пустую строку: {empty_keys}")

                replacements = {code: str(value) if value is not None else "" for code, value in card_data.items()}

                # Выбор шаблона и пути сохранения
                if template_choice == 1:
                    template_name = "Постановление экспертиза ФОНД"
                elif template_choice == 2:
                    template_name = "Постановление экспертиза ФОНД прямая идентификация"
                else:
                    raise ValueError("Неверный выбор шаблона.")

                save_path = MASSPRINT_SAVEPATH_DEFAULT
                save_path_obj = Path(save_path)
                if not save_path_obj.exists():
                    save_path_obj.mkdir(parents=True, exist_ok=True)
                    logger.info(f"Создана папка для сохранения: {save_path_obj}")

                template_folder = Path(TEMPLATES_FOLDER_DEFAULT)
                if not template_folder.exists():
                    logger.error(f"Папка с шаблонами не найдена: {template_folder}")
                    continue

                # Проверка существования файла шаблона
                template_exists = False
                found_ext = None
                for ext in [".doc", ".docx"]:
                    template_path = template_folder / f"{template_name}{ext}"
                    if template_path.exists():
                        template_exists = True
                        found_ext = ext
                        break
                if not template_exists:
                    logger.error(f"Файл шаблона '{template_name}' (с расширением .doc или .docx) не найден в папке {template_folder}")
                    continue
                logger.info(f"Шаблон найден: {template_name}{found_ext}")

                output_path = copy_template(card_data, template_name, save_path, template_folder)
                logger.info(f"Документ создан: {output_path}")

                replace_in_doc(output_path, replacements)

                print_document(output_path)

                logger.info(f"ФАЙЛ {index}/{len(excel_files)} УСПЕШНО ОБРАБОТАН: {os.path.basename(card_path)}")

            except Exception as e:
                logger.error(f"ОШИБКА при обработке файла {card_path}: {e}", exc_info=True)
            finally:
                if output_path:
                    logger.info(f"Итоговый документ: {output_path}")

        logger.info(f"=== ЗАВЕРШЕНИЕ ОБРАБОТКИ ВСЕХ ФАЙЛОВ ===")

    def choose_template():
        """Выбор типа шаблона пользователем."""
        print("Распечатать постановления по:")
        print("1 - родству")
        print("2 - прямой идентификации")
        choice = input("Введите номер шаблона: ")
        try:
            choice = int(choice)
            if choice not in [1, 2]:
                raise ValueError
            return choice
        except ValueError:
            print("Неверный выбор. Пожалуйста, введите 1 или 2.")
            return choose_template()

    # ---- Основная логика режима ----
    try:
        folder_path = Path.cwd()
        excel_files = get_excel_files(str(folder_path))
        if not excel_files:
            raise FileNotFoundError(f"В папке {folder_path} не найдено файлов Excel (.xlsx или .xls).")

        card_path = excel_files[0]
        card_data = read_card(card_path)
        logger.info(f"Обновленные данные из файла: {card_data}")

        template_choice = choose_template()
        process_multiple_cards(str(folder_path), template_choice)

    except Exception as e:
        logger.error(f"Произошла ошибка в режиме массовой печати: {e}", exc_info=True)


def process_rostov_cards(script_dir: Path) -> None:
    """
    Режим «Заведующая: заполнение Карт по ОБРАЗЦАМ ИЗ РОСТОВА».
    """
    logger.info("Запущен режим «Заведующая: заполнение Карт по ОБРАЗЦАМ ИЗ РОСТОВА»")

    default_input_dir = script_dir / ROSTOV_DIR_DEFAULT
    user_input = input(f"Введите путь к папке с исходным Excel-файлом (или Enter для '{default_input_dir}'): ").strip()
    input_dir = Path(user_input) if user_input else default_input_dir

    if not input_dir.exists():
        logger.error(f"Папка не найдена: {input_dir}")
        input("Нажмите Enter для завершения...")
        return
    logger.info(f"Папка с исходными данными: {input_dir}")

    default_output_dir = input_dir
    user_output = input(f"Введите путь для сохранения карт (или Enter для '{default_output_dir}'): ").strip()
    output_dir = Path(user_output) if user_output else default_output_dir

    try:
        output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Папка для сохранения карт: {output_dir}")
    except Exception as e:
        logger.error(f"Не удалось создать папку {output_dir}: {e}")
        return

    excel_files = list(input_dir.glob("*.xlsx")) + list(input_dir.glob("*.xlsm"))
    excel_files = [f for f in excel_files if not f.name.startswith('~$')]
    if not excel_files:
        logger.error(f"В папке {input_dir} не найдено файлов Excel.")
        return

    if len(excel_files) == 1:
        source_file = excel_files[0]
        logger.info(f"Найден файл: {source_file.name}")
    else:
        print("Найдено несколько файлов Excel. Выберите нужный:")
        for i, f in enumerate(excel_files, 1):
            print(f"{i}. {f.name}")
        while True:
            try:
                choice = int(input("Введите номер: ").strip())
                if 1 <= choice <= len(excel_files):
                    source_file = excel_files[choice-1]
                    break
                else:
                    print("Неверный номер.")
            except ValueError:
                print("Введите число.")
        logger.info(f"Выбран файл: {source_file.name}")

    template_path = find_template_excel(script_dir)
    if not template_path:
        logger.error("Не удалось найти шаблон карты (Excel-файл в папке программы).")
        return

    try:
        wb_source = openpyxl.load_workbook(source_file)
        ws_source = wb_source.active
    except Exception as e:
        logger.error(f"Не удалось открыть файл {source_file}: {e}")
        return

    # Определяем заголовки
    headers = {}
    for col in range(1, ws_source.max_column + 1):
        cell_value = ws_source.cell(row=1, column=col).value
        if cell_value:
            headers[cell_value.strip()] = col

    required_columns = [
        "Фамилия, Имя, Отчество родственника",
        "Дата рождения родственника",
        "Степень родства",
        "Фамилия, Имя, Отчество пропавшего",
        "Дата рождения пропавшего",
        "Дата отправки"
    ]
    for col_name in required_columns:
        if col_name not in headers:
            logger.error(f"В файле отсутствует обязательный столбец '{col_name}'. Обработка прервана.")
            return

    col_rel_fio = headers["Фамилия, Имя, Отчество родственника"]
    col_rel_birth = headers["Дата рождения родственника"]
    col_degree = headers["Степень родства"]
    col_missing_fio = headers["Фамилия, Имя, Отчество пропавшего"]
    col_missing_birth = headers["Дата рождения пропавшего"]
    col_send_date = headers["Дата отправки"]

    total_rows = 0
    success_rows = 0
    error_rows = 0

    card_filler = ExcelCardFiller(template_path)

    for row_idx in range(2, ws_source.max_row + 1):
        try:
            total_rows += 1
            logger.info(f"Обработка строки {row_idx}")

            rel_fio = normalize_text(ws_source.cell(row=row_idx, column=col_rel_fio).value)
            rel_birth = normalize_text(ws_source.cell(row=row_idx, column=col_rel_birth).value)
            degree = normalize_text(ws_source.cell(row=row_idx, column=col_degree).value)
            missing_fio = normalize_text(ws_source.cell(row=row_idx, column=col_missing_fio).value)
            missing_birth = normalize_text(ws_source.cell(row=row_idx, column=col_missing_birth).value)
            send_date_str = normalize_text(ws_source.cell(row=row_idx, column=col_send_date).value)

            if not degree:
                logger.warning(f"Строка {row_idx} не содержит степени родства, пропущена")
                continue

            send_day = ""
            send_month = ""
            if send_date_str is not None and send_date_str != "":
                try:
                    send_dt = parse_date_flexible(send_date_str)
                    if send_dt:
                        send_day = str(send_dt.day)
                        send_month = month_num_to_genitive(send_dt.month)
                        logger.debug(f"Дата отправки распознана: {send_dt.strftime('%Y-%m-%d')} -> день={send_day}, месяц={send_month}")
                    else:
                        logger.warning(f"Не удалось распознать дату отправки: {send_date_str}")
                except Exception as e:
                    logger.warning(f"Ошибка при разборе даты отправки: {e}")

            current_day, current_month = get_current_date()

            normalized_degree = normalize_degree(degree)
            is_direct = normalized_degree == 'личный генотип'

            if is_direct:
                code1_value = "СВО_Ростов_образец_прямая идентификация"
                fio_to_parse = missing_fio
                birth_to_parse = missing_birth
            else:
                code1_value = "СВО_Ростов_образец_родственники"
                fio_to_parse = rel_fio
                birth_to_parse = rel_birth

            surname, first_name, patronymic = parse_full_name(fio_to_parse)

            surname = add_dot_to_initials(surname)
            first_name = add_dot_to_initials(first_name)
            patronymic = add_dot_to_initials(patronymic)

            surname_gen = card_filler.convert_to_genitive(surname) if surname else ""
            first_name_gen = card_filler.convert_to_genitive(first_name) if first_name else ""
            patronymic_gen = card_filler.convert_to_genitive(patronymic) if patronymic else ""

            file_name_raw = rel_fio if rel_fio else missing_fio
            safe_name = sanitize_filename(file_name_raw)
            card_name = f"{safe_name}.xlsx"
            card_path = output_dir / card_name

            try:
                shutil.copy2(template_path, card_path)
                logger.info(f"Создана карта: {card_path}")
            except Exception as e:
                logger.error(f"Ошибка при копировании шаблона: {e}")
                error_rows += 1
                continue

            wb_card = openpyxl.load_workbook(card_path)
            ws_card = wb_card.active

            fill_data = {
                "1": code1_value,
                "ДП": send_day,
                "МП": send_month,
                "ДН": str(current_day),
                "МН": current_month,
                "ТЕР": "Ростова-на-Дону",
            }

            if is_direct:
                fill_data["ФИ"] = surname
                fill_data["ИИ"] = first_name
                fill_data["ОИ"] = patronymic
                fill_data["ФПИ"] = surname
                fill_data["ИПИ"] = first_name
                fill_data["ОПИ"] = patronymic
                fill_data["ФР"] = surname_gen
                fill_data["ИР"] = first_name_gen
                fill_data["ОР"] = patronymic_gen
                fill_data["ДР"] = missing_birth
                fill_data["ДРП"] = missing_birth
            else:
                fill_data["ФИ"] = surname
                fill_data["ИИ"] = first_name
                fill_data["ОИ"] = patronymic
                fill_data["ФР"] = surname_gen
                fill_data["ИР"] = first_name_gen
                fill_data["ОР"] = patronymic_gen
                fill_data["ДР"] = birth_to_parse

                missing_surname, missing_first, missing_patr = parse_full_name(missing_fio)

                missing_surname = add_dot_to_initials(missing_surname)
                missing_first = add_dot_to_initials(missing_first)
                missing_patr = add_dot_to_initials(missing_patr)

                fill_data["ФПИ"] = missing_surname
                fill_data["ИПИ"] = missing_first
                fill_data["ОПИ"] = missing_patr
                fill_data["ДРП"] = missing_birth

                degree_key = normalized_degree
                rod_value = ROD_MAPPING.get(degree_key, "родство не указано")
                fill_data["РОД"] = rod_value

            # Ячейки для очистки
            clear_codes = ["ФОР", "ИОР", "ООР", "ДРО", "РОД1", "ИНД", "АТО", "ВЧ", "ЛН", "ДБИ"]
            if not is_direct:
                # Для родства РОД оставляем, остальные очищаем
                for code in clear_codes:
                    fill_data[code] = None
            else:
                # Для прямой идентификации очищаем и РОД тоже
                clear_codes.append("РОД")
                for code in clear_codes:
                    fill_data[code] = None

            card_filler.fill_card(wb_card, fill_data)
            wb_card.save(card_path)
            logger.info(f"Карта {card_path} успешно заполнена")
            success_rows += 1

        except Exception as e:
            logger.error(f"Ошибка при обработке строки {row_idx}: {e}", exc_info=True)
            error_rows += 1

    logger.info("=" * 50)
    logger.info(f"Обработка завершена. Всего строк: {total_rows}, успешно: {success_rows}, ошибок: {error_rows}")
    logger.info("=" * 50)


# ========== ОСНОВНАЯ ФУНКЦИЯ ==========
def main():
    logger.info("Программа запущена")
    try:
        script_dir = Path(__file__).parent
        work_type = select_work_type()

        if work_type == "tobolsk_newborn":
            template_path = find_template_excel(script_dir)
            if not template_path:
                logger.error("Не удалось найти шаблон карты")
                input("Нажмите Enter для завершения...")
                return
            try:
                start_number = int(input("Введите начальный номер заключения: ").strip())
            except ValueError:
                logger.error("Номер заключения должен быть целым числом")
                input("Нажмите Enter для завершения...")
                return
            process_tobolsk_newborn(script_dir, template_path, start_number)

        elif work_type == "bone_cards":
            process_bone_cards(script_dir)

        elif work_type == "registry_cards":
            process_registry_cards(script_dir)

        elif work_type == "distribution_cards":
            process_distribution_cards(script_dir)

        elif work_type == "mass_fill_samples":
            process_mass_fill_cards_with_samples(script_dir)

        elif work_type == "mass_print_postanovleniy":
            process_mass_print_postanovleniy()

        elif work_type == "rostov_cards":
            process_rostov_cards(script_dir)

        else:
            logger.error(f"Неизвестный тип работы: {work_type}")

        logger.info("Обработка завершена")
        input("Нажмите Enter для завершения...")

    except KeyboardInterrupt:
        logger.info("Программа прервана пользователем")
        input("Нажмите Enter для завершения...")
    except Exception as e:
        logger.error(f"Критическая ошибка: {e}", exc_info=True)
        input("Нажмите Enter для завершения...")


if __name__ == "__main__":
    main()